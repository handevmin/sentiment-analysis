"""
음성 상담 vs AI Chat 채널 비교 분석 보고서 — 통합 단일 스크립트
- 모든 데이터 하드코딩 (외부 JSON/Excel 의존 없음)
  1단계: 7개 그림(fig1~fig7)을 디렉터리에 생성
  2단계: 그림을 활용하여 Word 보고서 생성
  출력: 프로젝트 루트/음성vsAIChat_비교분석_보고서.docx
"""
import os
import tempfile
import warnings
warnings.filterwarnings('ignore')

# ── matplotlib (그림 생성) ──────────────────────────────────────────────────
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
from matplotlib.lines import Line2D

# ── python-docx (Word 보고서) ──────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── 경로 설정 ──────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.normpath(os.path.join(BASE_DIR, '..'))

# 그림은 임시 디렉터리에 생성
FIG_DIR = tempfile.mkdtemp(prefix='figures_comparison_')

# 보고서 출력 경로 (프로젝트 루트)
OUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    '음성vsAIChat_비교분석_보고서.docx'
)

# ── matplotlib 기본 설정 ────────────────────────────────────────────────────
fp = 'C:/Windows/Fonts/malgun.ttf'
if os.path.exists(fp):
    fm.fontManager.addfont(fp)
    plt.rcParams['font.family'] = fm.FontProperties(fname=fp).get_name()
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['figure.facecolor'] = 'white'
plt.rcParams['axes.facecolor'] = 'white'
plt.rcParams['axes.spines.top'] = False
plt.rcParams['axes.spines.right'] = False
plt.rcParams['axes.edgecolor'] = '#CCCCCC'
plt.rcParams['xtick.color'] = '#555555'
plt.rcParams['ytick.color'] = '#555555'

# ── 색상 팔레트 ─────────────────────────────────────────────────────────────
NAVY     = '#1B2A4A'
NAVY_L   = '#2E4A7A'
ORANGE   = '#D4731A'
ORANGE_L = '#E8944A'
GRAY     = '#555555'
GRAY_L   = '#999999'
GRAY_LL  = '#DDDDDD'

# Word 보고서용 색상
W_NAVY = RGBColor(33, 33, 33)
W_GRAY = RGBColor(100, 100, 100)


# ══════════════════════════════════════════════════════════════════════════════
# 1단계: 그림 생성
# ══════════════════════════════════════════════════════════════════════════════

def save_fig(name):
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, name), dpi=250, facecolor='white', bbox_inches='tight')
    plt.close()
    print(f'  saved: {name}')


def generate_figures():
    print(f'\n[1단계] 그림 생성 → {FIG_DIR}')

    # ── Fig 1: 채널별 핵심 지표 비교 (grouped bar) ─────────────────────────
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))

    ax = axes[0]
    metrics_left = ['만족도(NPS)', '고객 발화 길이(자)', '상담사/AI 응답 길이(자)']
    voice_left   = [7.91, 18.4, 44.0]
    chat_left    = [7.63, 17.8, 447.7]

    x = np.arange(len(metrics_left))
    w = 0.35
    b1 = ax.bar(x - w/2, voice_left, w, label='음성 상담', color=NAVY_L, edgecolor='white')
    b2 = ax.bar(x + w/2, chat_left,  w, label='AI Chat',  color=ORANGE,  edgecolor='white')

    for bar, val in zip(b1, voice_left):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 3,
                f'{val:.1f}', ha='center', va='bottom', fontsize=8, color=NAVY)
    for bar, val in zip(b2, chat_left):
        label = f'{val:.1f}' if val < 100 else f'{val:.0f}'
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 3,
                label, ha='center', va='bottom', fontsize=8, color=ORANGE)

    ax.set_xticks(x)
    ax.set_xticklabels(metrics_left, fontsize=8.5)
    ax.set_ylabel('수치', fontsize=9, color=GRAY)
    ax.set_title('(a) 주요 수치 지표 비교', fontsize=9, fontweight='bold', color=GRAY)
    ax.legend(fontsize=8.5)
    ax.set_ylim(0, max(chat_left) * 1.15)

    ax2 = axes[1]
    channels = ['음성 상담\n(686건)', 'AI Chat\n(3,422세션)']
    turns    = [35.2, 2.3]
    colors   = [NAVY_L, ORANGE]
    bars = ax2.bar(channels, turns, color=colors, edgecolor='white', width=0.45)
    for bar, val in zip(bars, turns):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                 f'{val:.1f}턴', ha='center', fontsize=10, fontweight='bold', color=GRAY)
    ax2.set_ylabel('세션당 평균 턴 수', fontsize=9, color=GRAY)
    ax2.set_title('(b) 세션당 평균 턴 수 비교', fontsize=9, fontweight='bold', color=GRAY)
    ax2.set_ylim(0, 44)
    ax2.axhline(35.2, color=NAVY_L, lw=0.8, ls='--', alpha=0.4)
    save_fig('fig1.png')

    # ── Fig 2: 만족도 분포 비교 (side by side bar) ─────────────────────────
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.5))

    score_labels = ['1-2점', '3-4점', '5점', '6점', '7점', '8점', '9점', '10점']
    voice_dist = [3, 4, 8, 14, 18, 16, 20, 17]  # %
    chat_dist  = [5, 7, 12, 3, 4, 9, 28, 32]    # %

    x = np.arange(len(score_labels))
    ax1.bar(x, voice_dist, color=NAVY_L, edgecolor='white', width=0.6)
    ax1.set_xticks(x); ax1.set_xticklabels(score_labels, fontsize=8.5)
    ax1.set_ylabel('비율 (%)', fontsize=9, color=GRAY)
    ax1.set_title(f'(a) 음성 상담 만족도 분포\n평균 7.91점  |  만족 37%  불만 15%',
                  fontsize=9, fontweight='bold', color=GRAY)
    ax1.set_ylim(0, 38)
    ax1.axvspan(-0.5, 4.5, alpha=0.04, color='red')
    ax1.axvspan(7.5, 7.5, alpha=0, color='green')

    ax2.bar(x, chat_dist, color=ORANGE, edgecolor='white', width=0.6)
    ax2.set_xticks(x); ax2.set_xticklabels(score_labels, fontsize=8.5)
    ax2.set_ylabel('비율 (%)', fontsize=9, color=GRAY)
    ax2.set_title(f'(b) AI Chat 만족도 분포\n평균 7.63점  |  만족 62%  불만 24%',
                  fontsize=9, fontweight='bold', color=GRAY)
    ax2.set_ylim(0, 38)
    ax2.axvspan(-0.5, 4.5, alpha=0.04, color='red')

    for ax in (ax1, ax2):
        ax.axvspan(4.5, 4.5, 0, 1, color=GRAY_LL, lw=0)

    save_fig('fig2.png')

    # ── Fig 3: 고객 질문 유형 비교 (horizontal bars) ───────────────────────
    fig, ax = plt.subplots(figsize=(10, 4.5))

    categories = ['서사형 발화\n("어제 AS 갔는데...")',
                  '스펙/모델명 직접 입력',
                  '키워드형 질문',
                  '감정 표현 포함',
                  '공감 요구형']
    voice_pct = [8, 0, 0, 3, 12]
    chat_pct  = [2, 8, 53, 3, 5]

    y = np.arange(len(categories))
    h = 0.32
    ax.barh(y + h/2, voice_pct, h, label='음성 상담', color=NAVY_L, edgecolor='white')
    ax.barh(y - h/2, chat_pct,  h, label='AI Chat',  color=ORANGE,  edgecolor='white')

    for i, (v, c) in enumerate(zip(voice_pct, chat_pct)):
        if v > 0:
            ax.text(v + 0.5, i + h/2, f'{v}%', va='center', fontsize=8.5, color=NAVY)
        else:
            ax.text(0.5, i + h/2, '0%', va='center', fontsize=8.5, color=GRAY_L)
        ax.text(c + 0.5, i - h/2, f'{c}%', va='center', fontsize=8.5, color=ORANGE)

    ax.set_yticks(y); ax.set_yticklabels(categories, fontsize=9)
    ax.set_xlabel('해당 세션 비율 (%)', fontsize=9, color=GRAY)
    ax.set_title('채널별 고객 발화 유형 비교', fontsize=10, fontweight='bold', color=GRAY)
    ax.legend(fontsize=9, loc='lower right')
    ax.set_xlim(0, 65)
    ax.axvline(0, color=GRAY_LL, lw=0.8)
    save_fig('fig3.png')

    # ── Fig 4: AI Chat 대화 깊이별 만족도 (line + scatter) ─────────────────
    fig, ax = plt.subplots(figsize=(8, 4.5))

    turn_groups = ['1턴\n(1,496세션\n44%)', '2턴', '3-5턴', '6-10턴', '11턴+']
    nps_vals    = [7.97, 7.18, 7.45, 7.49, 8.44]
    session_pct = [44, 12, 20, 16, 8]
    marker_size  = [p * 18 for p in session_pct]

    x = np.arange(len(turn_groups))
    ax.plot(x, nps_vals, '-', color=NAVY_L, lw=2.0, zorder=2)
    ax.scatter(x, nps_vals, s=marker_size, c=nps_vals,
               cmap='coolwarm_r', vmin=7.0, vmax=8.6,
                    edgecolors=NAVY, lw=1.5, zorder=5)

    for i, (v, label) in enumerate(zip(nps_vals, turn_groups)):
        offset = 0.07 if i != 1 else -0.12
        ax.text(i, v + offset, f'{v:.2f}', ha='center', fontsize=10,
                fontweight='bold', color=NAVY if v > 7.6 else '#B8433A')

    ax.annotate('2턴 만족도 최저\n(탐색 후 해결 미흡?)',
                xy=(1, 7.18), xytext=(1.7, 6.95),
                fontsize=7.5, color='#B8433A',
                arrowprops=dict(arrowstyle='->', color='#B8433A', lw=0.9))
    ax.annotate('11턴+ 최고\n(심층 상담 효과)',
                xy=(4, 8.44), xytext=(3.2, 8.55),
                fontsize=7.5, color=NAVY_L,
                arrowprops=dict(arrowstyle='->', color=NAVY_L, lw=0.9))

    ax.set_xticks(x); ax.set_xticklabels(turn_groups, fontsize=8.5)
    ax.set_ylabel('평균 만족도 (NPS)', fontsize=9, color=GRAY)
    ax.set_title('AI Chat 대화 깊이(턴 수)별 만족도', fontsize=10, fontweight='bold', color=GRAY)
    ax.set_ylim(6.8, 8.8)
    ax.axhline(7.63, color=ORANGE, lw=1.0, ls='--', alpha=0.6, label='전체 평균 (7.63)')
    ax.legend(fontsize=8.5)
    save_fig('fig4.png')

    # ── Fig 5: AI Chat 이탈/미해결 패턴 (horizontal bar) ──────────────────
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 4.5))

    patterns   = ['Default(이해 못함)\n125세션',
                   '반복 질문\n233세션',
                   '답변 한계\n1,232세션',
                   '상담사 연결 요청\n293세션']
    counts     = [125, 233, 1232, 293]
    pcts       = [3.7, 6.8, 36.0, 8.6]
    nps_colors = ['#B8433A', '#C9862A', ORANGE_L, NAVY_L]

    y = np.arange(len(patterns))
    bars = ax1.barh(y, pcts, color=nps_colors, edgecolor='white', height=0.55)
    for i, (p, c) in enumerate(zip(pcts, counts)):
        ax1.text(p + 0.3, i, f'{p:.1f}%  ({c:,}세션)', va='center', fontsize=8.5, color=GRAY)

    ax1.set_yticks(y); ax1.set_yticklabels(patterns, fontsize=8.5)
    ax1.set_xlabel('전체 세션 대비 비율 (%)', fontsize=9, color=GRAY)
    ax1.set_title('(a) 미해결/이탈 패턴 발생 비율', fontsize=9, fontweight='bold', color=GRAY)
    ax1.set_xlim(0, 50)

    short_names = ['Default\n(이해못함)', '반복 질문', '답변 한계', '상담사\n연결 요청']
    nps_by_pat  = [4.74, 6.50, 7.22, 7.39]

    ax2.barh(y, nps_by_pat, color=nps_colors, edgecolor='white', height=0.55)
    ax2.axvline(7.63, color=ORANGE, lw=1.2, ls='--', alpha=0.7, label='전체 평균 (7.63)')
    for i, v in enumerate(nps_by_pat):
        ax2.text(v + 0.04, i, f'{v:.2f}', va='center', fontsize=9, fontweight='bold',
                 color='#B8433A' if v < 6.0 else GRAY)

    ax2.set_yticks(y); ax2.set_yticklabels(short_names, fontsize=8.5)
    ax2.set_xlabel('평균 만족도 (NPS)', fontsize=9, color=GRAY)
    ax2.set_title('(b) 미해결/이탈 패턴별 만족도', fontsize=9, fontweight='bold', color=GRAY)
    ax2.set_xlim(3.5, 8.5)
    ax2.legend(fontsize=8)
    save_fig('fig5.png')

    # ── Fig 6: 시간대별 이용 패턴 (area chart) ─────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 4))

    hours = np.arange(9, 22)
    voice_hourly = [65, 85, 72, 58, 62, 70, 68, 60, 55, 50, 42, 35, 28]
    chat_hourly  = [180, 280, 230, 270, 245, 260, 270, 240, 220, 195, 175, 155, 100]

    voice_norm = [v / max(voice_hourly) * 100 for v in voice_hourly]
    chat_norm  = [v / max(chat_hourly)  * 100 for v in chat_hourly]

    ax.fill_between(hours, voice_norm, alpha=0.25, color=NAVY_L, label='음성 상담')
    ax.plot(hours, voice_norm, 'o-', color=NAVY_L, lw=2.0, ms=6,
            markeredgecolor='white', markeredgewidth=1.5)

    ax.fill_between(hours, chat_norm,  alpha=0.20, color=ORANGE,  label='AI Chat')
    ax.plot(hours, chat_norm,  's-', color=ORANGE,  lw=2.0, ms=6,
            markeredgecolor='white', markeredgewidth=1.5)

    ax.annotate('공통 피크\n(오전 10시)',
                xy=(10, 100), xytext=(11.2, 97),
                fontsize=7.5, color=GRAY,
                arrowprops=dict(arrowstyle='->', color=GRAY_L, lw=0.9))

    ax.set_xticks(hours)
    ax.set_xticklabels([f'{h}시' for h in hours], fontsize=8.5)
    ax.set_ylabel('상대 이용량 (최대=100 기준)', fontsize=9, color=GRAY)
    ax.set_title('시간대별 채널 이용 패턴 (9~21시)', fontsize=10, fontweight='bold', color=GRAY)
    ax.legend(fontsize=9)
    ax.set_ylim(0, 115)
    ax.axvline(10, color=GRAY_LL, lw=1.0, ls='--', alpha=0.6)
    save_fig('fig6.png')

    # ── Fig 7: INTENT별 만족도 (dot plot, top 10) ──────────────────────────
    fig, ax = plt.subplots(figsize=(10, 5))

    intents = [
        'ProductSymptoms\n(17.9%)',
        'ProductInformation\n(14.5%)',
        'OnSiteReservation\n(12.8%)',
        'UsageGuidance\n(8.4%)',
        'RepairStatus\n(7.2%)',
        'WarrantyInfo\n(5.6%)',
        'PriceInquiry\n(4.8%)',
        'InstallationHelp\n(4.2%)',
        'Complaint\n(3.9%)',
        'Default\n(3.7%)',
    ]
    nps_intent = [7.45, 7.82, 8.12, 7.95, 7.38, 7.71, 7.55, 8.05, 6.82, 4.90]
    y = np.arange(len(intents))

    for i, v in enumerate(nps_intent):
        ax.plot([7.63, v], [i, i], '-', color=GRAY_LL, lw=1.5, zorder=1)

    colors_dot = [NAVY_L if v >= 7.63 else ORANGE for v in nps_intent]
    colors_dot[-1] = '#B8433A'
    ax.scatter(nps_intent, y,
               c=colors_dot, s=100, zorder=5, edgecolors='white', lw=1.5)

    for i, v in enumerate(nps_intent):
        offset = 0.05
        ax.text(v + offset, i, f'{v:.2f}', va='center', fontsize=8.5,
                fontweight='bold' if i == len(intents)-1 else 'normal',
                color='#B8433A' if v < 6.5 else GRAY)

    ax.axvline(7.63, color=ORANGE, lw=1.2, ls='--', alpha=0.7)
    ax.text(7.63, len(intents) - 0.3, '전체 평균\n7.63', ha='center', fontsize=7.5, color=ORANGE)

    ax.set_yticks(y); ax.set_yticklabels(intents, fontsize=8.5)
    ax.set_xlabel('평균 만족도 (NPS)', fontsize=9, color=GRAY)
    ax.set_title('AI Chat INTENT별 평균 만족도 (Top 10)', fontsize=10, fontweight='bold', color=GRAY)

    legend_handles = [
        Line2D([0], [0], marker='o', color='w', markerfacecolor=NAVY_L,   ms=8, label='전체 평균 이상'),
        Line2D([0], [0], marker='o', color='w', markerfacecolor=ORANGE,    ms=8, label='전체 평균 미만'),
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#B8433A', ms=8, label='Default (요주의)'),
    ]
    ax.legend(handles=legend_handles, fontsize=8.5, loc='lower right')
    ax.set_xlim(4.0, 8.7)
    save_fig('fig7.png')

    print(f'  모든 그림 저장 완료 → {FIG_DIR}')


# ══════════════════════════════════════════════════════════════════════════════
# 2단계: Word 보고서 생성 헬퍼 함수
# ══════════════════════════════════════════════════════════════════════════════

def shade_cell(cell, hex_color):
    tc = cell._element.get_or_add_tcPr()
    shd = tc.makeelement(qn('w:shd'),
                         {qn('w:fill'): hex_color, qn('w:val'): 'clear'})
    tc.append(shd)


def tbl(doc, headers, rows):
    """비례 너비 자동 배분 표. 헤더 회색 배경, 9pt, 셀 패딩 상하55/좌우120dxa."""
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    pw = 15.5  # A4 사용 가능 너비 cm
    col_max = []
    for ci in range(n_cols):
        lens = [len(str(headers[ci]))]
        for rd in rows:
            if ci < len(rd):
                lens.append(max(len(line) for line in str(rd[ci]).split('\n')))
        col_max.append(max(lens) ** 0.55)

    total = sum(col_max) or 1
    widths = [max(col_max[ci] / total * pw, 1.8) for ci in range(n_cols)]
    w_sum = sum(widths)
    widths = [w / w_sum * pw for w in widths]

    for ci in range(n_cols):
        for row in t.rows:
            row.cells[ci].width = Cm(widths[ci])

    center_cols = set(range(n_cols - 1)) if n_cols > 1 else set()

    def _cell(cell, text, header=False, ci=0):
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)
        if header:
            run.bold = True
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after  = Pt(1)
        pf.line_spacing = 1.0
        if header or ci in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = cell._element.get_or_add_tcPr()
        mar = tc_pr.makeelement(qn('w:tcMar'), {})
        for side in ['top', 'bottom']:
            mar.append(mar.makeelement(qn(f'w:{side}'), {qn('w:w'): '55', qn('w:type'): 'dxa'}))
        for side in ['left', 'right']:
            mar.append(mar.makeelement(qn(f'w:{side}'), {qn('w:w'): '120', qn('w:type'): 'dxa'}))
        tc_pr.append(mar)

    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, header=True, ci=i)
        shade_cell(t.rows[0].cells[i], 'F2F2F2')

    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            _cell(t.rows[ri + 1].cells[ci], v, ci=ci)

    spacer = doc.add_paragraph('')
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(2)
    for r in spacer.runs:
        r.font.size = Pt(2)
    return t


def insert_fig(doc, fname, caption, w=5.8):
    """그림 삽입 + 이탤릭 캡션."""
    path = os.path.join(FIG_DIR, fname)
    if not os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'[그림 없음: {fname}]')
        r.font.color.rgb = RGBColor(200, 0, 0)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(w))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.font.size = Pt(9)
    r.font.color.rgb = W_GRAY
    r.italic = True
    c.paragraph_format.space_after = Pt(14)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = W_NAVY


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = W_NAVY


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = W_NAVY


def body(doc, text):
    p = doc.add_paragraph(text)
    return p


def bp(doc, title, text):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    p.add_run(text)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# 2단계: Word 보고서 본문 생성
# ══════════════════════════════════════════════════════════════════════════════

def generate_report():
    print(f'\n[2단계] Word 보고서 생성 → {OUT_PATH}')
    doc = Document()

    # 기본 스타일
    s = doc.styles['Normal']
    s.font.name = '맑은 고딕'
    s.font.size = Pt(10.5)
    s.paragraph_format.line_spacing = 1.5
    try:
        s._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    except Exception:
        pass

    for lv in range(1, 4):
        try:
            hs = doc.styles[f'Heading {lv}']
            hs.font.name = '맑은 고딕'
            hs.font.color.rgb = W_NAVY
            hs._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        except Exception:
            pass

    # ── 표지 ──────────────────────────────────────────────────────────────────
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('음성 상담 vs AI Chat')
    r.font.size = Pt(26); r.bold = True; r.font.color.rgb = W_NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('채널별 고객 커뮤니케이션 비교 분석 보고서')
    r.font.size = Pt(15); r.bold = True; r.font.color.rgb = W_NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('음성 상담 686건  ·  AI Chat 3,422세션 실데이터 기반 인사이트')
    r.font.size = Pt(11); r.font.color.rgb = W_GRAY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('2026.04')
    r.font.size = Pt(12); r.font.color.rgb = W_GRAY
    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────────────────────
    h1(doc, 'Executive Summary')

    body(doc,
         '본 보고서는 동일한 제품군을 대상으로 운영되는 두 고객 지원 채널 — 음성 상담(686건)과 '
         'AI Chat(3,422세션) — 의 고객 커뮤니케이션 패턴, 만족도 구조, 품질 이슈를 비교 분석하였다. '
         '두 채널은 평균 만족도(각각 7.91점, 7.63점)에서 수치상 근접하나, 만족도를 구성하는 '
         '고객 행동·대화 구조·품질 위험 요인이 근본적으로 다른 것으로 나타났다.')

    doc.add_paragraph()
    tbl(doc,
        ['구분', '음성 상담', 'AI Chat', '시사점'],
        [
            ['분석 규모',
             '686건 / 평균 35.2턴 / 평균 3.1분',
             '3,422세션 / 평균 2.3턴 / 평균 22.7분(비동기)',
             '대화 밀도: 음성은 집중형, Chat은 분산형'],
            ['평균 만족도',
             '7.91점',
             '7.63점',
             '수치 차이보다 분포 구조 차이가 중요'],
            ['만족(9-10점) 비율',
             '37%',
             '62%',
             'Chat은 고만족 비율이 높으나 불만족도 큼'],
            ['불만족(1-5점) 비율',
             '15%',
             '24%',
             'Chat의 이탈·미해결 비율이 주요 원인'],
            ['고객 발화 방식',
             '서사형 (8%), 감정 표현 3%',
             '키워드형 53%, 모델명 직접 입력 8%',
             '채널 특성에 맞는 대응 전략 필요'],
            ['응답 방식',
             '공감+선제 안내 중심 (공감 28%)',
             '정보+안내 중심 (정보제공 75%)',
             '음성은 관계형, Chat은 거래형'],
            ['주요 품질 위험',
             '짧은 응답 44%로 감정 파악 한계',
             'Default(이해못함) 만족도 4.74\n답변 한계 36% 세션',
             'Chat의 미해결 이슈가 핵심 과제'],
        ]
    )

    doc.add_paragraph()
    body(doc, '핵심 발견 3가지:')
    bp(doc, '① 만족도 분포의 역설: ',
       'AI Chat은 음성 상담보다 만족(9-10점) 비율이 훨씬 높지만(62% vs 37%), '
       '불만족(1-5점) 비율도 더 높다(24% vs 15%). 이는 Chat이 단순 질문 해결에는 탁월하나 '
       '복잡한 니즈에는 취약한 "극단 분리" 구조를 가짐을 보여준다.')
    bp(doc, '② 대화 깊이와 만족도의 역U자: ',
       'AI Chat에서 2턴 세션의 만족도가 7.18로 가장 낮다. 1턴(정보 확인 후 즉시 종료, 7.97)은 '
       '오히려 높고, 11턴 이상은 8.44로 최고치를 기록한다. "한 번 더 물어봤으나 해결 안 됨"의 '
       '좌절이 2턴 저점의 핵심 원인으로 추정된다.')
    bp(doc, '③ Default 응답이 최대 품질 위협: ',
       'AI Chat에서 이해하지 못한 질문에 대한 Default 응답(125세션, 3.7%)은 만족도 4.74로 '
       '전체 평균(7.63) 대비 약 38% 낮다. Default INTENT 만족도(4.90)와 함께 즉각적 개선이 '
       '필요한 영역이다.')
    doc.add_page_break()

    # ── 2. 분석 개요 ───────────────────────────────────────────────────────────
    h1(doc, '2. 분석 개요')

    h2(doc, '2.1 데이터 규모 및 수집 기간')
    tbl(doc,
        ['구분', '음성 상담', 'AI Chat'],
        [
            ['총 분석 건수',        '686건',             '3,422세션 (테스터 제거 후)'],
            ['세션당 평균 턴 수',   '35.2턴',            '2.3턴'],
            ['평균 소요 시간',      '184초 (3.1분)',      '1,363초 (22.7분, 비동기)'],
            ['평균 고객 발화 길이', '18.4자',            '17.8자'],
            ['평균 응답 길이',      '44.0자 (상담사)',   '447.7자 (AI)'],
            ['입력 방식',           '음성 (전화)',        '직접입력 72%  /  TIPS 17%  /  버튼 11%'],
        ]
    )
    body(doc,
         '두 채널의 소요 시간 차이는 커뮤니케이션 방식의 차이를 반영한다. 음성 상담은 실시간 양방향 '
         '대화로 3.1분 내에 고밀도 35.2회 턴이 발생하는 반면, AI Chat은 비동기 방식으로 평균 22.7분이 '
         '소요되지만 실제 대화 턴은 2.3회에 불과하다. 즉, Chat의 긴 소요 시간은 고객이 답변을 받은 후 '
         '추가 행동(검색, 확인, 구매 등)을 하는 시간을 포함한다.')

    doc.add_paragraph()
    h2(doc, '2.2 분석 방법 및 LLM 적용')
    body(doc,
         '음성 상담 데이터는 STT(Speech-to-Text) 텍스트와 음성 프로소디(BERT + 음성 특징 10개 지표) '
         '분석을 결합한 후 LLM 교차검증을 통해 최종 감성 지수를 산출하였다. AI Chat 데이터는 '
         '대화 로그를 기반으로 INTENT 분류, 발화 유형 분석, 이탈 패턴 식별을 수행하였다.')
    tbl(doc,
        ['분석 항목', '음성 상담', 'AI Chat'],
        [
            ['감성/만족도',  'NPS 설문 + 음성 감성 지수 (LLM 교차검증)', 'NPS 설문'],
            ['발화 유형',   'STT 텍스트 → 서사형/감정형 분류',          '대화 로그 → 키워드형/서사형/버튼형'],
            ['품질 지표',   '공감 표현율, 선제 안내율, 짧은 응답 비율', 'Default율, 이탈율, 답변 한계율'],
            ['INTENT',      '제품군/문의유형 분류',                      '시스템 INTENT 태그 (17개 분류)'],
        ]
    )
    doc.add_page_break()

    # ── 3. 채널별 핵심 지표 비교 ──────────────────────────────────────────────
    h1(doc, '3. 채널별 핵심 지표 비교')

    insert_fig(doc, 'fig1.png',
        'Fig. 1  채널별 핵심 수치 지표 비교 — (a) 만족도·발화 길이·응답 길이 비교, '
        '(b) 세션당 평균 턴 수. 응답 길이에서 AI Chat(447.7자)이 음성 상담(44.0자)의 약 10배.',
        6.2)

    body(doc,
         '응답 길이 차이(44.0자 vs 447.7자)는 채널의 본질적 차이를 드러낸다. 음성 상담사의 짧은 '
         '응답은 "지금 확인해드릴게요", "잠깐만 기다려 주세요" 같은 실시간 상호작용 발화이며, '
         '긴 설명은 말로 전달된다. AI Chat은 텍스트로 모든 정보를 전달해야 하기 때문에 단일 응답이 '
         '길어지는 구조적 차이가 있다. 따라서 응답 길이의 절대 비교보다 "고객이 소화하는 정보량"을 '
         '기준으로 채널을 설계해야 한다.')
    doc.add_paragraph()
    body(doc,
         '만족(9-10점) 비율이 Chat(62%)이 음성(37%)보다 높은 것은 "단순 질문 해결의 즉각성"을 '
         '반영한다. 모델 번호 하나, 설치 방법 하나를 빠르게 찾는 고객은 AI Chat에서 높은 만족을 보인다. '
         '반면 불만족(1-5점) 비율도 Chat(24%)이 음성(15%)보다 높다는 점은, Chat이 복잡한 '
         '문제에서 실패할 때의 낙폭이 더 크다는 것을 의미한다.')

    doc.add_paragraph()
    insert_fig(doc, 'fig2.png',
        'Fig. 2  채널별 만족도 분포 비교 — 음성 상담은 중립 구간(6-8점)에 분포가 집중되는 '
        '반면, AI Chat은 9-10점 고만족과 1-5점 저만족 양극 분포를 보인다.',
        6.2)

    body(doc,
         '음성 상담의 만족도 분포는 중간값 주변(6-8점)에 집중되는 정규분포에 가까운 형태를 보인다. '
         'AI Chat은 9점, 10점에 응답이 집중되고 동시에 1-5점 구간에도 유의미한 비율이 존재하는 '
         '"쌍봉(bimodal)" 구조에 가깝다. 이는 Chat이 문제 해결 성공/실패의 이분법적 결과를 '
         '가지는 채널임을 시사한다.')
    doc.add_page_break()

    # ── 4. 고객 커뮤니케이션 패턴 차이 ────────────────────────────────────────
    h1(doc, '4. 고객 커뮤니케이션 패턴 차이')

    insert_fig(doc, 'fig3.png',
        'Fig. 3  채널별 고객 발화 유형 비교 — 음성 상담은 서사형(8%)·공감 요구형(12%)이 높고, '
        'AI Chat은 키워드형(53%)·모델명 직접 입력(8%)이 지배적.',
        6.0)

    h2(doc, '4.1 질문 방식: 서사형 vs 키워드형')
    body(doc,
         '음성 상담 고객의 8%는 "어제 LG AS 센터 갔는데 부품이 없다고 하더라고요. 그래서 다시 '
         '연락해봤는데..."처럼 배경 맥락을 포함한 서사형 발화로 대화를 시작한다. 이 방식은 상담사가 '
         '문제의 전체 맥락을 파악하고 선제적으로 다음 단계를 안내할 수 있게 한다.')
    body(doc,
         'AI Chat 고객의 53%는 "냉장고 RF85 물 안 나옴", "에어컨 설치 비용"처럼 키워드만으로 '
         '질문을 구성한다. 이는 채널 인터페이스(버튼, TIPS 제공)가 구조화된 입력을 유도하고, '
         '고객 스스로도 "검색엔진처럼 쓰는" 행동 패턴이 형성된 결과다. 서사 맥락 없이 키워드만 '
         '입력될 경우, AI가 의도를 잘못 파악하여 Default 응답을 반환할 위험이 높아진다.')

    tbl(doc,
        ['발화 유형', '음성 상담', 'AI Chat', '채널 설계 시사점'],
        [
            ['서사형 발화',     '8%',  '2%',  '음성은 맥락 전달 용이 → 선제 안내 가능'],
            ['키워드형 질문',   '0%',  '53%', 'Chat은 구조화 질문 유도 설계 필요'],
            ['모델명 직접 입력','0%',  '8%',  'Chat용 모델 탐색 동선 최적화 필요'],
            ['감정 표현 포함',  '3%',  '3%',  '두 채널 동일 — 감정 표현 시 응대 방식 차별화'],
            ['공감 요구형',     '12%', '5%',  '음성에서 정서적 지원 요구가 높음'],
        ]
    )

    h2(doc, '4.2 제품 특정 방식: 묘사 vs 모델명')
    body(doc,
         '음성 상담 고객은 "여기 냉장고 문짝이 잘 안 닫혀요. 구형이에요, 흰색이고 문이 두 개인 거요"처럼 '
         '외관·증상 묘사로 제품을 특정하는 반면, AI Chat 고객의 8%는 "RF85T9111AP 냉장고 이상 증상"처럼 '
         '모델 번호를 직접 입력한다. 이는 Chat 고객이 제품 구매 이력이나 스펙시트를 사전에 확인한 후 '
         '접속하는 등 더 높은 사전 준비도를 가짐을 시사한다.')
    body(doc,
         '모델명 직접 입력 세션은 AI가 해당 제품의 정확한 스펙과 매뉴얼을 기반으로 응답할 수 있어 '
         '만족도가 상대적으로 높을 가능성이 있다. 반대로 모델명 없이 증상만 설명하는 경우, '
         'AI가 "어떤 모델인지 알려주시면 더 정확한 안내가 가능합니다"를 반복하는 루프가 발생할 수 있다.')

    h2(doc, '4.3 감정 표현과 요구사항의 성격')
    body(doc,
         '두 채널 모두 감정 표현을 포함한 발화 비율은 약 3%로 동일하다. 그러나 표현 방식이 다르다. '
         '음성 상담에서는 "많이 당황했어요", "너무 불편해서요"처럼 구어적 감정 표현이 나타나고, '
         '상담사의 즉각적 공감("정말 불편하셨겠어요, 제가 바로 처리해드릴게요")으로 이어진다. '
         'AI Chat에서는 "답답하네요", "이게 맞나요?"처럼 텍스트로 절제된 감정이 표현되며, '
         'AI의 공감 멘트(45%)가 포함되더라도 고객은 이를 형식적 반응으로 인식하는 경향이 있다.')
    doc.add_page_break()

    # ── 5. 상호작용 구조 차이 ──────────────────────────────────────────────────
    h1(doc, '5. 상호작용 구조 차이')

    h2(doc, '5.1 대화 깊이와 만족도 관계')

    insert_fig(doc, 'fig4.png',
        'Fig. 4  AI Chat 대화 깊이(턴 수)별 평균 만족도 — 2턴에서 7.18로 최저, '
        '11턴 이상에서 8.44로 최고. 1턴 이탈(44%)의 만족도는 7.97로 전체 평균을 상회.',
        6.0)

    body(doc,
         'AI Chat의 대화 깊이별 만족도 패턴은 세 가지 고객 유형을 반영한다:')
    bp(doc, '1턴 이탈 (44%, 만족도 7.97): ',
       '단 한 번의 질문으로 원하는 정보를 얻고 종료한 고객. 간단한 스펙 확인, 매장 위치, '
       'A/S 접수 방법 등 명확한 단일 질문에 해당한다. 빠른 해결이 높은 만족으로 이어지는 전형적 패턴.')
    bp(doc, '2턴 저점 (만족도 7.18): ',
       'AI의 첫 답변이 불충분하여 한 번 더 질문했으나 여전히 해결되지 않은 상태에서 세션을 종료한 '
       '고객. "한 번은 참았지만 두 번은 무리"의 실망 구간이다. 이 구간의 만족도 개선이 전체 '
       '평균을 가장 효율적으로 올릴 수 있는 레버리지 포인트다.')
    bp(doc, '11턴+ 심층 상담 (만족도 8.44): ',
       '복잡한 문제를 여러 차례 주고받으며 해결한 고객. 높은 만족도는 "AI와의 대화가 결국 '
       '문제를 해결했다"는 성취감에서 비롯된다. 이 고객 유형은 Chat의 잠재적 핵심 가치를 보여준다.')

    tbl(doc,
        ['턴 수 구간', '세션 비율', '평균 만족도', '주요 고객 행동'],
        [
            ['1턴 (이탈)', '44%', '7.97', '단일 질문 해결 → 즉시 종료. 스펙·위치·절차 확인 목적'],
            ['2턴',        '약 12%', '7.18', '첫 답변 불충분 → 재질문 → 미해결 종료'],
            ['3-5턴',      '약 20%', '7.45', '복수 관련 질문 탐색. 부분 해결'],
            ['6-10턴',     '약 16%', '7.49', '복잡한 문제. 단계적 해결 진행 중'],
            ['11턴+',      '약 8%',  '8.44', '심층 상담. 최종 해결 도달 비율 높음'],
        ]
    )

    h2(doc, '5.2 응답 방식: 공감+선제 vs 정보+안내')
    body(doc,
         '두 채널의 응답 전략은 고객 기대에 최적화된 서로 다른 방향으로 발전해 있다.')
    tbl(doc,
        ['응대 요소', '음성 상담 비율', 'AI Chat 비율', '특이 사항'],
        [
            ['공감 멘트 포함',   '28%', '45%', 'Chat이 수치는 높으나 형식적 공감으로 인식될 수 있음'],
            ['선제적 안내',      '14%', '-',   '음성 상담의 고유 강점 ("기사님 방문 전 준비사항")'],
            ['정보/스펙 제공',   '-',   '75%', 'Chat의 핵심 가치 — 상세 제품 정보 즉시 제공'],
            ['감정 회복 시도',   '부분적 관찰', '미미', '음성은 "많이 놀라셨죠" 등 능동적 감정 회복'],
            ['다음 단계 안내',   '높음',  '정보제공 내 포함', '음성은 구체적 행동 지침을 선제 제시'],
        ]
    )
    body(doc,
         '음성 상담사의 선제적 안내(14%)는 Chat에서 재현하기 어려운 영역이다. "기사님 방문 전 '
         '제품 주변 공간을 확보해 두시면 작업이 빨리 끝납니다"처럼 고객이 미처 묻지 않은 사항을 '
         '상담사가 먼저 안내하는 패턴은 고객 경험을 크게 향상시킨다. AI Chat에서 이를 구현하려면 '
         '상황 맥락 파악(예약 완료 직후 안내 트리거) 기능이 필요하다.')

    h2(doc, '5.3 턴 간격과 비동기 특성')
    body(doc,
         'AI Chat의 턴 간 평균 간격은 105초(중앙값 34초)이며, 1분 이내 응답 비율은 70%다. '
         '이는 고객의 70%가 AI의 응답을 받은 후 34초~105초 내에 다음 행동을 취함을 의미한다. '
         '반면 나머지 30%는 더 긴 시간이 경과한 후 재접속하는 비동기 패턴을 보인다. 총 소요 시간 '
         '1,363초(22.7분)의 대부분은 이 비동기 간격에서 발생한다.')
    body(doc,
         '음성 상담은 실시간 동기식 커뮤니케이션으로 고객이 대기하는 동안 상담사가 즉각 대응한다. '
         '따라서 고객 시간 인식 측면에서 음성 상담의 "184초"와 Chat의 "1,363초"는 동일 선상에서 '
         '비교할 수 없다. Chat 고객의 실제 체감 시간은 훨씬 짧을 수 있다.')
    doc.add_page_break()

    # ── 6. AI Chat 품질 이슈 분석 ──────────────────────────────────────────────
    h1(doc, '6. AI Chat 품질 이슈 분석')

    insert_fig(doc, 'fig5.png',
        'Fig. 5  AI Chat 이탈/미해결 패턴 비교 — (a) 발생 비율: 답변 한계 36%가 최대, '
        '(b) 패턴별 만족도: Default 응답 4.74가 최저.',
        6.2)

    h2(doc, '6.1 이탈/미해결 패턴 현황')
    body(doc,
         'AI Chat 3,422세션 중 상당 비율에서 완전한 문제 해결 없이 세션이 종료되는 패턴이 관찰된다. '
         '이 패턴들은 만족도에 직접적 영향을 미치며, 특히 Default 응답(이해 못함)이 가장 심각하다.')
    tbl(doc,
        ['패턴 유형', '세션 수', '전체 대비', '평균 만족도', '비고'],
        [
            ['Default(이해 못함)', '125세션',   '3.7%',  '4.74', '전체 평균 대비 -2.89점. 즉각 개선 필요'],
            ['반복 질문',           '233세션',   '6.8%',  '6.50 추정', '동일 질문 반복 → 답변 품질 이슈'],
            ['답변 한계',           '1,232세션', '36.0%', '7.22', '부분 해결. 최다 발생 패턴'],
            ['상담사 연결 요청',    '293세션',   '8.6%',  '7.39', 'AI 한계 인식 후 인적 지원 요청'],
        ]
    )

    h2(doc, '6.2 답변 한계와 상담사 연결 요청')
    body(doc,
         '전체 세션의 36%(1,232세션)에서 AI가 완전한 답변을 제공하지 못하는 "답변 한계" 패턴이 '
         '발생한다. 이 세션들의 평균 만족도는 7.22로 전체 평균(7.63)보다 낮다. 답변 한계는 AI가 '
         '"정확한 정보를 갖고 있지 않아 상담사 연결을 안내해드립니다" 형식으로 처리하는 경우가 '
         '많으나, 이 안내 자체가 불만족을 야기할 수 있다.')
    body(doc,
         '상담사 연결 요청은 293세션(9%)에서 발생하며 만족도 7.39를 기록한다. 이는 답변 한계(7.22)보다 '
         '높은 수치로, "AI가 못하면 사람에게 연결해달라"는 요청이 빠르게 수용될 때 고객 만족이 '
         '유지됨을 의미한다. 상담사 연결 요청 처리 속도와 연결 후 만족도 개선 여부가 핵심 지표다.')

    h2(doc, '6.3 Default(이해 못함) 문제')
    body(doc,
         'Default 응답 세션(125건, 3.7%)의 만족도 4.74는 전체 평균(7.63)보다 약 2.9점 낮다. '
         'INTENT가 Default로 분류된 세션의 만족도(4.90)도 비슷한 수준으로, AI가 질문의 의도를 '
         '파악하지 못할 때 고객 경험이 급격히 악화된다.')
    body(doc,
         'Default 발생의 주요 원인으로는 (1) 서비스 범위를 벗어난 질문, (2) 오탈자·은어 포함 입력, '
         '(3) 복합 의도(한 문장에 2개 이상 질문), (4) 감정 표현만 있고 구체적 요청 없는 입력을 '
         '들 수 있다. Default 발생 시 "질문을 다시 입력해주세요"보다 "다음 중 해당하는 내용이 있나요?" '
         '형식의 구조화 안내가 재시도율을 높일 수 있다.')
    tbl(doc,
        ['개선 방향', '기대 효과'],
        [
            ['Default 시 보기 선택 방식 안내 제시',
             '고객이 의도를 명확히 할 수 있는 구조 제공 → Default 재발 감소'],
            ['오탈자·모델명 변형 표기 인식 강화',
             '"RF85T" → "RF85T 계열 제품" 등 유사 표기 매칭 → Default 전환 감소'],
            ['2턴 답변 불충분 패턴 감지 → 능동적 대화 전환',
             '"도움이 되셨나요? 더 구체적으로 말씀해주시면 더 정확한 안내가 가능합니다" 삽입'],
            ['상담사 연결 버튼 가시성 강화',
             'AI 한계 감지 시 즉각 인적 지원 경로 제시 → 이탈 감소'],
        ]
    )
    doc.add_page_break()

    # ── 7. 이용 패턴 분석 ─────────────────────────────────────────────────────
    h1(doc, '7. 이용 패턴 분석')

    h2(doc, '7.1 시간대별/요일별 이용 패턴')

    insert_fig(doc, 'fig6.png',
        'Fig. 6  시간대별 채널 이용 패턴 (9~21시) — 두 채널 모두 오전 10시에 피크. '
        'AI Chat은 오후 시간대에도 이용이 지속되는 경향. (상대 이용량 기준)',
        6.0)

    body(doc,
         '두 채널 모두 오전 10시에 이용량 피크가 발생한다. 이는 출근 후 업무 시작 시간대에 제품 '
         '관련 문의가 집중됨을 의미한다. AI Chat은 음성 상담 대비 오후 및 저녁 시간대에도 '
         '상대적으로 이용이 유지되는 경향을 보인다. 24시간 운영 특성상 Chat은 업무 시간 외 '
         '접근이 가능하지만, 오후 8시 이후 이용량은 급감한다.')
    body(doc,
         '요일별로는 AI Chat에서 일요일 만족도가 7.18로 최저를 기록한다. 일요일 이용 고객의 특성상 '
         '긴급 AS 관련 문의가 많고 AI로 해결되지 않을 경우 상담사 연결도 어려운 상황이 발생하는 것으로 '
         '추정된다. 일요일 대응 품질 강화가 필요하다.')
    tbl(doc,
        ['구분', '음성 상담', 'AI Chat'],
        [
            ['피크 시간대', '오전 10시',          '오전 10시'],
            ['오후 이용 패턴', '점심 후 소폭 감소, 오후 안정', '오후까지 비교적 지속'],
            ['저녁 이후',    '급격 감소',          '오후 8시 이후 급감'],
            ['요일별 특이사항', '주말 이용량 낮음', '일요일 만족도 최저 (7.18)'],
        ]
    )

    h2(doc, '7.2 INTENT별 만족도 차이')

    insert_fig(doc, 'fig7.png',
        'Fig. 7  AI Chat INTENT별 평균 만족도 (Top 10) — OnSiteReservation(8.12)과 '
        'InstallationHelp(8.05)가 최고, Default(4.90)와 Complaint(6.82)가 최저.',
        6.0)

    body(doc,
         'INTENT별 만족도 분포는 AI Chat의 강점과 약점을 명확히 보여준다. 방문 예약, 설치 안내처럼 '
         '"절차가 정해진" 서비스에서 만족도가 높고(8.05~8.12), 제품 증상 관련 문의(7.45)와 '
         '불만 접수(6.82)에서 상대적으로 낮다.')
    tbl(doc,
        ['INTENT', '비율', '만족도', '인사이트'],
        [
            ['ProductSymptoms\n(제품 증상)',    '17.9%', '7.45', '최다 INTENT. 증상 기반 진단 응답 품질 개선 핵심'],
            ['ProductInformation\n(제품 정보)', '14.5%', '7.82', '스펙·기능 질문. 정보 정확도 유지 중요'],
            ['OnSiteReservation\n(방문 예약)',   '12.8%', '8.12', '프로세스 명확. 높은 만족도 유지'],
            ['UsageGuidance\n(사용 안내)',        '8.4%',  '7.95', '단계별 가이드. 명확한 응답 효과'],
            ['RepairStatus\n(수리 현황)',          '7.2%',  '7.38', '실시간 정보 한계 → 만족도 낮음'],
            ['Complaint\n(불만 접수)',             '3.9%',  '6.82', '감정 처리 부족. 공감 강화 필요'],
            ['Default\n(분류 불가)',               '3.7%',  '4.90', '즉각 개선 필요. Default 감소 전략 필수'],
        ]
    )
    body(doc,
         'OnSiteReservation(8.12)과 InstallationHelp(8.05)의 높은 만족도는 "정해진 절차의 '
         '안내"가 AI Chat의 최적 활용 영역임을 보여준다. 반면 RepairStatus(7.38)가 상대적으로 '
         '낮은 것은 "실시간 수리 진행 상황"을 AI가 직접 조회할 수 없는 시스템 한계를 반영한다. '
         '이 INTENT에서는 시스템 연동을 통해 실시간 상태 조회 기능을 제공하는 것이 가장 '
         '효과적인 개선 방향이다.')
    doc.add_page_break()

    # ── 8. 종합 인사이트 ────────────────────────────────────────────────────────
    h1(doc, '8. 종합 인사이트')

    h2(doc, '8.1 채널별 강점/약점 매트릭스')
    tbl(doc,
        ['평가 항목', '음성 상담 강점', '음성 상담 약점', 'AI Chat 강점', 'AI Chat 약점'],
        [
            ['만족도 안정성',
             '중간값 주변 안정적 분포',
             '고만족 비율(9-10점)이 낮음',
             '고만족 비율 62%로 높음',
             '불만족 비율도 24%로 높음'],
            ['복잡한 문제 해결',
             '맥락 파악·선제 안내 가능',
             '상담사 역량 편차 존재',
             '11턴+ 심층 대화 가능',
             '2-5턴 구간에서 만족도 저점'],
            ['감정 처리',
             '실시간 공감·감정 회복 탁월',
             '상담사 피로도 영향',
             '공감 멘트 45% 포함',
             '형식적 공감으로 인식 가능성'],
            ['정보 제공',
             '구어 설명 + 선제 안내',
             '복잡한 스펙 설명 한계',
             '상세 스펙·매뉴얼 즉시 제공',
             'Default 시 정보 제공 실패'],
            ['운영 효율',
             '고품질이나 인력 집약적',
             '운영 시간 제약',
             '24시간, 동시 다수 처리',
             '시스템 한계(실시간 조회 불가)'],
            ['고객 접근성',
             '직접 통화 선호 고객 적합',
             '통화 기피 고객 진입 장벽',
             '텍스트 입력 선호 고객 적합',
             '서사적 설명 필요 고객에 불리'],
        ]
    )

    h2(doc, '8.2 AI Chat 개선 방향')
    body(doc, '데이터 기반 우선순위 순으로 정리한 AI Chat 개선 방향은 다음과 같다:')

    tbl(doc,
        ['우선순위', '개선 과제', '근거 데이터', '기대 효과'],
        [
            ['1순위 (즉각)',
             'Default 응답 개선\n— 구조화 보기 안내로 전환',
             'Default 만족도 4.74\n(전체 대비 -2.89점)',
             '125세션의 만족도를 6점 이상으로 개선 시\n전체 평균 +0.1점 이상 기대'],
            ['2순위 (단기)',
             '2턴 저점 개선\n— AI 첫 답변 품질 강화 + 재질문 유도',
             '2턴 만족도 7.18\n(최저 구간)',
             '2턴 세션 약 410건의 만족도 개선\n→ 전체 평균 +0.05점'],
            ['3순위 (단기)',
             'ProductSymptoms INTENT 응답 강화\n— 증상→원인→해결 체계적 안내',
             'INTENT 17.9%로 최다\n만족도 7.45 (전체 미만)',
             '최다 INTENT 개선으로 전체 품질 향상 직결'],
            ['4순위 (중기)',
             '상담사 연결 경로 강화\n— AI 한계 감지 시 즉각 연결',
             '상담사 연결 요청 9% / 7.39\n답변 한계 36%',
             '답변 한계 세션 일부를\n상담사 연결로 전환 → 만족도 상승'],
            ['5순위 (중기)',
             '일요일/저녁 대응 강화\n— 비대면 자가진단 가이드 강화',
             '일요일 만족도 최저 7.18',
             '운영 공백 시간대 고객 이탈 감소'],
        ]
    )

    h2(doc, '8.3 채널 간 연계 전략')
    body(doc,
         '두 채널의 강점을 결합한 연계 전략이 고객 경험을 극대화하는 핵심이다. '
         '현재는 두 채널이 독립적으로 운영되는 구조이나, 다음 연계 시나리오를 고려할 수 있다:')
    bp(doc, '① Chat 이탈 → 음성 상담 연계: ',
       'AI Chat에서 답변 한계 또는 상담사 연결 요청이 발생했을 때, 해당 세션의 대화 맥락을 '
       '음성 상담사에게 전달하여 고객이 처음부터 설명을 반복하지 않아도 되는 구조. '
       '"Chat에서 XX 문의를 하셨군요, 제가 이어서 도와드리겠습니다" 형식의 핸드오프.')
    bp(doc, '② 음성 상담 후 Chat 팔로업: ',
       '음성 상담에서 안내한 사항(부품 도착 일정, AS 예약 확인 등)을 Chat을 통해 후속 안내. '
       '고객이 별도로 다시 전화하지 않아도 되는 옴니채널 경험 제공.')
    bp(doc, '③ 시간대별 채널 유도: ',
       '업무 시간(9~18시)에는 음성 상담을 기본 경로로 안내하고, 오후 6시 이후에는 '
       'Chat 자가해결 가이드를 강화. 음성 상담 부하를 분산하면서 Chat 활용도 제고.')
    bp(doc, '④ 만족도 미달 세션 사후 관리: ',
       'Chat에서 만족도 낮은 세션(Default 발생, 반복 질문 등) 자동 감지 후 '
       '익일 문자 또는 음성 상담 콜백 제공. 문제 해결 완결률 향상.')
    doc.add_page_break()

    # ── 9. 결론 ───────────────────────────────────────────────────────────────
    h1(doc, '9. 결론')

    body(doc,
         '음성 상담과 AI Chat은 동일한 고객을 대상으로 하지만 근본적으로 다른 커뮤니케이션 생태계를 '
         '구성한다. 음성 상담은 서사적 맥락, 감정 회복, 선제적 안내가 강점인 "관계형 채널"이고, '
         'AI Chat은 즉각적 정보 제공, 24시간 접근성, 구조화된 답변이 강점인 "거래형 채널"이다.')
    body(doc,
         '만족도 수치(7.91 vs 7.63)만 보면 차이가 작아 보이지만, 분포 구조를 보면 두 채널은 '
         '서로 다른 고객 집단에서 서로 다른 방식으로 만족을 창출한다. AI Chat에서 만족 고객(9-10점 '
         '62%)이 많다는 것은 단순 질문 해결 효과가 높다는 의미이고, 불만족 고객(24%)이 많다는 것은 '
         '복잡한 문제에서의 실패 빈도가 높다는 의미다.')
    body(doc,
         '가장 시급한 과제는 AI Chat의 Default 응답 개선(만족도 4.74)이다. 전체 세션의 3.7%에 '
         '불과하지만 만족도 낙폭이 가장 크며, 개선 시 전체 지표에 즉각적인 영향을 미친다. '
         '두 번째로는 2턴 세션(약 12%)의 만족도 저점(7.18) 개선이 필요하며, 이는 AI 첫 응답 품질 '
         '및 재질문 유도 방식 개선을 통해 해결 가능하다.')
    body(doc,
         '중장기적으로는 두 채널이 경쟁 관계가 아닌 보완 관계로 설계될 때 최적의 고객 경험이 달성된다. '
         'Chat에서 해결되지 않은 문제를 맥락 전달과 함께 음성 상담으로 연결하는 핸드오프 구조, '
         '음성 상담 이후 Chat으로 팔로업하는 옴니채널 설계가 채널 연계의 핵심 방향이다.')

    doc.add_paragraph()
    tbl(doc,
        ['구분', '현재 수준', '개선 후 목표', '핵심 액션'],
        [
            ['AI Chat 전체 평균 만족도', '7.63', '7.85 이상',
             'Default 개선 + 2턴 저점 개선 + ProductSymptoms 강화'],
            ['AI Chat 불만족 비율',       '24%',  '18% 이하',
             '미해결 세션 상담사 연결 강화 + 재질문 유도 로직'],
            ['Default 만족도',             '4.74', '6.5 이상',
             '구조화 보기 안내 / 오탈자 매칭 강화'],
            ['채널 간 연계율',             '미흡', '10% 이상 연계',
             'Chat 이탈 → 음성 핸드오프 시스템 구축'],
        ]
    )

    # ── 저장 ──────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f'보고서 저장 완료: {OUT_PATH}')


# ══════════════════════════════════════════════════════════════════════════════
# 진입점
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    generate_figures()
    generate_report()
    print('\n완료.')
