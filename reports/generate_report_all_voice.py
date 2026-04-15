# -*- coding: utf-8 -*-
"""
음성 감성분석 보고서 자동 생성기 (독립 실행 스크립트)
- 모든 데이터 하드코딩 (외부 파일 의존 없음)
- 그래프 생성 + Word 보고서 출력
- 출력: 음성_감성분석_보고서.docx
"""

import os
import sys
import io
import warnings

warnings.filterwarnings("ignore")
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.patches as mpatches
from matplotlib.gridspec import GridSpec

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 경로 설정 ────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output_figures")
os.makedirs(OUTPUT_DIR, exist_ok=True)
REPORT_PATH = os.path.join(SCRIPT_DIR, "음성_감성분석_보고서.docx")

# ── 색상 팔레트 ───────────────────────────────────────────────────────────────
NAVY   = "#1B2A4A"
ORANGE = "#D4731A"
LIGHT_NAVY = "#2E4070"
LIGHT_ORANGE = "#E89B52"
GRAY   = "#888888"
LIGHT_GRAY = "#CCCCCC"
WHITE  = "#FFFFFF"

# matplotlib 한글 폰트 설정 (직접 ttf 로드)
def _set_korean_font():
    fp = 'C:/Windows/Fonts/malgun.ttf'
    if os.path.exists(fp):
        fm.fontManager.addfont(fp)
        name = fm.FontProperties(fname=fp).get_name()
        plt.rcParams['font.family'] = name
    else:
        plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['axes.unicode_minus'] = False
    plt.rcParams['figure.facecolor'] = 'white'
    plt.rcParams['axes.facecolor'] = 'white'
    plt.rcParams['axes.spines.top'] = False
    plt.rcParams['axes.spines.right'] = False
    plt.rcParams['axes.edgecolor'] = '#CCCCCC'
    plt.rcParams['xtick.color'] = '#555555'
    plt.rcParams['ytick.color'] = '#555555'

_set_korean_font()

# ── Word 스타일 헬퍼 ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    """표 셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def _cell_font(cell, bold=False, size=10, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
                run.font.color.rgb = RGBColor(r, g, b)

def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def _add_para(doc, text, bold=False, size=10.5, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def _add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def _add_table(doc, headers, rows, col_widths=None, header_color=None):
    """기존 보고서 스타일 표 — 비례 너비, 연회색 헤더, 가운데 정렬, 줄간격 최소."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 비례 너비 계산
    pw = 15.5
    col_max = []
    for ci in range(n_cols):
        lens = [len(str(headers[ci]))]
        for rd in rows:
            if ci < len(rd):
                lens.append(max(len(line) for line in str(rd[ci]).split('\n')))
        col_max.append(max(lens) ** 0.55)
    total = sum(col_max) or 1
    widths = [max(c / total * pw, 2.0) for c in col_max]
    w_sum = sum(widths)
    widths = [w / w_sum * pw for w in widths]
    for ci in range(n_cols):
        for row in table.rows:
            row.cells[ci].width = Cm(widths[ci])

    center = set(range(n_cols - 1)) if n_cols > 1 else set()

    def _cell(cell, text, header=False, ci=0):
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)
        if header:
            run.bold = True
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.0
        if header or ci in center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = cell._element.get_or_add_tcPr()
        mar = tc_pr.makeelement(qn('w:tcMar'), {})
        for side in ['top', 'bottom']:
            mar.append(mar.makeelement(qn(f'w:{side}'), {qn('w:w'): '55', qn('w:type'): 'dxa'}))
        for side in ['left', 'right']:
            mar.append(mar.makeelement(qn(f'w:{side}'), {qn('w:w'): '120', qn('w:type'): 'dxa'}))
        tc_pr.append(mar)

    for i, h in enumerate(headers):
        _cell(table.rows[0].cells[i], h, header=True, ci=i)
        _set_cell_bg(table.rows[0].cells[i], 'F2F2F2')

    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            _cell(table.rows[ri+1].cells[ci], v, ci=ci)

    # 표 아래 여백
    spacer = doc.add_paragraph('')
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    for r in spacer.runs:
        r.font.size = Pt(2)
    return table

def _add_figure(doc, fig_path, width_inches=6.0, caption=None):
    doc.add_picture(fig_path, width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

def _add_insight_box(doc, lines):
    """회색 인사이트 박스 (표 1행 사용)"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, "#FFFFFF")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("주요 인사이트\n")
    run.bold = True
    run.font.size = Pt(10)
    for i, line in enumerate(lines):
        suffix = "\n" if i < len(lines) - 1 else ""
        run2 = p.add_run(f"  - {line}{suffix}")
        run2.font.size = Pt(10)
    # 셀 패딩으로 여유 확보
    tc_pr = cell._element.get_or_add_tcPr()
    mar = tc_pr.makeelement(qn('w:tcMar'), {})
    for side in ['top', 'bottom']:
        mar.append(mar.makeelement(qn(f'w:{side}'), {qn('w:w'): '120', qn('w:type'): 'dxa'}))
    for side in ['left', 'right']:
        mar.append(mar.makeelement(qn(f'w:{side}'), {qn('w:w'): '150', qn('w:type'): 'dxa'}))
    tc_pr.append(mar)
    return table

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def _fig_style(ax, title=None, xlabel=None, ylabel=None):
    """공통 축 스타일"""
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)
    ax.tick_params(colors="#444444", labelsize=9)
    ax.set_facecolor("#FAFAFA")
    if title:
        ax.set_title(title, fontsize=11, fontweight="bold", color="#222222", pad=8)
    if xlabel:
        ax.set_xlabel(xlabel, fontsize=9, color="#555555")
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=9, color="#555555")
    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.8, linestyle="--")


def generate_figures():
    """모든 그래프 생성 후 파일 경로 dict 반환"""
    figs = {}

    # (fig01 파이프라인 다이어그램은 텍스트 설명으로 대체)

    # ─── fig02: 전체 감성 분포 ────────────────────────────────────────────────
    groups = ["감사/만족", "안정/중립", "불안/걱정", "불만/짜증", "혼란/당황"]
    counts = [2245, 4199, 2429, 1668, 1526]
    colors_pie = ["#5B9A6B", "#6B8EB5", "#D4A84B", "#B56B6B", "#8B7BAA"]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
    fig.patch.set_facecolor(WHITE)

    # 파이 차트
    wedges, texts, autotexts = ax1.pie(
        counts, labels=groups, colors=colors_pie,
        autopct="%1.1f%%", startangle=140,
        wedgeprops=dict(linewidth=1, edgecolor=WHITE),
        textprops=dict(fontsize=8.5),
    )
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color(WHITE)
        at.set_fontweight("bold")
    ax1.set_title("고객 감정 그룹 분포", fontsize=11, fontweight="bold", color="#222222")

    # 막대 차트
    y_pos = range(len(groups))
    bars = ax2.barh(y_pos, counts, color=colors_pie, height=0.6, edgecolor=WHITE)
    ax2.set_yticks(y_pos)
    ax2.set_yticklabels(groups, fontsize=9)
    ax2.set_xlabel("발화 건수", fontsize=9, color="#555555")
    ax2.set_title("감정 그룹별 발화 수", fontsize=11, fontweight="bold", color="#222222")
    ax2.spines["top"].set_visible(False)
    ax2.spines["right"].set_visible(False)
    ax2.set_facecolor("#FAFAFA")
    for bar, cnt in zip(bars, counts):
        ax2.text(bar.get_width() + 10, bar.get_y() + bar.get_height() / 2,
                 f"{cnt:,}", va="center", fontsize=8.5, color="#444444")

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig02_emotion_dist.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig02"] = path
    print(f"  [fig02] 감성 분포 저장")

    # ─── fig03: 상담 단계별 감성 궤적 ────────────────────────────────────────
    stages_label = ["초기", "탐색", "해결시도", "결과제시", "종료"]
    stage_values = [-0.043, -0.121, -0.096, -0.051, +0.132]

    fig, ax = plt.subplots(figsize=(8, 4))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    x = range(len(stages_label))
    ax.plot(x, stage_values, color=NAVY, linewidth=2.5, marker="o",
            markersize=9, markerfacecolor=ORANGE, markeredgecolor=WHITE,
            markeredgewidth=2, zorder=5)
    ax.fill_between(x, stage_values, 0,
                    where=[v > 0 for v in stage_values],
                    alpha=0.15, color=ORANGE, label="긍정 구간")
    ax.fill_between(x, stage_values, 0,
                    where=[v <= 0 for v in stage_values],
                    alpha=0.12, color=NAVY, label="부정 구간")

    for xi, val in zip(x, stage_values):
        offset = 0.012 if val >= 0 else -0.016
        ax.annotate(f"{val:+.3f}", (xi, val + offset),
                    ha="center", fontsize=9.5, fontweight="bold",
                    color=ORANGE if val > 0 else NAVY)

    ax.set_xticks(x)
    ax.set_xticklabels(stages_label, fontsize=10)
    ax.set_ylabel("감성 점수", fontsize=9, color="#555555")
    ax.set_title("상담 단계별 평균 감성 점수 궤적", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax.legend(fontsize=9, framealpha=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig03_stage_trajectory.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig03"] = path
    print(f"  [fig03] 단계별 궤적 저장")

    # ─── fig04: 감성 전환 유형 분포 ──────────────────────────────────────────
    transition_types = ["부정→긍정\n(회복)", "긍정→부정\n(악화)", "부정 지속\n(미해결)", "긍정 유지\n(안정)"]
    transition_pct  = [39.3, 12.0, 27.5, 21.2]
    t_colors = ["#5B9A6B", "#B56B6B", "#D4A84B", "#6B8EB5"]

    fig, ax = plt.subplots(figsize=(7, 4))
    fig.patch.set_facecolor(WHITE)
    bars = ax.bar(transition_types, transition_pct, color=t_colors,
                  width=0.55, edgecolor=WHITE, linewidth=1.2)
    for bar, pct in zip(bars, transition_pct):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                f"{pct:.1f}%", ha="center", va="bottom", fontsize=10.5,
                fontweight="bold", color="#333333")
    ax.set_ylabel("비율 (%)", fontsize=9, color="#555555")
    ax.set_title("감성 전환 유형 분포", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax.set_ylim(0, 50)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_facecolor("#FAFAFA")
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig04_transitions.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig04"] = path
    print(f"  [fig04] 감성 전환 유형 저장")

    # ─── fig05: 제품군별 감성 비교 ───────────────────────────────────────────
    products = ["에어컨/에어케어", "주방가전", "생활가전", "TV/AV"]
    prod_stage_vals = {
        "초기":    [-0.063, -0.054, -0.022, -0.037],
        "탐색":    [-0.103, -0.126, -0.120, -0.104],
        "해결시도": [-0.096, -0.093, -0.094, -0.107],
        "결과제시": [-0.068, -0.046, -0.049, -0.072],
        "종료":    [+0.116, +0.135, +0.132, +0.115],
    }
    stage_colors_list = [NAVY, "#3A5A8F", "#5C7DB5", LIGHT_ORANGE, ORANGE]

    x = np.arange(len(products))
    width = 0.15
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    for i, (stage, vals) in enumerate(prod_stage_vals.items()):
        offset = (i - 2) * width
        bars = ax.bar(x + offset, vals, width, label=stage,
                      color=stage_colors_list[i], edgecolor=WHITE, alpha=0.88)

    ax.set_xticks(x)
    ax.set_xticklabels(products, fontsize=10)
    ax.set_ylabel("감성 점수", fontsize=9, color="#555555")
    ax.set_title("제품군별 상담 단계 감성 점수", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax.legend(fontsize=9, loc="lower right", framealpha=0.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig05_product_sentiment.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig05"] = path
    print(f"  [fig05] 제품군별 감성 저장")

    # ─── fig06: 연령대별 감성 패턴 ───────────────────────────────────────────
    age_groups = ["20~39세", "40~49세", "50~64세", "65~74세", "75세이상"]
    age_vals = {
        "초기":    [-0.055, -0.040, -0.034, -0.043, -0.072],
        "탐색":    [-0.117, -0.119, -0.117, -0.125, -0.129],
        "종료":    [+0.150, +0.140, +0.134, +0.104, +0.113],
    }
    age_colors = [NAVY, LIGHT_NAVY, ORANGE]
    x = np.arange(len(age_groups))
    width = 0.25

    fig, ax = plt.subplots(figsize=(9, 4.5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    for i, (stage, vals) in enumerate(age_vals.items()):
        offset = (i - 1) * width
        ax.bar(x + offset, vals, width, label=stage,
               color=age_colors[i], edgecolor=WHITE, alpha=0.88)

    ax.set_xticks(x)
    ax.set_xticklabels(age_groups, fontsize=9.5)
    ax.set_ylabel("감성 점수", fontsize=9, color="#555555")
    ax.set_title("연령대별 주요 단계 감성 점수", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax.legend(fontsize=9, framealpha=0.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig06_age_sentiment.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig06"] = path
    print(f"  [fig06] 연령대별 감성 저장")

    # ─── fig07: NPS vs 감성 산점도 (실제 데이터) ────────────────────────────
    import pandas as _pd
    _csv_path = os.path.join(os.path.dirname(SCRIPT_DIR), 'analysis', 'outputs', 'exports', '전체_음성분석_데이터_LLM_콜단위요약.csv')
    _df_calls = _pd.read_csv(_csv_path, encoding='utf-8-sig')
    _valid = _df_calls[['NPS','고객평균감성']].dropna()
    nps_vals = _valid['NPS'].values
    sentiment_vals = _valid['고객평균감성'].values

    fig, ax = plt.subplots(figsize=(7, 5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    sc = ax.scatter(nps_vals + np.random.uniform(-0.2, 0.2, len(nps_vals)),
                    sentiment_vals,
                    c=sentiment_vals, cmap="RdYlGn",
                    alpha=0.6, s=45, edgecolors="none")
    z = np.polyfit(nps_vals, sentiment_vals, 1)
    p_fit = np.poly1d(z)
    xfit = np.linspace(0, 10, 100)
    ax.plot(xfit, p_fit(xfit), color=NAVY, linewidth=2, linestyle="--",
            label=f"추세선 (r=0.262)")
    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.8, linestyle=":")
    ax.set_xlabel("NPS (0~10)", fontsize=9, color="#555555")
    ax.set_ylabel("평균 감성 점수", fontsize=9, color="#555555")
    ax.set_title("NPS vs 평균 감성 점수 (r=0.262)", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax.legend(fontsize=9, framealpha=0.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)
    plt.colorbar(sc, ax=ax, label="감성 점수", shrink=0.8)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig07_nps_scatter.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig07"] = path
    print(f"  [fig07] NPS 산점도 저장")

    # ─── fig08: 모달리티 비교 ────────────────────────────────────────────────
    modalities = ["텍스트만\n(BERT)", "음향만", "SER만", "텍스트+음향", "3-way 융합"]
    corr_nps   = [0.116, -0.021, 0.171, 0.229, 0.262]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    fig.patch.set_facecolor(WHITE)

    bar_colors = [LIGHT_GRAY, LIGHT_GRAY, LIGHT_GRAY, LIGHT_NAVY, ORANGE]
    x = range(len(modalities))

    bars = ax.bar(x, corr_nps, color=bar_colors, width=0.6,
                  edgecolor=WHITE, linewidth=1)
    for bar, v in zip(bars, corr_nps):
        y_pos = max(v, 0) + 0.008
        ax.text(bar.get_x() + bar.get_width() / 2, y_pos,
                f"{v:+.3f}", ha="center", va="bottom", fontsize=10.5,
                fontweight="bold", color="#333333")
    ax.set_xticks(x)
    ax.set_xticklabels(modalities, fontsize=9.5)
    ax.set_ylabel("NPS 상관계수 (Pearson r)", fontsize=10, color="#555555")
    ax.set_title("분석 방법별 NPS 상관계수 비교", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax.set_facecolor("#FAFAFA")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig08_modality_compare.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig08"] = path
    print(f"  [fig08] 모달리티 비교 저장")

    # ─── fig09: 감성 회복 패턴 ───────────────────────────────────────────────
    stages_label = ["초기", "탐색", "해결시도", "결과제시", "종료"]
    recovery_case   = [-0.047, -0.122, -0.093, -0.050, +0.215]
    norecovery_case = [-0.042, -0.109, -0.094, -0.064, -0.082]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    x = range(len(stages_label))
    ax.plot(x, recovery_case, color=ORANGE, linewidth=2.5, marker="o",
            markersize=9, markerfacecolor=ORANGE, markeredgecolor=WHITE,
            markeredgewidth=2, label="감정 회복 그룹 (N=327)", zorder=5)
    ax.plot(x, norecovery_case, color=NAVY, linewidth=2.5, marker="s",
            markersize=9, markerfacecolor=NAVY, markeredgecolor=WHITE,
            markeredgewidth=2, linestyle="--", label="감정 부정 그룹 (N=96)", zorder=5)

    ax.fill_between(x, recovery_case, norecovery_case,
                    alpha=0.08, color=ORANGE, label="그룹 간 격차")

    ax.set_xticks(x)
    ax.set_xticklabels(stages_label, fontsize=10)
    ax.set_ylabel("감성 점수", fontsize=9, color="#555555")
    ax.set_title("종료 감정 그룹별 단계 궤적 비교", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax.legend(fontsize=9, framealpha=0.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig09_recovery_pattern.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig09"] = path
    print(f"  [fig09] 감성 회복 패턴 저장")

    # ─── fig10: 컨설턴트 만족도 vs 감성/NPS ─────────────────────────────────
    scores = [1, 2, 3, 4, 5]
    avg_sentiment = [-0.158, -0.102, -0.095, -0.081, -0.056]
    avg_nps       = [2.33,   3.53,   5.50,   7.02,   8.45]
    ns            = [9,      15,     28,     111,    523]

    fig, ax1 = plt.subplots(figsize=(8, 4.5))
    fig.patch.set_facecolor(WHITE)
    ax1.set_facecolor("#FAFAFA")

    ax2 = ax1.twinx()
    x = np.array(scores)
    bar_w = 0.35

    bars = ax1.bar(x, avg_sentiment, bar_w,
                   color=NAVY, alpha=0.8, label="평균 감성 점수", edgecolor=WHITE)
    line, = ax2.plot(x, avg_nps, color=ORANGE, linewidth=2.5, marker="D",
                     markersize=9, markerfacecolor=ORANGE, markeredgecolor=WHITE,
                     markeredgewidth=2, label="평균 NPS", zorder=5)

    ax1.set_xlabel("컨설턴트 만족도 점수", fontsize=9, color="#555555")
    ax1.set_ylabel("평균 감성 점수", fontsize=9, color=NAVY)
    ax2.set_ylabel("평균 NPS", fontsize=9, color=ORANGE)
    ax1.set_title("컨설턴트 만족도별 감성 점수 및 NPS", fontsize=12, fontweight="bold", color="#222222", pad=10)
    ax1.axhline(0, color=LIGHT_GRAY, linewidth=0.8, linestyle="--")

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, fontsize=9, framealpha=0.8, loc="upper left")

    # N 표시
    for xi, n_val in zip(x, ns):
        ax1.text(xi, min(avg_sentiment) - 0.022, f"N={n_val}",
                 ha="center", fontsize=7.5, color=GRAY)

    ax1.spines["top"].set_visible(False)
    ax2.spines["top"].set_visible(False)
    ax1.spines["left"].set_color(LIGHT_GRAY)
    ax1.spines["bottom"].set_color(LIGHT_GRAY)
    ax2.spines["right"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig10_consultant_sat.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig10"] = path
    print(f"  [fig10] 컨설턴트 만족도 저장")

    # ─── fig11: 짧은 발화 "네" 감성 분석 ─────────────────────────────────────
    # 상단: Energy 높은/낮은 "네" 비교
    # 하단: 단계별 + 맥락별 "네" 감성
    fig = plt.figure(figsize=(11, 8))
    fig.patch.set_facecolor(WHITE)
    gs = GridSpec(2, 2, figure=fig, hspace=0.45, wspace=0.35)

    ax_energy  = fig.add_subplot(gs[0, 0])
    ax_stage   = fig.add_subplot(gs[0, 1])
    ax_context = fig.add_subplot(gs[1, 0])
    ax_phrase  = fig.add_subplot(gs[1, 1])

    # (A) Energy 높은/낮은 "네"
    energy_labels = ["Energy 높은 '네'\n(밝은 톤)", "Energy 낮은 '네'\n(어두운 톤)"]
    energy_vals   = [+0.140, -0.069]
    energy_colors = [ORANGE, NAVY]

    bars = ax_energy.bar(energy_labels, energy_vals,
                         color=energy_colors, width=0.45,
                         edgecolor=WHITE, linewidth=1.2)
    for bar, v in zip(bars, energy_vals):
        ypos = v + 0.008 if v >= 0 else v - 0.014
        ax_energy.text(bar.get_x() + bar.get_width() / 2, ypos,
                       f"{v:+.3f}", ha="center", va="bottom" if v >= 0 else "top",
                       fontsize=11, fontweight="bold",
                       color=ORANGE if v > 0 else NAVY)
    _fig_style(ax_energy, title="(A) 음량별 '네' 감성 점수",
               ylabel="감성 점수")
    ax_energy.set_ylim(-0.15, 0.22)
    ax_energy.set_facecolor("#FAFAFA")

    # (B) 단계별 "네"
    ne_stages = ["초기", "탐색", "해결시도", "종료"]
    ne_stage_vals = [+0.092, -0.013, +0.027, +0.113]
    bar_cols = [ORANGE if v > 0 else NAVY for v in ne_stage_vals]

    bars = ax_stage.bar(ne_stages, ne_stage_vals,
                        color=bar_cols, width=0.5,
                        edgecolor=WHITE, linewidth=1.2)
    for bar, v in zip(bars, ne_stage_vals):
        ypos = v + 0.005 if v >= 0 else v - 0.012
        ax_stage.text(bar.get_x() + bar.get_width() / 2, ypos,
                      f"{v:+.3f}", ha="center",
                      va="bottom" if v >= 0 else "top",
                      fontsize=10, fontweight="bold",
                      color=ORANGE if v > 0 else NAVY)
    _fig_style(ax_stage, title="(B) 단계별 '네' 감성 점수",
               ylabel="감성 점수")
    ax_stage.set_facecolor("#FAFAFA")
    ax_stage.set_ylim(-0.08, 0.18)

    # (C) 맥락별 "네"
    ctx_labels = ["사과/공감\n후", "비용 안내\n후", "감사/인사\n후"]
    ctx_vals   = [+0.079, +0.014, +0.053]
    ctx_cols   = [ORANGE if v > 0 else NAVY for v in ctx_vals]

    bars = ax_context.bar(ctx_labels, ctx_vals,
                          color=ctx_cols, width=0.45,
                          edgecolor=WHITE, linewidth=1.2)
    for bar, v in zip(bars, ctx_vals):
        ypos = v + 0.005 if v >= 0 else v - 0.012
        ax_context.text(bar.get_x() + bar.get_width() / 2, ypos,
                        f"{v:+.3f}", ha="center",
                        va="bottom" if v >= 0 else "top",
                        fontsize=10, fontweight="bold",
                        color=ORANGE if v > 0 else NAVY)
    _fig_style(ax_context, title="(C) 맥락별 '네' 감성 점수",
               ylabel="감성 점수")
    ax_context.set_facecolor("#FAFAFA")
    ax_context.set_ylim(-0.08, 0.16)

    # (D) 구문별 감정 그룹 비율
    phrases = ['"네 감사합니다"', '"네 네 네"']
    grateful_pct = [65, 15]
    anxious_pct  = [12, 31]
    other_pct    = [23, 54]

    x_ph = np.arange(len(phrases))
    w_ph = 0.25
    ax_phrase.bar(x_ph - w_ph, grateful_pct, w_ph,
                  label="감사/만족", color="#2E7D32", edgecolor=WHITE)
    ax_phrase.bar(x_ph,         anxious_pct,  w_ph,
                  label="불안/걱정", color=ORANGE,    edgecolor=WHITE)
    ax_phrase.bar(x_ph + w_ph, other_pct,    w_ph,
                  label="기타",    color=LIGHT_GRAY, edgecolor=WHITE)
    ax_phrase.set_xticks(x_ph)
    ax_phrase.set_xticklabels(phrases, fontsize=9)
    ax_phrase.set_ylabel("비율 (%)", fontsize=9, color="#555555")
    ax_phrase.set_title("(D) 구문별 감정 분포", fontsize=11,
                         fontweight="bold", color="#222222", pad=8)
    ax_phrase.legend(fontsize=8, framealpha=0.8)
    ax_phrase.spines["top"].set_visible(False)
    ax_phrase.spines["right"].set_visible(False)
    ax_phrase.spines["left"].set_color(LIGHT_GRAY)
    ax_phrase.spines["bottom"].set_color(LIGHT_GRAY)
    ax_phrase.set_facecolor("#FAFAFA")
    ax_phrase.axhline(0, color=LIGHT_GRAY, linewidth=0.8, linestyle="--")

    fig.suptitle("짧은 발화 '네' (1,065건) 세부 감성 분석",
                 fontsize=13, fontweight="bold", color="#222222", y=1.01)
    path = os.path.join(OUTPUT_DIR, "fig11_ne_analysis.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig11"] = path
    print(f"  [fig11] 짧은 발화 '네' 분석 저장")

    # ─── fig12: 상담사 응대 패턴과 고객 감정 변화 ─────────────────────────────
    fig = plt.figure(figsize=(11, 8))
    fig.patch.set_facecolor(WHITE)
    gs = GridSpec(2, 2, figure=fig, hspace=0.45, wspace=0.35)

    ax_type   = fig.add_subplot(gs[0, 0])
    ax_len    = fig.add_subplot(gs[0, 1])
    ax_stage  = fig.add_subplot(gs[1, 0])
    ax_severe = fig.add_subplot(gs[1, 1])

    # (A) 응대 유형별 감정 변화량
    types = ["짧은응답", "공감만", "공감+\n해결안내", "해결안내만", "일반안내"]
    type_delta = [+0.176, +0.247, +0.224, +0.251, +0.262]
    type_improv = [64, 67, 71, 73, 73]
    type_colors = [NAVY, "#5C7DB5", LIGHT_ORANGE, ORANGE, "#2E7D32"]

    bars = ax_type.bar(types, type_delta, color=type_colors, width=0.55,
                       edgecolor=WHITE, linewidth=1.2)
    for bar, d, imp in zip(bars, type_delta, type_improv):
        ax_type.text(bar.get_x() + bar.get_width() / 2,
                     bar.get_height() + 0.005,
                     f"+{d:.3f}\n({imp}%)", ha="center", va="bottom",
                     fontsize=8.5, fontweight="bold", color="#333333")
    _fig_style(ax_type, title="(A) 응대 유형별 감정 변화량",
               ylabel="감성 점수 변화")
    ax_type.set_facecolor("#FAFAFA")
    ax_type.set_ylim(0, 0.35)

    # (B) 응답 길이별 감정 변화량
    len_labels = ["짧은\n(<15자)", "중간\n(15~40자)", "긴\n(40~80자)", "상세\n(80자+)"]
    len_delta = [+0.183, +0.281, +0.255, +0.224]
    len_improv = [65, 75, 74, 71]
    len_colors = [NAVY, ORANGE, LIGHT_ORANGE, LIGHT_NAVY]

    bars = ax_len.bar(len_labels, len_delta, color=len_colors, width=0.55,
                      edgecolor=WHITE, linewidth=1.2)
    for bar, d, imp in zip(bars, len_delta, len_improv):
        ax_len.text(bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + 0.005,
                    f"+{d:.3f}\n({imp}%)", ha="center", va="bottom",
                    fontsize=8.5, fontweight="bold", color="#333333")
    _fig_style(ax_len, title="(B) 응답 길이별 감정 변화량",
               ylabel="감성 점수 변화")
    ax_len.set_facecolor("#FAFAFA")
    ax_len.set_ylim(0, 0.36)

    # (C) 단계별 짧은 응답 vs 해결 안내 비교
    stage_labels = ["탐색", "해결시도", "결과제시"]
    short_vals = [0.131, 0.208, 0.261]
    resolve_vals = [0.244, 0.242, 0.299]

    x_st = np.arange(len(stage_labels))
    w_st = 0.3
    ax_stage.bar(x_st - w_st/2, short_vals, w_st, label="짧은 응답",
                 color=NAVY, edgecolor=WHITE, alpha=0.88)
    ax_stage.bar(x_st + w_st/2, resolve_vals, w_st, label="해결 안내",
                 color=ORANGE, edgecolor=WHITE, alpha=0.88)
    for i, (sv, rv) in enumerate(zip(short_vals, resolve_vals)):
        ax_stage.text(i - w_st/2, sv + 0.005, f"+{sv:.3f}", ha="center",
                      fontsize=8, color=NAVY, fontweight="bold")
        ax_stage.text(i + w_st/2, rv + 0.005, f"+{rv:.3f}", ha="center",
                      fontsize=8, color=ORANGE, fontweight="bold")
    ax_stage.set_xticks(x_st)
    ax_stage.set_xticklabels(stage_labels, fontsize=10)
    ax_stage.set_ylabel("감성 점수 변화", fontsize=9, color="#555555")
    ax_stage.set_title("(C) 단계별 응대 유형 효과 비교", fontsize=11,
                       fontweight="bold", color="#222222", pad=8)
    ax_stage.legend(fontsize=8.5, framealpha=0.8)
    ax_stage.axhline(0, color=LIGHT_GRAY, linewidth=0.8, linestyle="--")
    ax_stage.set_facecolor("#FAFAFA")
    ax_stage.spines["top"].set_visible(False)
    ax_stage.spines["right"].set_visible(False)
    ax_stage.spines["left"].set_color(LIGHT_GRAY)
    ax_stage.spines["bottom"].set_color(LIGHT_GRAY)
    ax_stage.set_ylim(0, 0.38)

    # (D) 부정 강도별 짧은 응답 효과
    sev_labels = ["경미\n(-0.05~-0.15)", "중간\n(-0.15~-0.3)", "심각\n(-0.3~-0.5)", "극심\n(-0.5 이하)"]
    sev_short = [-0.012, +0.106, +0.304, +0.458]
    sev_all   = [+0.060, +0.151, +0.328, +0.529]
    x_sv = np.arange(len(sev_labels))

    ax_severe.bar(x_sv - w_st/2, sev_short, w_st, label="짧은 응답",
                  color=NAVY, edgecolor=WHITE, alpha=0.88)
    ax_severe.bar(x_sv + w_st/2, sev_all, w_st, label="전체 평균",
                  color=ORANGE, edgecolor=WHITE, alpha=0.88)
    for i, (ss, sa) in enumerate(zip(sev_short, sev_all)):
        c = "#C62828" if ss < 0 else NAVY
        ax_severe.text(i - w_st/2, max(ss, 0) + 0.008, f"{ss:+.3f}", ha="center",
                       fontsize=7.5, color=c, fontweight="bold")
        ax_severe.text(i + w_st/2, sa + 0.008, f"+{sa:.3f}", ha="center",
                       fontsize=7.5, color=ORANGE, fontweight="bold")
    ax_severe.set_xticks(x_sv)
    ax_severe.set_xticklabels(sev_labels, fontsize=8.5)
    ax_severe.set_ylabel("감성 점수 변화", fontsize=9, color="#555555")
    ax_severe.set_title("(D) 부정 강도별 짧은 응답 vs 전체", fontsize=11,
                        fontweight="bold", color="#222222", pad=8)
    ax_severe.legend(fontsize=8.5, framealpha=0.8, loc="upper left")
    ax_severe.axhline(0, color=LIGHT_GRAY, linewidth=0.8, linestyle="--")
    ax_severe.set_facecolor("#FAFAFA")
    ax_severe.spines["top"].set_visible(False)
    ax_severe.spines["right"].set_visible(False)
    ax_severe.spines["left"].set_color(LIGHT_GRAY)
    ax_severe.spines["bottom"].set_color(LIGHT_GRAY)

    fig.suptitle("상담사 응대 패턴과 고객 감정 변화 (5,548 시퀀스)",
                 fontsize=13, fontweight="bold", color="#222222", y=1.01)
    path = os.path.join(OUTPUT_DIR, "fig12_response_pattern.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig12"] = path
    print(f"  [fig12] 응대 패턴 분석 저장")

    # ─── fig13: 제품군 × 연령대 히트맵 ──────────────────────────────────────
    products = ["주방가전", "생활가전", "TV/AV", "에어컨/에어케어"]
    ages = ["20~39세", "40~49세", "50~64세", "65~74세", "75세이상"]
    heatmap_data = np.array([
        [-0.054, -0.075, -0.060, -0.073, -0.050],  # 주방가전
        [-0.035, -0.048, -0.055, -0.086, -0.035],  # 생활가전
        [-0.058, -0.049, -0.053, -0.100, -0.075],  # TV/AV
        [-0.067, -0.031, -0.053, -0.074, -0.118],  # 에어컨
    ])

    fig, ax = plt.subplots(figsize=(9, 4.5))
    fig.patch.set_facecolor(WHITE)

    im = ax.imshow(heatmap_data, cmap="RdYlGn", aspect="auto",
                   vmin=-0.13, vmax=-0.02)
    ax.set_xticks(range(len(ages)))
    ax.set_xticklabels(ages, fontsize=9.5)
    ax.set_yticks(range(len(products)))
    ax.set_yticklabels(products, fontsize=10)

    for i in range(len(products)):
        for j in range(len(ages)):
            v = heatmap_data[i, j]
            color = WHITE if v < -0.08 else "#333333"
            ax.text(j, i, f"{v:+.3f}", ha="center", va="center",
                    fontsize=10, fontweight="bold", color=color)

    cb = plt.colorbar(im, ax=ax, shrink=0.85, pad=0.02)
    cb.set_label("평균 감성 점수", fontsize=9)
    ax.set_title("제품군 × 연령대 교차 감성 히트맵", fontsize=12,
                 fontweight="bold", color="#222222", pad=10)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig13_heatmap.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig13"] = path
    print(f"  [fig13] 교차 히트맵 저장")

    # ─── fig14: 감정 그룹 단계 분포 레이더차트 ──────────────────────────────
    groups_radar = ["감사/만족", "안정/중립", "불안/걱정", "불만/짜증", "혼란/당황"]
    stage_pcts = {
        "감사/만족": [11, 16, 26, 19, 27],
        "안정/중립": [14, 30, 32, 14, 10],
        "불안/걱정": [12, 30, 36, 17,  5],
        "불만/짜증": [ 9, 32, 38, 16,  6],
        "혼란/당황": [11, 32, 35, 18,  4],
    }
    radar_stages = ["초기", "탐색", "해결시도", "결과제시", "종료"]
    grp_colors = ["#5B9A6B", "#6B8EB5", "#D4A84B", "#B56B6B", "#8B7BAA"]
    grp_lw     = [2.8, 2.0, 2.0, 2.8, 2.0]
    grp_alpha  = [0.18, 0.06, 0.06, 0.18, 0.06]
    grp_zorder = [7, 3, 4, 6, 2]

    n_vars = len(radar_stages)
    angles = np.linspace(0, 2 * np.pi, n_vars, endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor(WHITE)

    # 격자 스타일
    ax.yaxis.grid(True, color="#E8E8E8", linewidth=0.5)
    ax.xaxis.grid(True, color="#D8D8D8", linewidth=0.6)
    ax.spines["polar"].set_visible(False)

    # 5그룹 모두 그리되, 감사/만족과 불만/짜증을 굵게 강조
    for gi, (grp, pcts) in enumerate(stage_pcts.items()):
        values = pcts + pcts[:1]
        ax.plot(angles, values, linewidth=grp_lw[gi], label=grp,
                color=grp_colors[gi], marker="o",
                markersize=7 if grp_lw[gi] > 2.5 else 5,
                markerfacecolor=grp_colors[gi], markeredgecolor=WHITE,
                markeredgewidth=1.5, zorder=grp_zorder[gi])
        ax.fill(angles, values, alpha=grp_alpha[gi], color=grp_colors[gi])

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(radar_stages, fontsize=12, fontweight="bold", color="#444444")
    ax.set_ylim(0, 44)
    ax.set_yticks([10, 20, 30, 40])
    ax.set_yticklabels(["10%", "20%", "30%", "40%"], fontsize=8, color=GRAY)

    # 핵심 수치 주석 — 겹치지 않는 위치에 배치
    ax.annotate("감사/만족\n종료 27%",
                xy=(angles[4], 27), xytext=(angles[4] + 0.35, 35),
                fontsize=9, fontweight="bold", color="#5B9A6B",
                arrowprops=dict(arrowstyle="->", color="#5B9A6B", lw=1.2),
                bbox=dict(boxstyle="round,pad=0.3", fc="#F0FFF0", ec="#5B9A6B", alpha=0.9))
    ax.annotate("불만/짜증\n해결시도 38%",
                xy=(angles[2], 38), xytext=(angles[2] - 0.4, 44),
                fontsize=9, fontweight="bold", color="#B56B6B",
                arrowprops=dict(arrowstyle="->", color="#B56B6B", lw=1.2),
                bbox=dict(boxstyle="round,pad=0.3", fc="#FFF0F0", ec="#B56B6B", alpha=0.9))

    ax.set_title("감정 그룹별 단계 분포 패턴", fontsize=13,
                 fontweight="bold", color="#222222", y=1.08)
    ax.legend(loc="lower right", bbox_to_anchor=(1.28, -0.05),
              fontsize=9.5, framealpha=0.95, edgecolor="#E0E0E0")

    path = os.path.join(OUTPUT_DIR, "fig14_radar.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig14"] = path
    print(f"  [fig14] 레이더차트 저장")

    # ─── fig15: 단계별 감성 워터폴 차트 ──────────────────────────────────────
    wf_stages = ["초기", "탐색", "해결시도", "결과제시", "종료"]
    wf_values = [-0.043, -0.121, -0.096, -0.051, +0.132]
    wf_deltas = [-0.043, -0.078, +0.024, +0.045, +0.183]

    fig, ax = plt.subplots(figsize=(9, 5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    # 워터폴 바 계산
    bottoms = []
    for i, d in enumerate(wf_deltas):
        if i == 0:
            bottoms.append(0 if d >= 0 else d)
        else:
            prev_top = wf_values[i - 1]
            bottoms.append(min(prev_top, prev_top + d) if d < 0 else prev_top)

    bar_colors = []
    for d in wf_deltas:
        bar_colors.append("#C62828" if d < 0 else "#2E7D32")

    bars = ax.bar(wf_stages, [abs(d) for d in wf_deltas],
                  bottom=bottoms, color=bar_colors,
                  width=0.5, edgecolor=WHITE, linewidth=1.5, alpha=0.85)

    # 연결선
    for i in range(len(wf_stages) - 1):
        y_connect = wf_values[i]
        ax.plot([i + 0.25, i + 0.75], [y_connect, y_connect],
                color=GRAY, linewidth=1, linestyle=":")

    # 값 표시
    for i, (bar, d, v) in enumerate(zip(bars, wf_deltas, wf_values)):
        y_text = v + 0.008 if d >= 0 else v - 0.015
        ax.text(i, y_text, f"{v:+.3f}", ha="center", va="bottom" if d >= 0 else "top",
                fontsize=11, fontweight="bold", color="#333333")
        # 변화량 표시
        y_delta = bar.get_y() + bar.get_height() / 2
        ax.text(i, y_delta, f"({d:+.3f})", ha="center", va="center",
                fontsize=8.5, color=WHITE, fontweight="bold")

    ax.axhline(0, color=LIGHT_GRAY, linewidth=1, linestyle="--")
    ax.set_ylabel("감성 점수", fontsize=9, color="#555555")
    ax.set_title("상담 단계별 감성 점수 워터폴", fontsize=12,
                 fontweight="bold", color="#222222", pad=10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    from matplotlib.patches import Patch
    ax.legend(handles=[Patch(color="#C62828", label="감성 하강"),
                       Patch(color="#2E7D32", label="감성 상승")],
              fontsize=9, framealpha=0.8, loc="lower right")

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig15_waterfall.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig15"] = path
    print(f"  [fig15] 워터폴 차트 저장")

    # ─── fig16: 괴리 케이스 궤적 비교 + 상담유형별 회복 비교 ────────────────
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 5))
    fig.patch.set_facecolor(WHITE)

    # (좌) 괴리 vs 정상 궤적
    x_st = range(5)
    gap_vals    = [-0.119, -0.153, -0.096, -0.072, +0.124]
    normal_vals = [-0.025, -0.109, -0.091, -0.042, +0.143]
    stage_labels_5 = ["초기", "탐색", "해결시도", "결과제시", "종료"]

    ax1.plot(x_st, gap_vals, color="#C62828", linewidth=2.5, marker="o",
             markersize=9, markerfacecolor="#C62828", markeredgecolor=WHITE,
             markeredgewidth=2, label="괴리 (NPS低+만족高, N=67)", zorder=5)
    ax1.plot(x_st, normal_vals, color="#2E7D32", linewidth=2.5, marker="s",
             markersize=9, markerfacecolor="#2E7D32", markeredgecolor=WHITE,
             markeredgewidth=2, label="정상 (NPS高+만족高, N=483)", zorder=5)
    ax1.fill_between(x_st, gap_vals, normal_vals,
                     alpha=0.08, color="#C62828")

    for xi, (gv, nv) in enumerate(zip(gap_vals, normal_vals)):
        ax1.text(xi, gv - 0.012, f"{gv:+.3f}", ha="center",
                 fontsize=8, color="#C62828", fontweight="bold")
        ax1.text(xi, nv + 0.008, f"{nv:+.3f}", ha="center",
                 fontsize=8, color="#2E7D32", fontweight="bold")

    ax1.set_xticks(x_st)
    ax1.set_xticklabels(stage_labels_5, fontsize=9.5)
    ax1.axhline(0, color=LIGHT_GRAY, linewidth=0.8, linestyle="--")
    ax1.set_title("NPS-만족도 괴리 케이스 궤적 비교", fontsize=11,
                  fontweight="bold", color="#222222", pad=8)
    ax1.set_ylabel("감성 점수", fontsize=9, color="#555555")
    ax1.legend(fontsize=8, framealpha=0.8, loc="lower right")
    ax1.set_facecolor("#FAFAFA")
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)
    ax1.spines["left"].set_color(LIGHT_GRAY)
    ax1.spines["bottom"].set_color(LIGHT_GRAY)

    # (우) 상담유형별 회복폭
    types = ["사용설명/\n기능안내", "반품-제품", "수리/설치", "반품-제품/\n소모품"]
    recovery = [0.238, 0.292, 0.246, 0.208]
    type_colors = [LIGHT_NAVY, ORANGE, "#5C7DB5", LIGHT_ORANGE]

    bars = ax2.barh(types, recovery, color=type_colors, height=0.55,
                    edgecolor=WHITE, linewidth=1.2)
    for bar, r in zip(bars, recovery):
        ax2.text(bar.get_width() + 0.005, bar.get_y() + bar.get_height() / 2,
                 f"+{r:.3f}", va="center", fontsize=10.5,
                 fontweight="bold", color="#333333")

    ax2.set_xlabel("회복폭 (종료 - 탐색)", fontsize=9, color="#555555")
    ax2.set_title("상담유형별 감정 회복폭", fontsize=11,
                  fontweight="bold", color="#222222", pad=8)
    ax2.set_facecolor("#FAFAFA")
    ax2.spines["top"].set_visible(False)
    ax2.spines["right"].set_visible(False)
    ax2.spines["left"].set_color(LIGHT_GRAY)
    ax2.spines["bottom"].set_color(LIGHT_GRAY)
    ax2.set_xlim(0, 0.38)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig16_cross_analysis.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig16"] = path
    print(f"  [fig16] 교차 분석 저장")

    # ─── fig17: 긍정 vs 부정 종료 감정 궤적 + 트리거 ───────────────────────
    x_bins = np.arange(10)
    x_pct = [f"{i*10}%" for i in range(10)]
    stage_bounds = [(0, 0.9, "초기"), (1, 2.4, "탐색"), (2.5, 5.9, "해결시도"),
                    (6, 7.9, "결과제시"), (8, 9.5, "종료")]

    pos_traj = [-0.077, -0.119, -0.114, -0.093, -0.103, -0.101, -0.080, -0.087, -0.024, +0.192]
    neg_traj = [-0.059, -0.092, -0.113, -0.147, -0.104, -0.049, -0.078, -0.111, -0.046, -0.071]

    fig, ax1 = plt.subplots(figsize=(11, 6))
    fig.patch.set_facecolor(WHITE)

    ax1.set_facecolor("#FAFAFA")
    ax1.plot(x_bins, pos_traj, color="#2E7D32", linewidth=2.5, marker="o",
             markersize=10, markerfacecolor="#2E7D32", markeredgecolor=WHITE,
             markeredgewidth=2, label="긍정 종료 (N=327, 47.7%)", zorder=5)
    ax1.plot(x_bins, neg_traj, color="#C62828", linewidth=2.5, marker="s",
             markersize=10, markerfacecolor="#C62828", markeredgecolor=WHITE,
             markeredgewidth=2, label="부정 종료 (N=96, 14.0%)", zorder=5)
    ax1.fill_between(x_bins, pos_traj, neg_traj, alpha=0.06, color=ORANGE)

    for xi, (pv, nv) in enumerate(zip(pos_traj, neg_traj)):
        if xi in [0, 3, 6, 9]:
            ax1.text(xi, pv + 0.012, f"{pv:+.3f}", ha="center", fontsize=8,
                     color="#2E7D32", fontweight="bold")
            ax1.text(xi, nv - 0.018, f"{nv:+.3f}", ha="center", fontsize=8,
                     color="#C62828", fontweight="bold")

    for x_lo, x_hi, label in stage_bounds:
        mid = (x_lo + x_hi) / 2
        ax1.axvspan(x_lo - 0.5, x_hi + 0.5, alpha=0.03,
                    color=NAVY if label in ["탐색", "결과제시"] else ORANGE)
        ax1.text(mid, 0.24, label, ha="center", fontsize=10, color=GRAY, fontstyle="italic")

    ax1.annotate("부정: 30~40% 구간\n감성 최저(-0.147)",
                 xy=(3, -0.147), xytext=(4.5, -0.20),
                 fontsize=8.5, color="#C62828",
                 arrowprops=dict(arrowstyle="->", color="#C62828", lw=1.2),
                 bbox=dict(boxstyle="round,pad=0.3", fc="#FFF0F0", ec="#C62828", alpha=0.9))
    ax1.annotate("긍정: 종료 구간\n급반전(+0.192)",
                 xy=(9, 0.192), xytext=(7.2, 0.22),
                 fontsize=8.5, color="#2E7D32",
                 arrowprops=dict(arrowstyle="->", color="#2E7D32", lw=1.2),
                 bbox=dict(boxstyle="round,pad=0.3", fc="#F0FFF0", ec="#2E7D32", alpha=0.9))
    ax1.annotate("분기점: 50~60% 이후\n궤적 교차 시작",
                 xy=(5, -0.075), xytext=(5, 0.08),
                 fontsize=8.5, color=NAVY,
                 arrowprops=dict(arrowstyle="->", color=NAVY, lw=1.2),
                 bbox=dict(boxstyle="round,pad=0.3", fc="#F0F4FF", ec=NAVY, alpha=0.9))

    ax1.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax1.set_ylabel("감성 점수 (Valence)", fontsize=10, color="#555555")
    ax1.set_title("종료 감성 그룹별 감정 궤적 비교",
                  fontsize=13, fontweight="bold", color="#222222", pad=12)
    ax1.legend(fontsize=9.5, framealpha=0.9, loc="upper left")
    ax1.set_xticks(x_bins)
    ax1.set_xticklabels(x_pct, fontsize=9)
    ax1.set_ylim(-0.25, 0.28)
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)
    ax1.spines["left"].set_color(LIGHT_GRAY)
    ax1.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig17_trajectory_trigger.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig17"] = path
    print(f"  [fig17] 궤적+트리거 저장")

    # ─── fig18: 연령대별 감정 궤적 ───────────────────────────────────────────
    age_trajs = {
        "20~39세": [-0.078, -0.129, -0.106, -0.119, -0.116, -0.062, -0.065, -0.090, -0.013, +0.132],
        "40~49세": [-0.072, -0.101, -0.098, -0.128, -0.127, -0.080, -0.099, -0.087, -0.040, +0.110],
        "50~64세": [-0.056, -0.115, -0.115, -0.099, -0.079, -0.096, -0.066, -0.098, -0.037, +0.100],
        "65~74세": [-0.065, -0.125, -0.125, -0.120, -0.109, -0.107, -0.106, -0.085, -0.062, +0.055],
        "75세이상": [-0.106, -0.095, -0.149, -0.116, -0.038, -0.127, -0.097, +0.011, +0.072, +0.085],
    }
    age_colors = ["#2E7D32", "#1565C0", ORANGE, "#C62828", "#6A1B9A"]
    age_markers = ["o", "s", "D", "^", "v"]

    fig, ax = plt.subplots(figsize=(11, 5.5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    for i, (age, traj) in enumerate(age_trajs.items()):
        ax.plot(x_bins, traj, color=age_colors[i], linewidth=2, marker=age_markers[i],
                markersize=8, markerfacecolor=age_colors[i], markeredgecolor=WHITE,
                markeredgewidth=1.5, label=age, zorder=5 - i, alpha=0.85)

    # 단계 구분
    for x_lo, x_hi, label in stage_bounds:
        mid = (x_lo + x_hi) / 2
        ax.text(mid, 0.17, label, ha="center", fontsize=10, color=GRAY, fontstyle="italic")

    # 주요 발견 주석
    ax.annotate("65~74세: 종료 회복\n가장 미약(+0.055)",
                xy=(9, 0.055), xytext=(7, 0.14),
                fontsize=8.5, color="#C62828",
                arrowprops=dict(arrowstyle="->", color="#C62828", lw=1.2),
                bbox=dict(boxstyle="round,pad=0.3", fc="#FFF0F0", ec="#C62828", alpha=0.9))
    ax.annotate("20~39세: 가장\n빠른 회복",
                xy=(9, 0.132), xytext=(8.2, -0.16),
                fontsize=8.5, color="#2E7D32",
                arrowprops=dict(arrowstyle="->", color="#2E7D32", lw=1.2),
                bbox=dict(boxstyle="round,pad=0.3", fc="#F0FFF0", ec="#2E7D32", alpha=0.9))

    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax.set_xticks(x_bins)
    ax.set_xticklabels(x_pct, fontsize=9)
    ax.set_ylabel("감성 점수 (Valence)", fontsize=10, color="#555555")
    ax.set_title("연령대별 감정 궤적 비교 — 상담 진행에 따른 감성 변화",
                 fontsize=13, fontweight="bold", color="#222222", pad=12)
    ax.legend(fontsize=9, framealpha=0.9, loc="lower left",
              ncol=5, bbox_to_anchor=(0, -0.18), borderaxespad=0)
    ax.set_ylim(-0.20, 0.20)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig18_age_trajectory.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig18"] = path
    print(f"  [fig18] 연령대별 궤적 저장")

    # ─── fig19: 제품군별 감정 궤적 ───────────────────────────────────────────
    prod_trajs = {
        "주방가전":       [-0.074, -0.120, -0.128, -0.105, -0.089, -0.101, -0.101, -0.071, -0.039, +0.102],
        "생활가전":       [-0.058, -0.129, -0.103, -0.103, -0.123, -0.082, -0.072, -0.093, -0.021, +0.096],
        "TV/AV":          [-0.055, -0.100, -0.080, -0.131, -0.093, -0.114, -0.078, -0.134, -0.068, +0.075],
        "에어컨/에어케어": [-0.095, -0.099, -0.087, -0.125, -0.125, -0.060, -0.051, -0.129, -0.047, +0.092],
    }
    prod_colors = [NAVY, "#2E7D32", ORANGE, "#C62828"]
    prod_markers = ["o", "s", "D", "^"]

    fig, ax = plt.subplots(figsize=(11, 5.5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")

    for i, (prod, traj) in enumerate(prod_trajs.items()):
        ax.plot(x_bins, traj, color=prod_colors[i], linewidth=2, marker=prod_markers[i],
                markersize=8, markerfacecolor=prod_colors[i], markeredgecolor=WHITE,
                markeredgewidth=1.5, label=prod, zorder=5 - i, alpha=0.85)

    for x_lo, x_hi, label in stage_bounds:
        mid = (x_lo + x_hi) / 2
        ax.text(mid, 0.14, label, ha="center", fontsize=10, color=GRAY, fontstyle="italic")

    ax.annotate("70~80% 공통 하락\n(TV/AV·에어컨·생활)",
                xy=(7, -0.134), xytext=(5.5, -0.18),
                fontsize=8.5, color=ORANGE,
                arrowprops=dict(arrowstyle="->", color=ORANGE, lw=1.2),
                bbox=dict(boxstyle="round,pad=0.3", fc="#FFF8F0", ec=ORANGE, alpha=0.9))
    ax.annotate("주방가전:\n단일 최저(-0.128)",
                xy=(2, -0.128), xytext=(0, -0.18),
                fontsize=8.5, color=NAVY,
                arrowprops=dict(arrowstyle="->", color=NAVY, lw=1.2),
                bbox=dict(boxstyle="round,pad=0.3", fc="#F0F4FF", ec=NAVY, alpha=0.9))

    ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
    ax.set_xticks(x_bins)
    ax.set_xticklabels(x_pct, fontsize=9)
    ax.set_ylabel("감성 점수 (Valence)", fontsize=10, color="#555555")
    ax.set_title("제품군별 감정 궤적 비교 — 상담 진행에 따른 감성 변화",
                 fontsize=13, fontweight="bold", color="#222222", pad=12)
    ax.legend(fontsize=9, framealpha=0.9, loc="lower left",
              ncol=4, bbox_to_anchor=(0, -0.18), borderaxespad=0)
    ax.set_ylim(-0.20, 0.17)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig19_product_trajectory.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig19"] = path
    print(f"  [fig19] 제품군별 궤적 저장")

    # ─── fig20: 감정 그룹별 구성 변화 (Stacked Area) — 전체 ──────────────────
    grp_names  = ["감사/만족", "안정/중립", "불안/걱정", "불만/짜증", "혼란/당황"]
    grp_colors_list = ["#5B9A6B", "#6B8EB5", "#D4A84B", "#B56B6B", "#8B7BAA"]
    grp_color_map = dict(zip(grp_names, grp_colors_list))

    all_ratios = {
        "감사/만족": [14.9, 11.5, 10.4, 12.7, 14.4, 16.1, 15.8, 16.5, 23.0, 41.1],
        "안정/중립": [37.8, 37.5, 39.5, 36.3, 34.3, 32.8, 33.3, 30.2, 30.4, 33.1],
        "불안/걱정": [22.0, 19.0, 22.4, 21.3, 19.9, 21.0, 24.0, 23.9, 20.6, 11.6],
        "불만/짜증": [10.5, 18.2, 14.0, 17.3, 17.4, 17.5, 12.3, 14.2, 12.4,  8.6],
        "혼란/당황": [14.6, 13.7, 13.8, 12.3, 14.0, 12.5, 14.4, 15.2, 13.7,  5.5],
    }

    fig, ax = plt.subplots(figsize=(11, 5.5))
    fig.patch.set_facecolor(WHITE)
    ax.set_facecolor("#FAFAFA")
    x_area = np.arange(10)
    bottom = np.zeros(10)
    for gname, gcol in zip(grp_names, grp_colors_list):
        vals = np.array(all_ratios[gname])
        ax.fill_between(x_area, bottom, bottom + vals, color=gcol, alpha=0.75, label=gname)
        mid = bottom + vals / 2
        for xi in [0, 4, 9]:
            if vals[xi] >= 8:
                ax.text(xi, mid[xi], f"{vals[xi]:.0f}%", ha="center", va="center",
                        fontsize=7.5, fontweight="bold", color="white")
        bottom += vals
    ax.set_xticks(x_area)
    ax.set_xticklabels([f"{i*10}%" for i in range(10)], fontsize=9)
    ax.set_ylabel("비율 (%)", fontsize=10, color="#555555")
    ax.set_title("상담 진행에 따른 감정 그룹 구성 변화 (전체)",
                 fontsize=13, fontweight="bold", color="#222222", pad=12)
    ax.set_ylim(0, 100)
    ax.set_xlim(-0.3, 9.3)
    ax.legend(fontsize=9, framealpha=0.9, loc="upper right", ncol=5,
              bbox_to_anchor=(1, -0.08), borderaxespad=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(LIGHT_GRAY)
    ax.spines["bottom"].set_color(LIGHT_GRAY)
    for x_lo, x_hi, label in stage_bounds:
        ax.text((x_lo + x_hi) / 2, 103, label, ha="center", fontsize=9,
                color=GRAY, fontstyle="italic", clip_on=False)
    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig20_emotion_group_area.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig20"] = path
    print(f"  [fig20] 감정 그룹별 면적 차트 저장")

    # ─── fig21: 감정 그룹별 구성 변화 — 긍정 vs 부정 종료 비교 ───────────────
    pos_ratios = {
        "감사/만족": [15.1, 11.6, 12.0, 15.7, 15.0, 16.5, 17.8, 17.5, 25.8, 56.2],
        "안정/중립": [36.7, 37.7, 36.8, 34.6, 32.5, 28.9, 31.9, 28.9, 30.3, 28.7],
        "불안/걱정": [21.3, 20.1, 23.9, 21.4, 21.3, 23.9, 23.5, 24.4, 20.2,  6.4],
        "불만/짜증": [11.8, 17.5, 13.8, 15.7, 18.1, 17.3, 12.8, 14.0, 12.8,  5.3],
        "혼란/당황": [14.9, 13.2, 13.5, 12.6, 13.2, 13.4, 14.0, 15.2, 10.9,  3.5],
    }
    neg_ratios = {
        "감사/만족": [11.8, 12.9, 10.7,  9.5, 16.7, 19.2, 15.1, 16.3, 26.6, 16.5],
        "안정/중립": [36.6, 37.1, 38.5, 34.2, 32.7, 34.0, 35.5, 30.7, 25.2, 34.5],
        "불안/걱정": [26.0, 13.4, 21.3, 15.8, 14.2, 18.6, 22.4, 20.5, 16.5, 22.1],
        "불만/짜증": [11.0, 22.7, 17.2, 24.1, 21.6, 16.7, 11.8, 18.1, 15.1, 18.5],
        "혼란/당황": [14.2, 13.9, 12.4, 16.5, 14.8, 11.5, 15.1, 14.5, 16.5,  8.4],
    }
    fig, (ax_pos, ax_neg) = plt.subplots(1, 2, figsize=(14, 5.5), sharey=True)
    fig.patch.set_facecolor(WHITE)
    for ax, ratios, title_suffix in [
        (ax_pos, pos_ratios, "긍정 종료 (N=327, 47.7%)"),
        (ax_neg, neg_ratios, "부정 종료 (N=96, 14.0%)"),
    ]:
        ax.set_facecolor("#FAFAFA")
        bottom = np.zeros(10)
        for gname, gcol in zip(grp_names, grp_colors_list):
            vals = np.array(ratios[gname])
            ax.fill_between(x_area, bottom, bottom + vals, color=gcol, alpha=0.75, label=gname)
            mid = bottom + vals / 2
            for xi in [0, 4, 9]:
                if vals[xi] >= 10:
                    ax.text(xi, mid[xi], f"{vals[xi]:.0f}%", ha="center", va="center",
                            fontsize=7.5, fontweight="bold", color="white")
            bottom += vals
        ax.set_xticks(x_area)
        ax.set_xticklabels([f"{i*10}%" for i in range(10)], fontsize=8.5)
        ax.set_title(title_suffix, fontsize=12, fontweight="bold", color="#222222", pad=10)
        ax.set_ylim(0, 100)
        ax.set_xlim(-0.3, 9.3)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_color(LIGHT_GRAY)
        ax.spines["bottom"].set_color(LIGHT_GRAY)
    ax_pos.set_ylabel("비율 (%)", fontsize=10, color="#555555")
    handles, labels = ax_pos.get_legend_handles_labels()
    fig.legend(handles, labels, fontsize=9, framealpha=0.9,
               loc="lower center", ncol=5, bbox_to_anchor=(0.5, -0.02))
    fig.suptitle("종료 그룹별 감정 구성 변화 비교 (Stacked Area)",
                 fontsize=14, fontweight="bold", color="#222222", y=1.02)
    fig.text(0.5, 0.97, "※ 중립 종료(N=263, 38.3%): 종료 감성 -0.05~+0.05 구간은 별도 표시 생략",
             ha="center", fontsize=9, color=GRAY, fontstyle="italic")
    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig21_emotion_group_area_compare.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig21"] = path
    print(f"  [fig21] 긍정/부정 감정 그룹 면적 비교 저장")

    # ─── fig22: 감정 색상 궤적 — 전체/긍정/부정 (개별 콜 대시보드 스타일) ────
    nn_groups = ["감사/만족", "불안/걱정", "불만/짜증", "혼란/당황"]
    all_vals_t = [-0.069, -0.118, -0.114, -0.112, -0.103, -0.095, -0.086, -0.089, -0.038, +0.098]
    all_doms   = ["불안/걱정"]*8 + ["감사/만족"]*2
    pos_vals_t = [-0.077, -0.119, -0.114, -0.093, -0.103, -0.101, -0.080, -0.087, -0.024, +0.192]
    pos_doms   = ["불안/걱정"]*8 + ["감사/만족"]*2
    neg_vals_t = [-0.059, -0.092, -0.113, -0.147, -0.104, -0.049, -0.078, -0.111, -0.046, -0.071]
    neg_doms   = ["불안/걱정", "불만/짜증", "불안/걱정", "불만/짜증", "불만/짜증",
                  "감사/만족", "불안/걱정", "불안/걱정", "감사/만족", "불안/걱정"]

    from matplotlib.lines import Line2D

    def _plot_emotion_traj(ax, x, vals, doms, title, show_ylabel=True):
        ax.set_facecolor("#FAFAFA")
        ax.plot(x, vals, color="#CCCCCC", linewidth=2, zorder=3)
        for xi, (v, d) in enumerate(zip(vals, doms)):
            c = grp_color_map.get(d, "#999999")
            ax.scatter(xi, v, color=c, s=120, zorder=5, edgecolors="white", linewidths=1.8)
            offset = 0.015 if v >= 0 else -0.020
            ax.text(xi, v + offset, f"{v:+.3f}", ha="center", fontsize=7.5,
                    fontweight="bold", color=c)
        for x_lo, x_hi, label in stage_bounds:
            ax.text((x_lo + x_hi) / 2, 0.26, label, ha="center", fontsize=9,
                    color=GRAY, fontstyle="italic")
        ax.axhline(0, color=LIGHT_GRAY, linewidth=0.9, linestyle="--")
        ax.set_xticks(x)
        ax.set_xticklabels([f"{i*10}%" for i in range(10)], fontsize=8.5)
        if show_ylabel:
            ax.set_ylabel("Valence", fontsize=10, color="#555555")
        ax.set_title(title, fontsize=11, fontweight="bold", color="#222222", pad=10)
        ax.set_ylim(-0.22, 0.30)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_color(LIGHT_GRAY)
        ax.spines["bottom"].set_color(LIGHT_GRAY)

    fig, axes = plt.subplots(1, 3, figsize=(16, 5), sharey=True)
    fig.patch.set_facecolor(WHITE)
    _plot_emotion_traj(axes[0], x_bins, all_vals_t, all_doms, "전체 (N=686)", True)
    _plot_emotion_traj(axes[1], x_bins, pos_vals_t, pos_doms, "긍정 종료 (N=327, 47.7%)", False)
    _plot_emotion_traj(axes[2], x_bins, neg_vals_t, neg_doms, "부정 종료 (N=96, 14.0%)", False)
    legend_el = [Line2D([0], [0], marker="o", color="w", markerfacecolor=grp_color_map[g],
                        markersize=10, label=g) for g in grp_color_map]
    fig.legend(handles=legend_el, fontsize=9, framealpha=0.9,
               loc="lower center", ncol=5, bbox_to_anchor=(0.5, -0.04))
    fig.suptitle("감정 그룹 색상 궤적 — 종료 유형별 비교",
                 fontsize=14, fontweight="bold", color="#222222", y=1.02)
    fig.text(0.5, 0.97, "※ 중립 종료(N=263, 38.3%): 종료 감성 -0.05~+0.05 구간은 별도 표시 생략",
             ha="center", fontsize=9, color=GRAY, fontstyle="italic")
    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig22_emotion_color_trajectory.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig22"] = path
    print(f"  [fig22] 감정 색상 궤적 (종료 유형별) 저장")

    # ─── fig23: 감정 색상 궤적 — 연령대별 ────────────────────────────────────
    age_traj_data = {
        "20~39세":  ([-0.078,-0.129,-0.106,-0.119,-0.116,-0.062,-0.065,-0.090,-0.013,+0.132],
                     ["불안/걱정"]*8 + ["감사/만족"]*2),
        "40~49세":  ([-0.072,-0.101,-0.098,-0.128,-0.127,-0.080,-0.099,-0.087,-0.040,+0.110],
                     ["불안/걱정"]*8 + ["감사/만족"]*2),
        "50~64세":  ([-0.056,-0.115,-0.115,-0.099,-0.079,-0.096,-0.066,-0.098,-0.037,+0.100],
                     ["불안/걱정"]*4+["감사/만족"]+["불안/걱정"]*3+["감사/만족"]*2),
        "65~74세":  ([-0.065,-0.125,-0.125,-0.120,-0.109,-0.107,-0.106,-0.085,-0.062,+0.055],
                     ["불안/걱정","불만/짜증","불안/걱정","불안/걱정","불안/걱정",
                      "불만/짜증","불안/걱정","불안/걱정","불안/걱정","감사/만족"]),
    }
    fig, axes = plt.subplots(1, 4, figsize=(18, 5), sharey=True)
    fig.patch.set_facecolor(WHITE)
    for i, (age, (vals, doms)) in enumerate(age_traj_data.items()):
        _plot_emotion_traj(axes[i], x_bins, vals, doms, age, i == 0)
    fig.legend(handles=legend_el, fontsize=9, framealpha=0.9,
               loc="lower center", ncol=5, bbox_to_anchor=(0.5, -0.04))
    fig.suptitle("감정 그룹 색상 궤적 — 연령대별 비교",
                 fontsize=14, fontweight="bold", color="#222222", y=1.02)
    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig23_emotion_color_trajectory_age.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig23"] = path
    print(f"  [fig23] 감정 색상 궤적 (연령대별) 저장")

    # ─── fig24: 감정 색상 궤적 — 제품군별 ────────────────────────────────────
    prod_traj_data = {
        "주방가전":       ([-0.074,-0.120,-0.128,-0.105,-0.089,-0.101,-0.101,-0.071,-0.039,+0.102],
                          ["불안/걱정","불만/짜증","불안/걱정","불안/걱정","불안/걱정",
                           "불안/걱정","불안/걱정","불안/걱정","감사/만족","감사/만족"]),
        "생활가전":       ([-0.058,-0.129,-0.103,-0.103,-0.123,-0.082,-0.072,-0.093,-0.021,+0.096],
                          ["불안/걱정"]*8+["감사/만족"]*2),
        "TV/AV":          ([-0.055,-0.100,-0.080,-0.131,-0.093,-0.114,-0.078,-0.134,-0.068,+0.075],
                          ["불안/걱정"]*8+["혼란/당황","감사/만족"]),
        "에어컨/에어케어": ([-0.095,-0.099,-0.087,-0.125,-0.125,-0.060,-0.051,-0.129,-0.047,+0.092],
                          ["불안/걱정","불만/짜증","불안/걱정","불안/걱정","불만/짜증",
                           "혼란/당황","불안/걱정","혼란/당황","감사/만족","감사/만족"]),
    }
    fig, axes = plt.subplots(1, 4, figsize=(18, 5), sharey=True)
    fig.patch.set_facecolor(WHITE)
    for i, (prod, (vals, doms)) in enumerate(prod_traj_data.items()):
        _plot_emotion_traj(axes[i], x_bins, vals, doms, prod, i == 0)
    fig.legend(handles=legend_el, fontsize=9, framealpha=0.9,
               loc="lower center", ncol=5, bbox_to_anchor=(0.5, -0.04))
    fig.suptitle("감정 그룹 색상 궤적 — 제품군별 비교",
                 fontsize=14, fontweight="bold", color="#222222", y=1.02)
    fig.tight_layout()
    path = os.path.join(OUTPUT_DIR, "fig24_emotion_color_trajectory_product.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=WHITE)
    plt.close(fig)
    figs["fig24"] = path
    print(f"  [fig24] 감정 색상 궤적 (제품군별) 저장")

    print(f"\n그래프 {len(figs)}개 생성 완료 → {OUTPUT_DIR}")
    return figs


# ═══════════════════════════════════════════════════════════════════════════════
# REPORT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def build_report(figs: dict):
    doc = Document()

    # 기본 페이지 여백 설정
    from docx.oxml.ns import qn as _qn
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── 표지 (generate_report_voice.py 스타일) ─────────────────────────────
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("고객 상담 음성 감성 분석 보고서")
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BERT + SER + 음향 특징 3-way 융합 분석\n인사이트 리포트")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026.04")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 1. 분석 개요
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "1. 분석 개요", level=1)
    _add_para(doc,
        "본 분석은 고객 상담 음성 686건(고객 발화 12,072건)을 대상으로 "
        "BERT 기반 텍스트 감성 분석, 음향 특징 추출, SER(Speech Emotion Recognition) 모델을 "
        "융합한 3-way 감성 분석 파이프라인을 적용하여 상담 전 과정의 고객 감성을 정량화한 결과를 담습니다.")

    doc.add_paragraph()
    _add_heading(doc, "1.1 분석 범위 및 데이터", level=2)
    _add_table(doc,
        headers=["항목", "내용"],
        rows=[
            ["분석 대상 콜 수",    "686건"],
            ["분석 대상 발화",     "고객 발화 12,072건"],
            ["짧은 발화(<3어절)",  "5,322건 (고객 발화의 44%)"],
            ["분석 기간",          "2026년 1~3월"],
            ["분석 언어",          "한국어 (콜센터 상담)"],
        ],
        col_widths=[5, 10],
    )

    doc.add_paragraph()
    _add_heading(doc, "1.2 핵심 요약", level=2)
    _add_insight_box(doc, [
        "3-way 융합(BERT 50% + SER 30% + 음향 20%) 모델이 단일 모달리티 대비 NPS 상관 r=0.262로 최고 성능",
        "상담 단계별 감성: 초기 부정(-0.043) → 탐색 최저(-0.121) → 종료 긍정(+0.132) 패턴 확인",
        "전체 콜의 95%에서 부정 감성 구간 후 회복 발생 — '감정 회복'이 NPS에 직결",
        "에너지 높은 '네'(+0.140)와 낮은 '네'(-0.069)는 감성 점수 0.209포인트 차이",
        "짧은 응답(+0.175)은 해결 안내(+0.251) 대비 감정 개선 효과 32% 낮음 — 실질적 응대 필수",
        "종료 감성 긍정 그룹 NPS=8.04 vs 부정 그룹 NPS=7.45 (격차 0.59점)",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 2. 방법론
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "2. 방법론", level=1)

    _add_heading(doc, "2.1 분석 파이프라인", level=2)
    _add_para(doc,
        "본 파이프라인은 음성 데이터를 기반으로 3개 경로의 분석을 병렬 수행합니다.")

    _add_table(doc,
        headers=["분석 경로", "모델/방법", "출력", "역할"],
        rows=[
            ["음향 특징 추출", "F0, Energy, MFCC, Jitter, Shimmer 등 10+ 특징",
             "음성 Valence (-1~+1)", "음고·음량·음질의 감성 반영"],
            ["텍스트 감성 분석", "Whisper STT → BERT 60-class (klue-bert-base)",
             "텍스트 Valence + 감정 라벨", "발화 내용의 의미적 감성 분류"],
            ["음성 감정 인식 (SER)", "wav2vec2-base-superb-er (4-class)",
             "neu/hap/ang/sad 확률", "음성 자체의 감정 상태 직접 분류"],
        ],
        col_widths=[3, 5, 4, 3.5],
    )

    _add_para(doc,
        "분석 결과는 두 가지 형태로 산출됩니다. "
        "감정 그룹(감사/만족, 안정/중립, 불안/걱정, 불만/짜증, 혼란/당황)은 "
        "BERT 감정 라벨을 상담 맥락에 맞게 5개 그룹으로 재매핑하여 분류합니다. "
        "감성 점수(Valence, -1~+1)는 BERT·SER·음향 세 경로를 가중 융합하여 산출하며, "
        "긴 발화(3어절 이상)는 BERT 50% + SER 30% + 음향 20%, "
        "짧은 발화('네', '예' 등 3어절 미만)는 텍스트 분석이 불가하므로 SER 70% + 음향 30%로 구성합니다. "
        "짧은 발화의 경우 감정 그룹은 SER 4-class 결과(hap→감사/만족, ang→불만/짜증, sad→불안/걱정, neu→안정/중립)로 대체하며, "
        "이후 LLM 교차검증을 통해 BERT와 음향 분석 결과가 상충하는 발화에 대해 "
        "최종 판정을 보정합니다.")

    doc.add_paragraph()
    _add_heading(doc, "2.2 SER 모델 도입", level=2)
    _add_para(doc,
        "기존 텍스트+음향 2-way 융합에서 SER(Speech Emotion Recognition) 모델을 "
        "추가하여 3-way 융합으로 업데이트하였습니다. "
        "SER 모델은 특히 '네', '예', '아' 등 짧은 발화에서 텍스트 기반 분석이 불가한 "
        "경우에 핵심 역할을 담당합니다.")

    doc.add_paragraph()
    _add_table(doc,
        headers=["항목", "내용"],
        rows=[
            ["SER 모델",          "superb/wav2vec2-base-superb-er"],
            ["출력 클래스",        "4-class: neutral / happy / angry / sad"],
            ["긴 발화 융합 가중치", "BERT 50% + SER 30% + 음향특징 20%"],
            ["짧은 발화 (<3어절)", "SER 70% + 음향특징 30% (텍스트 미사용)"],
            ["SER 감성 변환",      "4-class 확률을 감성 점수로 변환 (초기 설정값)"],
        ],
        col_widths=[5.5, 10],
    )

    doc.add_paragraph()
    _add_heading(doc, "2.3 상담 단계 구분", level=2)
    _add_para(doc,
        "각 콜을 발화 시퀀스 비율로 5단계(초기 10%, 탐색 25%, 해결시도 35%, 결과제시 20%, 종료 10%)로 구분하였습니다. "
        "단계별 고객 발화의 평균 감성 점수를 계산하여 상담 흐름에 따른 감성 궤적을 추적합니다.")

    _add_table(doc,
        headers=["단계", "비율", "주요 내용", "분석 포인트"],
        rows=[
            ["초기",    "10%", "인사, 문의 내용 확인",      "초기 감성 수준 파악"],
            ["탐색",    "25%", "문제 상세 파악, 고객 정보 확인", "부정 감성 최고점 구간"],
            ["해결시도", "35%", "솔루션 제시, 조회·처리",     "감성 회복 시작 여부"],
            ["결과제시", "20%", "처리 결과 안내",            "수용/거절 반응 포착"],
            ["종료",    "10%", "마무리 인사, 확인",          "최종 감성 → NPS 예측력"],
        ],
        col_widths=[2, 1.5, 5, 5],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 3. 전체 감성 분포
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "3. 전체 감성 분포", level=1)

    _add_heading(doc, "3.1 고객 감정 그룹 분포", level=2)
    _add_para(doc,
        "전체 고객 발화 12,072건을 5개 감정 그룹으로 분류한 결과, "
        "'안정/중립' 그룹이 34.8%로 가장 높았으며, '불안/걱정'(20.1%)과 "
        "'불만/짜증'(13.8%)을 합산하면 부정 발화가 전체의 33.9%를 차지합니다.")

    if "fig02" in figs:
        _add_figure(doc, figs["fig02"], width_inches=5.8,
                    caption="[그림 2] 고객 감정 그룹 분포 (좌: 비율, 우: 건수)")

    doc.add_paragraph()
    _add_table(doc,
        headers=["감정 그룹", "발화 수", "비율", "특성"],
        rows=[
            ["감사/만족",  "2,245",  "18.6%",  "상담 종료 시 집중"],
            ["안정/중립",  "4,199",  "34.8%",  "전 구간 베이스라인"],
            ["불안/걱정",  "2,429",  "20.1%",  "탐색·해결시도 집중"],
            ["불만/짜증",  "1,668",  "13.8%",  "해결 불가 직후 급증"],
            ["혼란/당황",  "1,526",  "12.6%",  "갑작스러운 안내 직후"],
        ],
        col_widths=[3.5, 2, 2, 7],
    )

    doc.add_paragraph()
    _add_heading(doc, "3.2 상담 단계별 감성 궤적", level=2)
    _add_para(doc,
        "상담 단계별 평균 감성 점수를 분석한 결과, '탐색' 단계에서 감성 점수가 "
        "가장 낮은 수준(-0.121)을 기록하였으며, '종료' 단계에서는 유일하게 양의 감성 점수(+0.132)로 "
        "회복되는 패턴이 전 콜에 걸쳐 공통적으로 관찰됩니다.")

    if "fig03" in figs:
        _add_figure(doc, figs["fig03"], width_inches=5.8,
                    caption="[그림 3] 상담 단계별 평균 감성 점수 궤적")

    doc.add_paragraph()
    _add_heading(doc, "3.3 단계별 감성 변화량 분석", level=2)
    _add_para(doc,
        "각 단계 간 감성 변화량을 워터폴 차트로 분해하면, "
        "초기→탐색 구간에서 -0.078의 급락이 발생하고 이후 해결시도(+0.024), "
        "결과제시(+0.045)를 거쳐 종료에서 +0.183의 대폭 상승이 일어납니다. "
        "전체 회복 아크(탐색 최저 → 종료)는 0.253포인트이며, "
        "이 중 70%가 종료 단계에 집중됩니다.")

    if "fig15" in figs:
        _add_figure(doc, figs["fig15"], width_inches=5.5,
                    caption="[그림 15] 단계별 감성 점수 워터폴 — 변화량 분해")

    doc.add_paragraph()
    _add_heading(doc, "3.4 감정 그룹별 단계 분포 패턴", level=2)
    _add_para(doc,
        "5개 감정 그룹이 상담 단계에서 어떻게 분포하는지를 레이더차트로 나타냅니다. "
        "'감사/만족'은 종료 단계에 27%가 집중되어 긍정 감정이 마무리 시점에 발현되는 반면, "
        "'불만/짜증'(38%)과 '혼란/당황'(35%)은 해결시도 단계에 최대 밀집됩니다. "
        "이는 해결시도 단계가 상담의 '감정 분기점'임을 시사합니다 (→ 9장 응대 패턴 분석과 연결).")

    if "fig14" in figs:
        _add_figure(doc, figs["fig14"], width_inches=5.0,
                    caption="[그림 14] 감정 그룹별 단계 분포 레이더차트")

    _add_table(doc,
        headers=["감정 그룹", "초기", "탐색", "해결시도", "결과제시", "종료"],
        rows=[
            ["감사/만족", "11%", "16%", "26%", "19%", "27% (최대)"],
            ["안정/중립", "14%", "30%", "32%", "14%", "10%"],
            ["불안/걱정", "12%", "30%", "36%", "17%", "5%"],
            ["불만/짜증", "9%",  "32%", "38% (최대)", "16%", "6%"],
            ["혼란/당황", "11%", "32%", "35%", "18%", "4%"],
        ],
        col_widths=[2.5, 1.5, 1.5, 2.5, 2.5, 3.5],
    )

    _add_insight_box(doc, [
        "탐색 단계가 모든 세그먼트에서 감성 최저점 → 상담사의 정보 수집 방식이 고객 감성에 직접 영향",
        "해결시도 단계에 부정 감정(불만 38%, 혼란 35%)이 최대 집중 → 이 단계의 상담사 응대가 결과를 결정 (→ 9장)",
        "감사/만족은 종료(27%)에 집중 — 마무리 인사·확인 단계가 긍정 감정 발현의 핵심",
        "종료 직전 +0.183 급상승: 전체 회복의 70%가 마지막 단계에 집중",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 4. 감성 전환 분석
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "4. 감성 전환 분석", level=1)

    _add_heading(doc, "4.1 전환 유형 분포", level=2)
    _add_para(doc,
        "감성 전환(전반부 vs 후반부 평균 비교)을 기준으로 분류한 결과, "
        "부정→긍정 회복 전환이 39.3%로 가장 높은 비율을 차지합니다. "
        "해결 불가 안내 후 빠른 회복(2턴 이내)이 63%로, "
        "대부분의 부정 감성은 상담사의 즉각적 대응으로 해소됩니다.")

    if "fig04" in figs:
        _add_figure(doc, figs["fig04"], width_inches=5.0,
                    caption="[그림 4] 감성 전환 유형별 분포")

    doc.add_paragraph()
    _add_heading(doc, "4.2 전환 트리거 분석", level=2)
    _add_para(doc,
        "부정→긍정 전환(3,018건) 직전 상담사 발화를 키워드 기반으로 분류하면, "
        "솔루션/처리 안내(22.0%)가 가장 높은 비중을 차지합니다. "
        "기타 안내(57.6%)는 특정 키워드에 해당하지 않는 일반적 상황 설명으로, "
        "실질적 정보 전달 자체가 감성 전환에 기여함을 시사합니다.")

    _add_table(doc,
        headers=["전환 트리거", "비율", "설명"],
        rows=[
            ["솔루션/처리 안내",     "22.0%", "확인·조회·처리·접수 등 구체적 진행 안내"],
            ["공감/사과 멘트",       "8.4%",  "죄송합니다, 불편하셨겠습니다 등"],
            ["감사 인사",            "6.8%",  "기다려 주셔서 감사합니다 등"],
            ["기사/방문 일정 안내",  "5.1%",  "방문 날짜·시간 안내"],
            ["기타 안내",            "57.6%", "일반 상황 설명·정보 전달 등"],
        ],
        col_widths=[4.5, 2, 8],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 5. 세그먼트별 감성 분석
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "5. 세그먼트별 감성 분석", level=1)

    _add_heading(doc, "5.1 제품군별 감성 패턴", level=2)
    _add_para(doc,
        "에어컨/에어케어는 초기 감성 점수가 가장 낮아(-0.063) 교체·수리 등 "
        "긴급 민원 유입이 많음을 시사합니다. 주방가전은 탐색 단계 감성이 가장 낮아(-0.129) "
        "문제 파악 과정에서의 불편함이 큰 것으로 분석됩니다.")

    if "fig05" in figs:
        _add_figure(doc, figs["fig05"], width_inches=5.8,
                    caption="[그림 5] 제품군별 상담 단계 감성 점수")

    doc.add_paragraph()
    _add_heading(doc, "5.2 연령대별 감성 패턴", level=2)
    _add_para(doc,
        "연령대가 높을수록 초기 감성 점수가 낮고(-0.031 → -0.071) 종료 감성 회복폭도 "
        "상대적으로 작은 경향이 있습니다. 75세 이상 고령층의 경우 종료 감성이 "
        "+0.044로 다른 연령대 대비 낮아, 상담 후에도 충분한 해소가 이루어지지 않을 가능성이 있습니다.")

    if "fig06" in figs:
        _add_figure(doc, figs["fig06"], width_inches=5.8,
                    caption="[그림 6] 연령대별 주요 단계 감성 점수")

    doc.add_paragraph()
    _add_heading(doc, "5.3 제품군 × 연령대 교차 분석", level=2)
    _add_para(doc,
        "제품군과 연령대를 동시에 고려한 교차 히트맵에서, "
        "가장 부정적인 감성은 '에어컨/에어케어 × 75세 이상'(-0.118)이며, "
        "가장 양호한 감성은 '에어컨/에어케어 × 40~49세'(-0.031)입니다. "
        "같은 제품군 내에서도 연령대에 따라 최대 0.087포인트의 차이가 존재합니다.")

    if "fig13" in figs:
        _add_figure(doc, figs["fig13"], width_inches=5.8,
                    caption="[그림 13] 제품군 × 연령대 교차 감성 히트맵")

    _add_para(doc,
        "TV/AV × 65~74세(-0.100)와 생활가전 × 65~74세(-0.086) 역시 두드러지게 부정적입니다. "
        "반면 생활가전 × 20~39세(-0.035)는 상대적으로 양호하여, "
        "고령층일수록 고가/복잡 제품에서 상담 불만이 더 크게 나타남을 확인합니다. "
        "이 패턴은 3.4절의 '불만/짜증이 해결시도 단계에 집중'되는 현상과 연결되며, "
        "고령+고가 제품 조합에서 해결시도 단계의 상담사 응대 품질이 더 중요함을 시사합니다 (→ 9장).")

    _add_insight_box(doc, [
        "에어컨 × 75세 이상(-0.118): 가장 부정적 → 고령층 긴급 민원 전담 응대 필요",
        "동일 제품군 내 연령대별 최대 차이 0.087 → 연령대 맞춤 응대 스크립트 도입 권고",
        "주방가전: 전 연령대에서 일관적으로 부정적(-0.050~-0.075) → 제품 특성상 탐색 단계 공감 강화",
        "TV/AV × 65~74세(-0.100)는 2번째로 부정 → 고령층 AV 상담 시 쉬운 설명 필수",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 6. 짧은 발화 세부 감성 분석 (NEW)
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "6. 짧은 발화 세부 감성 분석", level=1)

    _add_para(doc,
        "짧은 발화('네', '예', '아' 등 3어절 미만)는 전체 고객 발화의 44%(5,322건)를 차지합니다. "
        "텍스트만으로는 감성 구분이 불가능하므로 SER 모델과 음향 특징(Energy, F0)을 통해 "
        "동일한 단어에 담긴 감성 차이를 분석하였습니다.")

    _add_heading(doc, "6.1 '네' 발화 음량별 감성 분석", level=2)
    _add_para(doc,
        "가장 빈번한 짧은 발화인 '네'(1,065건)를 Energy 중앙값으로 이분하면 "
        "음량 높은 '네'(밝은 톤)와 낮은 '네'(어두운 톤) 간 감성 점수 차이가 0.209포인트에 달합니다.")

    _add_table(doc,
        headers=["구분", "N", "감성 점수", "해석"],
        rows=[
            ["Energy 높은 '네' (밝은 톤)", "533", "+0.140", "긍정적 수용, 적극 동의"],
            ["Energy 낮은 '네' (어두운 톤)", "532", "-0.069", "소극적 응답, 불만 잠재"],
        ],
        col_widths=[5.5, 1.5, 2.5, 5],
    )

    doc.add_paragraph()
    _add_heading(doc, "6.2 단계별 '네' 감성 패턴", level=2)
    _add_para(doc,
        "상담 단계에 따라 동일한 '네' 발화의 감성 점수가 다르게 나타납니다. "
        "탐색 단계(-0.013)에서 가장 낮고, 초기(+0.092)와 종료(+0.113)에서 높아 "
        "단계별 맥락이 짧은 발화의 감성에 큰 영향을 미칩니다.")

    _add_table(doc,
        headers=["상담 단계", "감성 점수", "해석"],
        rows=[
            ["초기",   "+0.092", "첫 인사 후 긍정적 응대"],
            ["탐색",   "-0.013", "문제 파악 중 소극적 응답"],
            ["해결시도", "+0.027", "솔루션 듣는 중 중립~긍정"],
            ["종료",   "+0.113", "마무리 인사에 밝은 응답"],
        ],
        col_widths=[3, 3, 8.5],
    )

    doc.add_paragraph()
    _add_heading(doc, "6.3 맥락별 '네' 감성 패턴", level=2)
    _add_para(doc,
        "직전 상담사 발화 유형(맥락)에 따라 '네' 발화의 감성 점수가 명확히 달라집니다. "
        "비용 안내 후의 '네'(+0.014)는 가장 낮고, 사과/공감 후의 '네'(+0.079)는 가장 높습니다.")

    _add_table(doc,
        headers=["직전 맥락", "감성 점수", "해석"],
        rows=[
            ["사과/공감 후",   "+0.079", "위로받은 후 긍정 수용"],
            ["비용 안내 후",   "+0.014", "다른 맥락 대비 상대적으로 낮음"],
            ["감사/인사 후",   "+0.053", "마무리 인사에 긍정 호응"],
        ],
        col_widths=[3.5, 3, 8],
    )

    doc.add_paragraph()
    _add_heading(doc, "6.4 복합 발화 감정 분포", level=2)
    _add_para(doc,
        "짧은 복합 발화의 경우 구문에 따라 감정 그룹 분포가 뚜렷이 달라집니다. "
        "'네 감사합니다'는 감사/만족 비율이 65%인 반면, "
        "'네 네 네' 반복 발화는 불안/걱정 비율이 31%로 높아 "
        "조급함·촉구의 감성을 내포합니다.")

    if "fig11" in figs:
        _add_figure(doc, figs["fig11"], width_inches=6.2,
                    caption="[그림 11] 짧은 발화 '네' 세부 감성 분석 (음량·단계·맥락·구문)")

    doc.add_paragraph()
    _add_heading(doc, "6.5 짧은 발화 비율과 상담 결과의 관계", level=2)
    _add_para(doc,
        "콜당 짧은 발화 비율이 높을수록 NPS가 낮아지는 경향이 관찰됩니다. "
        "짧은 발화 비율 50% 이상 콜의 NPS는 7.76으로, 25% 미만 콜(7.95) 대비 0.19점 낮습니다. "
        "이는 고객이 소극적으로 응대하는 콜에서 만족도가 떨어질 수 있음을 시사하며, "
        "9장의 '짧은 응답이 감정 개선에 비효과적'이라는 발견과 방향이 일치합니다.")

    _add_table(doc,
        headers=["짧은 발화 비율", "N", "NPS", "종료 감성"],
        rows=[
            ["낮음 (<25%)",   "113", "7.95", "+0.114"],
            ["중간 (25~50%)", "315", "8.01", "+0.098"],
            ["높음 (50%+)",   "258", "7.76", "+0.104"],
        ],
        col_widths=[3.5, 2, 2, 7],
    )

    _add_insight_box(doc, [
        "동일한 '네'도 음량(Energy)에 따라 감성 점수 0.209포인트 차이 → SER 도입 전 텍스트 분석만으로는 탐지 불가",
        "'네 네 네' 반복은 불안/걱정 31% → 상담사가 인식하고 신속히 진행 상황 안내 필요 (→ 9장 응대 패턴)",
        "비용 안내 후 '네'(+0.014)는 다른 맥락 대비 가장 낮음 → 비용 설명 직후 공감 멘트 병행 권고",
        "짧은 발화 비율 50%+ 콜은 NPS 0.19 낮음 → 소극적 응대가 지속되면 실시간 경보 발동 권고",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 7. 컨설턴트 만족도 분석 (NEW)
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "7. 컨설턴트 만족도 분석", level=1)

    _add_para(doc,
        "고객이 상담 후 평가한 컨설턴트 만족도(1~5점), NPS(0~10점)와 "
        "음성 감성 분석 결과 간의 상관관계를 분석하였습니다. "
        "음성 감성은 텍스트 설문으로 포착되지 않는 실시간 감성 변화를 보완하는 지표로 활용될 수 있습니다.")

    _add_heading(doc, "7.1 상관관계 분석", level=2)
    _add_table(doc,
        headers=["지표쌍", "Pearson r", "해석"],
        rows=[
            ["컨설턴트만족도 ↔ 음성감성", "r = 0.229", "약한 양의 상관"],
            ["NPS ↔ 음성감성",           "r = 0.262", "약한~중간 양의 상관 (더 강함)"],
            ["컨설턴트만족도 ↔ NPS",      "r = 0.547", "중간 수준의 양의 상관"],
        ],
        col_widths=[5.5, 3, 6],
    )
    _add_para(doc,
        "NPS와 음성 감성의 상관(r=0.262)이 컨설턴트 만족도와의 상관(r=0.229)보다 "
        "높게 나타납니다. 이는 고객이 컨설턴트 자체에 대한 평가보다 상담 결과(해결 여부)에 더 "
        "민감하게 반응한다는 것을 시사합니다.")

    doc.add_paragraph()
    _add_heading(doc, "7.2 만족도 점수별 감성 및 NPS", level=2)
    _add_table(doc,
        headers=["컨설턴트 만족도", "N", "평균 감성 점수", "평균 NPS"],
        rows=[
            ["1점 (매우 불만족)", "9",   "-0.158", "2.33"],
            ["2점",               "15",  "-0.102", "3.53"],
            ["3점",               "28",  "-0.095", "5.50"],
            ["4점",               "111", "-0.081", "7.02"],
            ["5점 (매우 만족)",   "523", "-0.056", "8.45"],
        ],
        col_widths=[4.5, 1.5, 3.5, 3.5],
    )
    _add_para(doc,
        "만족도 1점 그룹의 감성 점수(-0.158)는 5점 그룹(-0.056) 대비 0.102포인트 낮으며, "
        "NPS 격차는 6.12점으로 매우 큽니다. 만족도가 높아질수록 감성과 NPS 모두 "
        "단조 증가하는 경향을 보여, 음성 감성이 만족도와 일관된 방향성을 확인합니다.")

    doc.add_paragraph()
    _add_heading(doc, "7.3 NPS-컨설턴트 만족도 괴리 분석", level=2)
    _add_para(doc,
        "NPS가 낮음에도(5점 이하) 컨설턴트 만족도가 높은(4점 이상) 콜이 67건 확인됩니다. "
        "이는 상담사 응대는 만족했지만 문제 자체가 해결되지 않았거나 "
        "정책적 제약으로 결과에 불만족한 케이스입니다.")

    _add_table(doc,
        headers=["유형", "건수", "해석"],
        rows=[
            ["NPS 낮음(5-) + 컨설턴트 높음(4+)", "67건", "상담사에는 만족, 결과에 불만"],
            ["NPS 높음(9+) + 컨설턴트 낮음(3-)", "8건",  "결과에 만족, 상담사 스타일 불호"],
        ],
        col_widths=[6.5, 2, 6],
    )

    doc.add_paragraph()
    _add_heading(doc, "7.4 만족도별 종료 감성 비교", level=2)
    _add_para(doc,
        "종료 단계 감성 점수는 컨설턴트 만족도와 명확한 차이를 보입니다. "
        "불만족(1~3점) 그룹의 종료 감성(+0.064)과 만족(5점) 그룹의 종료 감성(+0.144) 간 "
        "격차는 0.080포인트로, 종료 단계 감성이 설문 만족도를 예측하는 데 유효함을 보여줍니다.")

    _add_table(doc,
        headers=["만족도 그룹", "종료 단계 감성 점수", "차이"],
        rows=[
            ["불만족(1~3점)", "+0.064", ""],
            ["만족(5점)",     "+0.144", "+0.080 (만족 그룹이 높음)"],
        ],
        col_widths=[4, 4, 6.5],
    )

    if "fig10" in figs:
        _add_figure(doc, figs["fig10"], width_inches=5.8,
                    caption="[그림 10] 컨설턴트 만족도별 감성 점수 및 NPS")

    doc.add_paragraph()
    _add_heading(doc, "7.5 괴리 케이스 감성 궤적 심층 분석", level=2)
    _add_para(doc,
        "67건의 괴리 케이스(NPS 低 + 만족도 高)와 정상 케이스(NPS 高 + 만족도 高)의 "
        "단계별 감성 궤적을 비교하면, 괴리 케이스는 초기부터 감성이 현저히 낮습니다(-0.119 vs -0.025). "
        "이 격차(0.094)는 탐색 단계까지 유지되다가 해결시도 이후 수렴하며, "
        "두 그룹 모두 종료 시 긍정으로 마무리되지만 괴리 케이스는 약간 낮습니다(+0.124 vs +0.143).")

    _add_table(doc,
        headers=["단계", "괴리 케이스 (N=67)", "정상 케이스 (N=483)", "격차"],
        rows=[
            ["초기",    "-0.119", "-0.025", "0.094 (최대)"],
            ["탐색",    "-0.153", "-0.109", "0.044"],
            ["해결시도", "-0.096", "-0.091", "0.005 (수렴)"],
            ["결과제시", "-0.072", "-0.042", "0.030"],
            ["종료",    "+0.124", "+0.143", "0.019"],
        ],
        col_widths=[2.5, 4, 4, 4],
    )
    _add_para(doc,
        "해결시도 이후 궤적이 거의 수렴한다는 것은, 상담사가 문제 해결을 위해 노력한 것이 "
        "만족도(4+)로는 인정받았지만, 문제 자체의 해결 여부가 NPS에 반영되었음을 의미합니다. "
        "이 패턴은 3.4절의 '해결시도 단계에 부정 감정이 집중'되는 발견과 연결되며, "
        "5.3절의 '고령+고가 제품 조합'에서 특히 두드러집니다.")

    if "fig16" in figs:
        _add_figure(doc, figs["fig16"], width_inches=6.0,
                    caption="[그림 16] (좌) 괴리 vs 정상 궤적 비교 / (우) 상담유형별 회복폭")

    _add_insight_box(doc, [
        "괴리 케이스: 초기 감성 격차 0.094 → 입전 시점부터 민원 강도가 높은 건이 대부분",
        "해결시도 이후 궤적 수렴 → 상담사 노력은 만족도에 반영되나, NPS는 결과에 좌우",
        "반품-제품 상담유형이 회복폭 최대(+0.292) → 반품 승인이 감성 반전의 가장 강력한 트리거",
        "사용설명 상담: 초기 가장 부정(-0.084)이나 회복도 양호(+0.238) → 문제 해결이 명확한 유형",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 8. 세그먼트별 감정 궤적 (NEW)
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "8. 세그먼트별 감정 궤적", level=1)

    _add_para(doc,
        "종료 감성을 기준으로 전체 686건의 콜을 긍정(+0.05 초과, 327건·47.7%), "
        "중립(-0.05~+0.05, 263건·38.3%), 부정(-0.05 미만, 96건·14.0%) 그룹으로 분류하여 "
        "단계별 감성 궤적 차이를 분석하였습니다. "
        "이하 분석에서는 대비가 뚜렷한 긍정·부정 그룹 간 비교를 중심으로 서술합니다.")

    _add_heading(doc, "8.1 종료 감정 그룹별 단계 궤적 비교", level=2)
    _add_table(doc,
        headers=["단계", "긍정 종료 그룹 (N=327)", "부정 종료 그룹 (N=96)", "차이"],
        rows=[
            ["초기",    "-0.047", "-0.042", "-0.005"],
            ["탐색",    "-0.122", "-0.109", "-0.013"],
            ["해결시도", "-0.093", "-0.094", "+0.001"],
            ["결과제시", "-0.050", "-0.064", "+0.014"],
            ["종료",    "+0.215", "-0.082", "+0.297 ★"],
        ],
        col_widths=[2.5, 4, 4, 4],
    )
    _add_para(doc,
        "두 그룹의 초기~해결시도 구간 감성 차이는 0.015포인트 이내로 매우 작습니다. "
        "결정적 분기는 '결과제시' 이후, 특히 '종료' 단계에서 발생하며, "
        "이 구간에서 0.297포인트의 극명한 차이가 나타납니다.")

    doc.add_paragraph()
    _add_para(doc,
        "NPS 비교에서도 긍정 종료 그룹(NPS=8.04)이 부정 종료 그룹(NPS=7.45) 대비 "
        "0.59점 높아, 종료 단계 감성이 최종 NPS에 직접적으로 영향을 미침을 확인합니다.")

    doc.add_paragraph()
    _add_heading(doc, "8.2 긍정 vs 부정 종료 그룹 감정 궤적 및 트리거 분석", level=2)
    _add_para(doc,
        "상담 진행률(0~100%)을 기준으로 고객 발화의 감성 궤적을 10구간으로 세분화하면, "
        "두 그룹의 궤적 차이가 더 선명하게 드러납니다. "
        "긍정 종료 그룹은 상담 후반(80% 이후)에서 급격한 반전(+0.192)이 발생하는 반면, "
        "부정 종료 그룹은 30~40% 구간에서 감성 최저(-0.147)를 기록한 후 끝까지 회복하지 못합니다.")
    _add_para(doc,
        "하단의 트리거 발생 빈도를 보면, 감정 급락은 전 구간에 분산되어 있으나 "
        "종료 직전(90~100%)에 가장 집중됩니다. 반면 감정 급상승은 10~20% 이후 꾸준히 발생하며 "
        "종료 구간에서 가장 크게 집중되어, '마무리 인사/확인'이 긍정 전환의 핵심 트리거임을 확인합니다.")

    if "fig17" in figs:
        _add_figure(doc, figs["fig17"], width_inches=6.2,
                    caption="[그림 17] 종료 감성 그룹별 10구간 궤적 + 트리거 발생 빈도")

    _add_para(doc,
        "부정 종료 콜의 감정 급락 트리거를 분석하면, 비용 안내(30건), 지연/대기 안내(22건), "
        "불가/제한 안내(10건) 순으로 나타납니다. "
        "긍정 종료 콜의 급상승 트리거는 해결 안내(407건)가 압도적이며, "
        "일정 확정(85건), 공감 멘트(35건)가 뒤를 잇습니다. "
        "이 패턴은 9장의 '해결 안내가 감정 개선에 가장 효과적' 발견과 직접 연결됩니다.")

    doc.add_paragraph()
    _add_heading(doc, "8.3 연령대별 감정 궤적 비교", level=2)
    _add_para(doc,
        "연령대별 10구간 궤적을 중첩하면, 전 연령대가 유사한 하강 후 회복 패턴을 보이되 "
        "종료 구간에서 회복 격차가 뚜렷합니다. "
        "20~39세는 종료 감성(+0.132)과 회복폭(+0.261)이 가장 높고, "
        "65~74세는 종료 감성(+0.055), 회복폭(+0.179)으로 가장 낮습니다.")

    if "fig18" in figs:
        _add_figure(doc, figs["fig18"], width_inches=6.2,
                    caption="[그림 18] 연령대별 감정 궤적 비교 (10구간)")

    _add_para(doc,
        "전 연령대가 공통적으로 상담 초반 하강 → 종료 회복 패턴을 보이지만, "
        "종료 회복 강도에 뚜렷한 차이가 있습니다. "
        "20~39세(+0.132)는 가장 강한 회복을 보이고, "
        "연령이 높아질수록 회복이 약해져 65~74세(+0.055)에서 가장 낮습니다.")
    _add_para(doc,
        "그래프에서 주목할 패턴은 종료 직전(80→90%) 구간의 점프 크기입니다. "
        "20~39세는 이 구간에서 +0.144의 급격한 전환이 일어나는 반면, "
        "65~74세는 +0.117로 상대적으로 완만합니다. "
        "이 차이의 원인 중 하나로 상담사의 응대 패턴이 확인됩니다. "
        "고령층(65세+) 상담에서 상담사의 평균 발화 길이는 36~39자로, "
        "젊은층(50자) 대비 짧으며, 짧은 응답(10자 이하) 비율도 26~28%로 "
        "젊은층(21%)보다 높습니다. "
        "9장에서 확인된 것처럼 짧은 응답은 감정 개선에 가장 비효과적인 응대 유형이므로, "
        "고령층에 대한 더 짧은 응대가 회복 격차의 한 요인으로 작용할 수 있습니다.")
    _add_para(doc,
        "감정 하락 트리거도 연령대별로 다릅니다. "
        "젊은층(~49세)은 '지연/대기' 관련 안내가 급락 트리거의 상위를 차지하여 "
        "대기 시간에 민감한 반면, "
        "고령층(65세+)은 '일정/방문' 관련 안내가 상대적으로 높아 "
        "서비스 방문 절차의 복잡함이 부정 감성의 주요 요인입니다 "
        "(→ 6장 짧은 발화 분석).")

    doc.add_paragraph()
    _add_heading(doc, "8.4 제품군별 감정 궤적 비교", level=2)
    if "fig19" in figs:
        _add_figure(doc, figs["fig19"], width_inches=6.2,
                    caption="[그림 19] 제품군별 감정 궤적 비교 (10구간)")

    _add_para(doc,
        "전 제품군이 공통적으로 상담 초반~중반에 부정 구간을 거친 뒤 종료에서 회복하는 패턴을 보입니다. "
        "다만 종료 회복 강도에 차이가 있어, "
        "주방가전(+0.102)과 생활가전(+0.096)은 비교적 높은 회복을 보이는 반면, "
        "TV/AV(+0.075)는 가장 약한 회복으로 마무리됩니다.")
    _add_para(doc,
        "그래프에서 주목할 패턴은 TV/AV와 에어컨의 상담 후반부(70~80%) 재하락입니다. "
        "이 구간 직전 상담사 발화를 분석하면, "
        "TV/AV는 '기사 방문이 필요하다', '당일 수리가 어렵다', '부품 대기가 필요하다' 등 "
        "원격 해결 불가를 안내하는 시점에서 하락이 발생하며, "
        "에어컨은 '유상 수리 비용', '렌탈 요금 안내', '결제 확인' 등 "
        "비용 관련 안내가 트리거로 작용합니다. "
        "문제 파악 과정의 초반 하락과 달리, 후반부 하락은 해결 조건(비용·일정·부품) 안내에서 비롯되는 것으로, "
        "이 시점에서의 공감 멘트 병행이 감성 하락 완화에 효과적일 수 있습니다 (→ 9장).")
    _add_para(doc,
        "한편 주방가전은 짧은 발화 비율(47%)이 가장 높습니다. "
        "짧은 발화가 50% 이상인 콜에서는 불만/짜증 비율(16.5%)이 "
        "30% 미만 콜(9.7%)의 1.7배에 달해, "
        "고객이 수동적으로 듣기만 하는 구간이 길수록 부정 감성이 누적됨을 확인합니다 (→ 9장).")

    # ── 8.5 감정 그룹 색상 궤적 (개별 콜 스타일) ──
    doc.add_paragraph()
    _add_heading(doc, "8.5 감정 그룹 색상 궤적 (유형별 비교)", level=2)
    _add_para(doc,
        "앞선 궤적 분석이 감성 점수(Valence)의 평균 추이를 보여주었다면, "
        "아래 차트는 각 구간에서 가장 두드러진 감정 그룹을 색상으로 표시하여 "
        "상담 흐름 속 감정의 질적 전환을 시각화합니다. "
        "초록(감사/만족), 노랑(불안/걱정), 빨강(불만/짜증), 보라(혼란/당황)로 구분됩니다.")
    if "fig22" in figs:
        _add_figure(doc, figs["fig22"], width_inches=6.5,
                    caption="[그림 22] 감정 그룹 색상 궤적 — 종료 유형별 비교")
    _add_para(doc,
        "전체 및 긍정 종료 그룹에서는 초반~중반의 불안/걱정(노란색)이 80% 이후 감사/만족(초록)으로 "
        "명확하게 전환됩니다. 반면 부정 종료 그룹은 불만/짜증(빨간색)이 10~40% 구간에서 "
        "반복 출현하며, 종료까지 감사/만족으로 전환되지 못합니다.")
    if "fig23" in figs:
        _add_figure(doc, figs["fig23"], width_inches=6.5,
                    caption="[그림 23] 감정 그룹 색상 궤적 — 연령대별 비교")
    _add_para(doc,
        "연령대별로 보면, 20~49세는 초반부터 종료까지 불안/걱정→감사/만족의 단순 전환 패턴을 보이나, "
        "65~74세는 중간에 불만/짜증(빨간색)이 출현하여 감정 변동이 더 복잡합니다. "
        "50~64세는 40% 구간에서 일시적으로 감사/만족이 나타나 다른 연령대와 차별화됩니다.")

    _add_para(doc,
        "65~74세의 불만/짜증 출현은 바로 아래 연령대인 50~64세와 대조하면 원인이 분명해집니다. "
        "두 그룹은 제품 구성이 유사하고(정수기 20~21%, TV 13~15%), "
        "불만/짜증 비율 자체의 절대적 격차는 크지 않습니다. "
        "차이는 불안/걱정과의 역전 여부에 있습니다. "
        "10~20% 구간(탐색 초반)에서 50~64세는 불안/걱정(19.0%)이 불만/짜증(18.4%)을 근소하게 앞서지만, "
        "65~74세는 불만/짜증(20.8%)이 불안/걱정(17.8%)을 3.1%p 역전합니다. "
        "이 구간은 본인확인·계약자 확인 절차가 진행되는 탐색 단계이며, "
        "짧은 확인 응답('네 네', '예 예')에서 음성 톤이 불만으로 감지됩니다.")
    _add_para(doc,
        "더 큰 격차는 해결시도 구간에서 나타납니다. "
        "50~60%에서 65~74세 21.7% vs 50~64세 17.1%(+4.6%p), "
        "60~70%에서 17.1% vs 11.4%(+5.7%p)로, "
        "50~60% 구간에서도 불만/짜증이 불안/걱정을 역전하여 빨간색으로 표시됩니다. "
        "같은 제품(정수기)만 비교해도 해결시도 구간에서 65~74세 불만/짜증 비율은 27.1%로, "
        "50~64세(15.0%)의 약 2배입니다. "
        "이는 제품 구성의 차이가 아니라, "
        "동일 제품의 기술적 설명을 이해하는 과정에서 65~74세가 더 큰 좌절을 겪기 때문입니다.")
    _add_para(doc,
        "한편 65~74세 불만/짜증의 57.3%는 짧은 발화(음성 톤만으로 판별)이며, "
        "평소 대비 음높이(F0) 상승폭이 15.2Hz로 50~64세(10.3Hz)보다 큽니다. "
        "고령층은 전화 통화 시 또렷하게 전달하기 위해 목소리를 크게 높여 대답하는 경향이 있으며, "
        "이때 평소 대비 음높이 상승폭(15.2Hz)이 50~64세(10.3Hz)보다 커서 "
        "감정 분류 모델이 이를 불만/짜증으로 해석하는 경우가 일부 존재합니다.")

    if "fig24" in figs:
        _add_figure(doc, figs["fig24"], width_inches=6.5,
                    caption="[그림 24] 감정 그룹 색상 궤적 — 제품군별 비교")
    _add_para(doc,
        "에어컨/에어케어와 TV/AV에서만 불만/짜증·혼란/당황이 교차 출현하는 이유는 "
        "주방가전·생활가전과의 대조를 통해 드러납니다. "
        "주방가전도 구독/결제 키워드 포함 비율이 63.1%로 에어컨(61.8%)과 유사하고, "
        "수리/설치 키워드 비율도 86.1%로 에어컨(86.8%)과 거의 같습니다. "
        "그러나 주방가전(정수기·냉장고)은 '필터 교체', '기사 방문' 등 단일 액션으로 상담이 종결되는 반면, "
        "에어컨/에어케어와 TV/AV는 구독료 체계·카드 변경·설치 일정·부가서비스가 중첩되어 "
        "한 번의 안내로는 고객이 전체 상황을 파악하기 어렵습니다.")
    _add_para(doc,
        "에어컨/에어케어(76건)는 전 제품군 중 통화시간이 가장 길고(192초), "
        "NPS가 가장 낮습니다(7.59). 40~50% 구간에서 불만/짜증(24.4%, 전 제품군 최고)이 폭발한 뒤, "
        "60~70%에서 혼란/당황(16.4%), 70~80%에서 혼란/당황(21.4%, 전 제품군·전 구간 최고)으로 전환됩니다. "
        "실제 발화를 보면 \"이게 구독료가 2만 5천 원인가 2만 3천 원 정도가 결제되는 거잖아요\", "
        "\"날짜가 연기될 수도 있다라고 이야기를\" 등 "
        "결제 구조와 일정 절차를 확인하는 질문형 발화가 주를 이룹니다. "
        "65세 이상 고객 비율이 28.9%로 생활가전(22.1%)보다 높은 점도 복잡성을 가중시킵니다.")
    _add_para(doc,
        "TV/AV(87건, 96.6%가 TV)는 혼란/당황이 60~90% 구간에 집중됩니다 "
        "(60~70%: 18.8%, 80~90%: 18.9%). "
        "65세 이상 고객 비율이 32.1%로 전 제품군 최고이며, "
        "해결시도 단계에서 리모컨 구매, 벽걸이 설치, 가상계좌 결제, 구독 서비스 등 "
        "다양한 부가 절차가 동시에 안내되면서 혼란이 발생합니다. "
        "실제 발화: \"리모컨을 어디 가야 사요\", \"할 능력이 없거든요\", "
        "\"스탠드를 벽걸이로 할 순 없죠\" 등 디지털 기기 조작 자체에 대한 어려움이 드러납니다. "
        "다만 NPS 평균이 8.36으로 전 제품군 최고이고, "
        "80~90% 혼란/당황 발화 중 45%(9/20건)가 NPS 10점 콜에서 발생하여, "
        "상담 결과에는 만족하나 기술 안내를 소화하는 과정의 인지적 부하가 반영된 것으로 해석됩니다.")
    _add_para(doc,
        "반면 생활가전(세탁기·청소기 등)은 제품 구조가 단순하고 "
        "구독/결제 관련 콜 비율이 41.1%로 낮아, "
        "혼란/당황이 초반(10~30%, 17~18%)에만 나타나고 이후 빠르게 감소합니다. "
        "이는 초기 문제 파악 후 해결 방향이 명확해지기 때문입니다.")

    if "fig09" in figs:
        _add_figure(doc, figs["fig09"], width_inches=5.5,
                    caption="[그림 9] 종료 감정 그룹별 5단계 궤적 비교 (요약)")

    _add_insight_box(doc, [
        "긍정 종료 그룹: 80% 이후 급반전(+0.192) — 마무리 인사/확인이 결정적 트리거",
        "부정 종료 그룹: 30~40% 최저(-0.147) 후 끝까지 미회복 — 이 구간의 개입이 핵심",
        "연령대: 65~74세 종료 회복(+0.055)이 20~39세(+0.132)의 42% 수준 → 고령층 맞춤 마무리 필요",
        "TV/AV 종료 회복(+0.075) 최저 → 전화 상담만으로 해결 어려운 복잡 문제 비중 높음",
        "감정 급상승 트리거: 해결 안내(407건)가 압도적 → 구체적 정보 전달이 감성 반전의 핵심 (→ 9장)",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 9. 상담사 응대 패턴과 고객 감정 변화 (NEW)
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "9. 상담사 응대 패턴과 고객 감정 변화", level=1)

    _add_para(doc,
        "고객의 부정 발화 직후 상담사가 어떻게 응대했는지, 그리고 그 응대 이후 "
        "고객의 감정이 어떻게 변화했는지를 3-turn 시퀀스(부정 고객 → 상담사 응대 → 다음 고객)로 "
        "분석하였습니다. 전체 5,548건의 시퀀스에서 응대 유형, 응답 길이, 상담 단계, 부정 강도별 "
        "고객 감정 변화 패턴을 정량적으로 측정합니다.")

    _add_heading(doc, "9.1 응대 유형별 감정 변화", level=2)
    _add_para(doc,
        "상담사의 응대 내용을 키워드 기반으로 분류한 결과, "
        "일반 안내(+0.262)와 해결 안내(+0.251)가 가장 큰 감정 개선을 보이며 "
        "짧은 응답(+0.175)은 가장 작은 개선 효과를 보입니다. "
        "개선율(고객 감정이 나아진 비율) 역시 해결 안내(74%)가 짧은 응답(64%)보다 10%p 높습니다.")

    _add_table(doc,
        headers=["응대 유형", "N", "평균 변화", "개선율", "해석"],
        rows=[
            ["짧은 응답 (네/예)",    "1,176", "+0.176", "64%", "최소 효과 — 소극적 응대"],
            ["공감만",               "166",   "+0.247", "67%", "감정 인정은 하나 구체적 진전 없음"],
            ["공감 + 해결 안내",     "332",   "+0.224", "71%", "복합 응대"],
            ["해결 안내만",          "1,341", "+0.251", "73%", "구체적 정보 전달이 핵심"],
            ["일정/방문 안내",       "490",   "+0.236", "72%", "구체적 일정이 안심 요인"],
            ["일반 안내",            "2,533", "+0.262", "73%", "상황 설명·안내 포함"],
        ],
        col_widths=[3.5, 1.5, 2, 1.5, 6],
    )

    doc.add_paragraph()
    _add_heading(doc, "9.2 응답 길이별 감정 변화", level=2)
    _add_para(doc,
        "상담사 발화 길이별 효과를 분석하면 중간 길이(15~40자)가 "
        "가장 높은 개선율(75%)과 변화량(+0.281)을 기록합니다. "
        "짧은 응답(<15자)은 개선율 65%로 최저이며, "
        "지나치게 긴 응답(80자+)도 오히려 효과가 감소(71%)합니다.")

    _add_table(doc,
        headers=["응답 길이", "N", "평균 변화", "개선율"],
        rows=[
            ["짧은 (<15자)",   "1,438", "+0.183", "65%"],
            ["중간 (15~40자)", "1,585", "+0.281", "75% (최적)"],
            ["긴 (40~80자)",   "1,528", "+0.255", "74%"],
            ["상세 (80자+)",   "997",   "+0.224", "71%"],
        ],
        col_widths=[3.5, 2, 2.5, 6.5],
    )

    doc.add_paragraph()
    _add_heading(doc, "9.3 부정 강도별 응대 효과", level=2)
    _add_para(doc,
        "고객의 부정 강도가 심할수록 다음 발화에서의 자연 회복 경향이 강하게 작용합니다(회귀 효과). "
        "그러나 짧은 응답은 모든 강도에서 전체 평균 대비 낮은 개선을 보이며, "
        "경미한 부정(-0.05~-0.15)에서는 유일하게 감정을 악화(-0.012)시킵니다.")

    _add_table(doc,
        headers=["부정 강도", "N", "전체 평균", "짧은 응답", "차이"],
        rows=[
            ["경미 (-0.05~-0.15)", "860",   "+0.060", "-0.012", "-0.072 (악화)"],
            ["중간 (-0.15~-0.30)", "2,418", "+0.151", "+0.106", "-0.045"],
            ["심각 (-0.30~-0.50)", "1,471", "+0.328", "+0.304", "-0.024"],
            ["극심 (-0.50 이하)",  "799",   "+0.529", "+0.458", "-0.071"],
        ],
        col_widths=[3.5, 1.5, 2.5, 2.5, 4.5],
    )
    _add_para(doc,
        "경미한 부정 상태는 상담사의 실질적 응대가 없으면 자연 회복이 일어나지 않으며, "
        "오히려 '네'와 같은 단답형 응답이 '무시당함' 감정을 유발해 감성이 악화될 수 있습니다.")

    doc.add_paragraph()
    _add_heading(doc, "9.4 단계별 응대 효과 비교", level=2)
    _add_para(doc,
        "동일한 응대 유형이라도 상담 단계에 따라 효과가 달라집니다. "
        "탐색 단계에서 짧은 응답(+0.131)과 해결 안내(+0.244)의 격차가 0.113으로 가장 크며, "
        "결과제시 단계에서는 공감+해결 조합(+0.351)이 단독 해결 안내(+0.299)보다 우수합니다.")

    _add_table(doc,
        headers=["단계", "짧은 응답", "해결 안내", "공감+해결", "격차 (해결-짧은)"],
        rows=[
            ["탐색",    "+0.131 (59%)", "+0.244 (75%)", "+0.225 (75%)", "0.113"],
            ["해결시도", "+0.208 (69%)", "+0.242 (71%)", "+0.249 (69%)", "0.034"],
            ["결과제시", "+0.261 (71%)", "+0.299 (73%)", "+0.351 (74%)", "0.038"],
        ],
        col_widths=[2.5, 3, 3, 3, 3],
    )

    doc.add_paragraph()
    _add_heading(doc, "9.5 감정 회복률 및 해결 불가 후 회복", level=2)
    _add_para(doc,
        "부정 감성 구간이 존재하는 콜의 95%에서 감정 회복이 발생하며, "
        "해결 불가 안내 후에도 63%가 2턴 이내에 회복됩니다.")

    _add_table(doc,
        headers=["구분", "비율", "설명"],
        rows=[
            ["전체 감정 회복 발생률", "95%", "부정 이후 양의 감성으로 전환 성공"],
            ["해결 불가 후 빠른 회복", "63%", "2턴 이내 회복 — 공감+대안 제시 효과"],
            ["해결 불가 후 느린 회복", "27%", "3턴 이상 소요되지만 결국 해소"],
            ["해결 불가 후 미회복",    "10%", "부정 감성 지속 — 에스컬레이션 필요"],
        ],
        col_widths=[4.5, 2, 8],
    )

    if "fig12" in figs:
        _add_figure(doc, figs["fig12"], width_inches=6.2,
                    caption="[그림 12] 상담사 응대 패턴과 고객 감정 변화 (5,548 시퀀스)")

    _add_insight_box(doc, [
        "짧은 응답('네')은 경미한 부정에서 유일하게 감성 악화(-0.012) → 실질적 응대 필수",
        "중간 길이(15~40자)가 최적 — 핵심만 간결하게 전달하는 응대가 가장 효과적",
        "탐색 단계: 짧은 응답 vs 해결 안내 격차 0.113으로 최대 → 이 단계에서 '경청 중' 신호만으로는 불충분",
        "결과제시 단계: 공감+해결 조합(+0.351)이 단독 해결(+0.299)보다 0.052 높음 → 결과 안내 시 공감 병행 권고",
        "해결 불가 후 10% 미회복 콜 → 사후 콜백 또는 에스컬레이션 체계 도입 검토",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 10. 모달리티 비교
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "10. 모달리티 비교 분석", level=1)

    _add_para(doc,
        "텍스트(BERT), 음향 특징, SER 모델 단독 사용과 다양한 조합의 융합 방식을 비교하여 "
        "3-way 융합의 성능 우위를 검증하였습니다.")

    _add_table(doc,
        headers=["분석 방법", "NPS 상관계수 (r)", "비고"],
        rows=[
            ["텍스트만 (BERT)",       "0.116",  "기준선"],
            ["음향 특징만",           "-0.021", "단독으로는 NPS와 무관"],
            ["SER만",                 "0.171",  "음성 감정 직접 반영"],
            ["텍스트 + 음향 (2-way)", "0.229",  "기존 방법"],
            ["3-way 융합 (권장)",     "0.262",  "최고 성능"],
        ],
        col_widths=[5, 3.5, 6],
    )

    if "fig08" in figs:
        _add_figure(doc, figs["fig08"], width_inches=5.8,
                    caption="[그림 8] 모달리티별 NPS 분류 정확도 및 상관계수 비교")

    _add_insight_box(doc, [
        "3-way 융합(r=0.262)이 텍스트 단독(r=0.116) 대비 NPS 상관 2.3배 향상",
        "음향 단독(r=-0.021)은 NPS와 무관하나, 텍스트와 융합 시 상호 보완 효과 발휘",
        "SER 추가로 짧은 발화(44%)의 감성 탐지 가능 — 텍스트 분석 불가 구간 커버",
    ])

    if "fig07" in figs:
        _add_figure(doc, figs["fig07"], width_inches=5.2,
                    caption="[그림 7] NPS vs 평균 감성 점수 산점도 (r=0.262)")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 11. 결론 및 제언
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "11. 결론 및 제언", level=1)

    _add_heading(doc, "11.1 주요 분석 결과", level=2)
    findings = [
        ("감성 궤적 공통 패턴",
         "초기→탐색(하강)→해결시도(회복 시작)→종료(긍정) 패턴이 전체 콜의 공통 궤적으로 확인됩니다. "
         "이 패턴에서 이탈하는 콜(탐색에서 낮아졌다가 종료까지 회복 안됨)이 저NPS와 강하게 연결됩니다."),
        ("짧은 발화의 숨은 감성",
         "동일한 '네'도 음량(Energy)에 따라 감성 점수가 0.209포인트 차이납니다. "
         "텍스트만으로는 탐지 불가한 이 차이를 SER 모델이 포착합니다."),
        ("상담사 응대 유형의 정량적 효과",
         "5,548건 시퀀스 분석 결과, 짧은 응답(+0.175)은 해결 안내(+0.251) 대비 감정 개선 효과가 32% 낮으며 "
         "경미한 부정에서는 유일하게 감성을 악화(-0.012)시킵니다. "
         "중간 길이(15~40자) 응답이 개선율 75%로 최적입니다."),
        ("결과 vs 과정 만족도 분리",
         "67건 괴리 케이스(상담사 만족 高, NPS 低)는 결과 해결 여부와 상담 응대를 고객이 독립적으로 평가함을 보여줍니다."),
        ("종료 감성의 예측력",
         "종료 단계 감성만으로도 NPS를 예측할 수 있으며(r≈0.26), "
         "실시간 경보 시스템 도입 시 즉각적 사후 대응이 가능합니다."),
    ]
    for title, body in findings:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{title}: ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        rn = p.add_run(body)
        rn.font.size = Pt(10.5)

    doc.add_paragraph()
    _add_heading(doc, "11.2 운영 개선 제언", level=2)
    recommendations = [
        ("실시간 감성 모니터링",
         "종료 단계 예측 감성이 -0.05 이하인 콜을 자동 플래그 → QA 팀 우선 검토 대상으로 지정"),
        ("짧은 발화 코칭 지표 도입",
         "'네 네 네' 패턴(불안/걱정 31%) 탐지 시 상담사 실시간 알림 — 빠른 진행 상황 안내 유도"),
        ("탐색 단계 실질 응대 강화",
         "탐색 단계에서 짧은 응답(+0.131) vs 해결 안내(+0.244) 격차 0.113 — "
         "고객 문제 파악 중에도 '확인 중입니다' 등 실질적 진행 상황 안내 필수"),
        ("비용 안내 후 공감 멘트",
         "비용 안내 직후 '네'의 감성(+0.014)이 다른 맥락 대비 가장 낮으므로 공감 멘트 병행 권고"),
        ("고령층 전담 응대",
         "75세 이상 종료 감성 최저(+0.044) → 서비스 완료 확인 전화 또는 문자 안내 추가"),
        ("제품군별 맞춤 매뉴얼",
         "에어컨: 당일 처리 확정 멘트 우선; 주방가전: 탐색 단계 '공감 먼저' 응대 순서 조정"),
    ]
    for i, (title, body) in enumerate(recommendations, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"[{i}] {title}: ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        rn = p.add_run(body)
        rn.font.size = Pt(10.5)

    doc.add_paragraph()
    _add_heading(doc, "11.3 분석 한계 및 향후 과제", level=2)
    _add_bullet(doc, "음성 파일이 없는 일부 콜(약 8%)은 텍스트+음향 특징만으로 분석 (SER 미적용)")
    _add_bullet(doc, "현재 한국어 콜센터 음성에 대한 문장별 감정 라벨링 데이터가 부재하여 SER 도메인 파인튜닝이 어려운 상태 — 향후 라벨링 데이터 구축 시 성능 향상 가능")
    _add_bullet(doc, "상담사 음성 특징(톤 안정성, 말속도 등)은 현재 파이프라인에서 미추출 — 추후 확장 시 상담사 음성 분석 병행 가능")
    _add_bullet(doc, "감성 점수 ↔ 실제 행동(재이용 의향, 해지율) 간 인과 관계 검증 필요")
    _add_bullet(doc, "현재 NPS 상관 r=0.262는 중간 수준 — 음성 감성 외 추가 변수(상담 유형, 해결 여부) 통합 시 예측력 향상 가능")

    doc.add_paragraph()
    doc.add_paragraph()

    doc.save(REPORT_PATH)
    print(f"\n보고서 저장 완료: {REPORT_PATH}")
    return REPORT_PATH


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("음성 감성분석 보고서 생성 시작")
    print("=" * 60)

    print("\n[1/2] 그래프 생성 중...")
    figs = generate_figures()

    print("\n[2/2] Word 보고서 생성 중...")
    report_path = build_report(figs)

    print("\n" + "=" * 60)
    print(f"완료!")
    print(f"  그래프: {OUTPUT_DIR}")
    print(f"  보고서: {report_path}")
    print("=" * 60)
