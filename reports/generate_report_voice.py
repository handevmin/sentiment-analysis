"""
고객 상담 음성 감성 분석 보고서
- 피겨 생성(fig1~fig8, fig10) + Word 보고서 생성을 하나의 파일에서 처리
"""
import os, warnings, tempfile
warnings.filterwarnings('ignore')

# ── 하드코딩 통계 데이터 ──────────────────────────────────────────────────────
STATS = {
    "n_calls": 686,
    "n_turns": 24153,
    "n_customer": 12072,
    "short": 5322,
    "short_pct": 0.441,
    "changed": 2572,
    "change_rate": 0.381,
    "avg_duration": 184.1,
    "nps_corr": 0.262,
    "groups": {
        "안정/중립": 4199,
        "불안/걱정": 2429,
        "감사/만족": 2245,
        "불만/짜증": 1668,
        "혼란/당황": 1526,
    },
    "stages": {
        "초기":     {"mean": -0.043,  "std": 0.2841, "n": 1451},
        "탐색":     {"mean": -0.1207, "std": 0.2858, "n": 3354},
        "해결시도": {"mean": -0.0963, "std": 0.289,  "n": 3956},
        "결과제시": {"mean": -0.0512, "std": 0.3076, "n": 1969},
        "종료":     {"mean":  0.1322, "std": 0.3087, "n": 1342},
    },
    "transitions": 7877,
    "n2p": 1591,
    "p2n": 1462,
}

# ── 경로 설정 ─────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DOCX = os.path.join(os.path.dirname(BASE_DIR), '음성_감성분석_보고서.docx')

# ═════════════════════════════════════════════════════════════════════════════
# PART 1: 피겨 생성 (fig1~fig8, fig10)
# ═════════════════════════════════════════════════════════════════════════════

def generate_figures(fig_dir):
    """fig_dir 에 fig1.png ~ fig8.png, fig10.png 를 생성한다."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    import numpy as np
    from matplotlib.patches import Patch

    fp = 'C:/Windows/Fonts/malgun.ttf'
    if os.path.exists(fp):
        fm.fontManager.addfont(fp)
        plt.rcParams['font.family'] = fm.FontProperties(fname=fp).get_name()
    plt.rcParams['axes.unicode_minus'] = False
    plt.rcParams['figure.facecolor'] = 'white'
    plt.rcParams['axes.facecolor'] = 'white'
    plt.rcParams['axes.spines.top'] = False
    plt.rcParams['axes.spines.right'] = False
    plt.rcParams['axes.edgecolor'] = '#CCCCCC'
    plt.rcParams['xtick.color'] = '#555555'
    plt.rcParams['ytick.color'] = '#555555'

    # ── 통일 색상 팔레트 ──────────────────────────────────────────────────────
    NAVY    = '#1B2A4A'
    NAVY_L  = '#2E4A7A'
    NAVY_LL = '#4A6FA5'
    ORANGE   = '#D4731A'
    ORANGE_L = '#E8944A'
    EMO = {
        '감사/만족': '#2D8E5F',
        '안정/중립': '#5B7B94',
        '불안/걱정': '#C9862A',
        '불만/짜증': '#B8433A',
        '혼란/당황': '#7E5BA6',
    }
    STAGE_C = ['#4A6FA5', '#3B5998', '#2E4A7A', '#1B2A4A', '#2D8E5F']
    GRAY    = '#555555'
    GRAY_L  = '#999999'
    GRAY_LL = '#DDDDDD'

    def save(name):
        plt.tight_layout()
        plt.savefig(os.path.join(fig_dir, name), dpi=250,
                    facecolor='white', bbox_inches='tight')
        plt.close()
        print(f'  생성: {name}')

    # ═══ Fig 1: 고객 감정 궤적 ═══════════════════════════════════════════════
    fig, ax = plt.subplots(figsize=(10, 3.8))
    turns = [
        (6.1,  '감사/만족',  0.50), (8.5,   '안정/중립',  0.10),
        (16.3, '불안/걱정', -0.30), (31.3,  '불안/걱정', -0.20),
        (33.7, '안정/중립',  0.00), (45.1,  '안정/중립',  0.10),
        (86.0, '혼란/당황', -0.10), (94.7,  '안정/중립',  0.20),
        (104.1,'감사/만족',  0.40), (116.5, '불안/걱정', -0.35),
        (124.7,'안정/중립',  0.00), (129.5, '불안/걱정', -0.25),
        (135.6,'불안/걱정', -0.30), (155.0, '안정/중립',  0.05),
        (163.0,'감사/만족',  0.15),
    ]
    stages = [(0,17,'초기'),(17,43,'탐색'),(43,103,'해결시도'),(103,137,'결과제시'),(137,172,'종료')]
    for s, e, lb in stages:
        ax.axvspan(s, e, alpha=0.03, color=GRAY_L)
        ax.text((s+e)/2, 0.72, lb, ha='center', fontsize=8, color=GRAY_L, fontstyle='italic')

    ts = [t[0] for t in turns]
    fv = [t[2] for t in turns]
    cs = [EMO[t[1]] for t in turns]
    ax.plot(ts, fv, '-', color=GRAY_LL, lw=1.2)
    ax.scatter(ts, fv, c=cs, s=70, zorder=5, edgecolors='white', lw=1.5)
    ax.axhline(0, color=GRAY_LL, lw=0.8)

    ax.annotate('기사 연락 불가', xy=(16.3,-0.30), xytext=(28,-0.55), fontsize=7,
                color=EMO['불안/걱정'],
                arrowprops=dict(arrowstyle='->', color=EMO['불안/걱정'], lw=0.8))
    ax.annotate('감사 표현', xy=(104.1,0.40), xytext=(80,0.60), fontsize=7,
                color=EMO['감사/만족'],
                arrowprops=dict(arrowstyle='->', color=EMO['감사/만족'], lw=0.8))
    ax.annotate('지연 우려', xy=(116.5,-0.35), xytext=(132,-0.55), fontsize=7,
                color=EMO['불만/짜증'],
                arrowprops=dict(arrowstyle='->', color=EMO['불만/짜증'], lw=0.8))

    ax.set_xlabel('Time (sec)', fontsize=9, color=GRAY)
    ax.set_ylabel('Valence', fontsize=9, color=GRAY)
    ax.set_ylim(-0.7, 0.8)
    ax.set_xlim(0, 172)
    leg = [Patch(fc=c, ec='white', label=g) for g, c in EMO.items()]
    ax.legend(handles=leg, fontsize=7, ncol=5, loc='upper center',
              bbox_to_anchor=(0.5, 1.13), frameon=False)
    save('fig1.png')

    # ═══ Fig 2: 분석 방식 + 감정 분포 ═══════════════════════════════════════
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    ax1.pie([6750, 5322],
            labels=[f'텍스트+음성 융합\n(6,750건, 56%)', f'음성 전용\n(5,322건, 44%)'],
            colors=[NAVY_L, ORANGE],
            autopct='', startangle=90,
            textprops={'fontsize': 9, 'color': GRAY},
            wedgeprops={'edgecolor': 'white', 'linewidth': 2})
    ax1.set_title('(a) 분석 방식별 고객 발화', fontsize=9, fontweight='bold', color=GRAY)

    gn   = ['안정/중립', '불안/걱정', '감사/만족', '불만/짜증', '혼란/당황']
    gc   = [4199, 2429, 2245, 1668, 1526]
    gcol = [EMO[g] for g in gn]
    ax2.barh(gn[::-1], gc[::-1], color=gcol[::-1], edgecolor='white', height=0.55)
    for i, (g, c) in enumerate(zip(gn[::-1], gc[::-1])):
        ax2.text(c+50, i, f'{c:,}건 ({c/12072:.0%})', va='center', fontsize=8, color=GRAY)
    ax2.set_xlabel('발화 수', fontsize=9, color=GRAY)
    ax2.set_title('(b) 감정 그룹 분포 (N=12,072)', fontsize=9, fontweight='bold', color=GRAY)
    save('fig2.png')

    # ═══ Fig 3: 단계별 감성 (영역 차트) ════════════════════════════════════
    fig, ax = plt.subplots(figsize=(7, 3.8))
    sn = ['초기', '탐색', '해결시도', '결과제시', '종료']
    sv = [-0.043, -0.121, -0.096, -0.051, +0.132]
    x  = np.arange(5)

    ax.fill_between(x, 0, sv, where=[v >= 0 for v in sv],
                    alpha=0.15, color=EMO['감사/만족'], interpolate=True)
    ax.fill_between(x, 0, sv, where=[v < 0 for v in sv],
                    alpha=0.15, color=EMO['불안/걱정'], interpolate=True)
    ax.plot(x, sv, 'o-', color=NAVY, lw=2.5, ms=9,
            markeredgecolor='white', markeredgewidth=2, zorder=5)
    ax.axhline(0, color=GRAY_LL, lw=0.8)

    for i, v in enumerate(sv):
        ax.text(i, v+(0.015 if v >= 0 else -0.020),
                f'{v:+.3f}', ha='center', fontsize=9, fontweight='bold', color=GRAY)

    ax.set_xticks(x)
    ax.set_xticklabels(sn, fontsize=10)
    ax.set_ylabel('Mean Valence', fontsize=9, color=GRAY)
    ax.set_ylim(-0.18, 0.18)
    save('fig3.png')

    # ═══ Fig 4: 판정 변경 패턴 ═══════════════════════════════════════════════
    fig, ax = plt.subplots(figsize=(8, 4.5))
    patterns = [
        '혼란/당황 → 안정/중립', '감사/만족 → 안정/중립', '혼란/당황 → 불안/걱정',
        '불안/걱정 → 안정/중립', '안정/중립 → 불안/걱정', '감사/만족 → 불안/걱정',
        '불만/짜증 → 안정/중립', '감사/만족 → 불만/짜증',
    ]
    counts    = [481, 438, 316, 298, 197, 163, 145, 91]
    colors_p  = [NAVY_L, NAVY_L, ORANGE, NAVY_L, ORANGE, ORANGE, NAVY_L, EMO['불만/짜증']]
    ax.barh(patterns[::-1], counts[::-1], color=colors_p[::-1],
            edgecolor='white', height=0.6)
    for i, (p, c) in enumerate(zip(patterns[::-1], counts[::-1])):
        ax.text(c+10, i, f'{c}건', va='center', fontsize=8, fontweight='bold', color=GRAY)
    ax.set_xlabel('변경 건수', fontsize=9, color=GRAY)
    ax.set_title('AI 교차검증으로 판정 변경된 패턴 (총 2,572건, 38.1%)',
                 fontsize=9, fontweight='bold', color=GRAY, loc='left')
    save('fig4.png')

    # ═══ Fig 5: NPS 그룹별 감성 (도넛 + 수치) ════════════════════════════════
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4),
                                   gridspec_kw={'width_ratios': [1, 1.3]})

    nps_labels = ['비추천\n(1-5점)', '중립\n(6-8점)', '추천\n(9-10점)']
    nps_counts = [104, 325, 257]
    nps_colors = [EMO['불만/짜증'], ORANGE_L, EMO['감사/만족']]
    wedges, texts, autotexts = ax1.pie(
        nps_counts, labels=nps_labels, colors=nps_colors,
        autopct='%1.0f%%', startangle=90, pctdistance=0.75,
        textprops={'fontsize': 9, 'color': GRAY},
        wedgeprops={'edgecolor': 'white', 'linewidth': 2, 'width': 0.45})
    for at in autotexts:
        at.set_fontsize(9); at.set_fontweight('bold')
    ax1.set_title('(a) NPS 분포 (N=686)', fontsize=9, fontweight='bold', color=GRAY)

    groups_l = ['추천 (9-10)', '중립 (6-8)', '비추천 (1-5)']
    vals     = [-0.0466, -0.0670, -0.0972]
    y        = np.arange(3)
    ax2.hlines(y, 0, vals,
               color=[EMO['감사/만족'], ORANGE_L, EMO['불만/짜증']], lw=3)
    ax2.scatter(vals, y,
                c=[EMO['감사/만족'], ORANGE_L, EMO['불만/짜증']],
                s=120, zorder=5, edgecolors='white', lw=2)
    for i, v in enumerate(vals):
        ax2.text(v-0.005, i, f'{v:+.04f}',
                 va='center', ha='right', fontsize=9, fontweight='bold', color=GRAY)
    ax2.axvline(0, color=GRAY_LL, lw=0.8)
    ax2.set_yticks(y)
    ax2.set_yticklabels(groups_l, fontsize=9)
    ax2.set_xlabel('Mean Valence', fontsize=9, color=GRAY)
    ax2.set_title(f'(b) NPS별 감성 (r = 0.262)', fontsize=9, fontweight='bold', color=GRAY)
    ax2.set_xlim(-0.12, 0.01)
    save('fig5.png')

    # ═══ Fig 6: 연령대 + 성별 (버블 차트) ═══════════════════════════════════
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.5))

    ages_l = ['20~39세', '40~49세', '50~64세', '65~74세', '75세+']
    nps_a  = [8.21, 7.76, 7.95, 8.11, 8.31]
    val_a  = [-0.052, -0.061, -0.058, -0.080, -0.056]
    cnt_a  = [101, 135, 264, 141, 29]
    size_a = [c*2.5 for c in cnt_a]

    ax1.scatter(nps_a, val_a, s=size_a, c=NAVY_L, alpha=0.6,
                edgecolors=NAVY, lw=1.5, zorder=5)
    for i in range(5):
        ax1.text(nps_a[i], val_a[i]+0.004,
                 f'{ages_l[i]}\n(N={cnt_a[i]})',
                 ha='center', va='bottom', fontsize=7.5, color=GRAY)
    ax1.axhline(0, color=GRAY_LL, lw=0.8)
    ax1.axvline(8.0, color=GRAY_LL, lw=0.8, ls='--')
    ax1.set_xlabel('NPS', fontsize=9, color=GRAY)
    ax1.set_ylabel('Valence', fontsize=9, color=GRAY)
    ax1.set_xlim(7.5, 8.6); ax1.set_ylim(-0.10, 0.01)
    ax1.set_title('(a) 연령대별 NPS × Valence × 건수',
                  fontsize=9, fontweight='bold', color=GRAY)

    genders_l = ['남성', '여성']
    nps_g     = [8.19, 7.87]
    val_g     = [-0.071, -0.057]
    cnt_g     = [263, 407]
    colors_g  = [NAVY_L, ORANGE]
    for i in range(2):
        ax2.scatter(nps_g[i], val_g[i], s=cnt_g[i]*2, c=colors_g[i],
                    alpha=0.6, edgecolors=colors_g[i], lw=1.5, zorder=5)
        ax2.text(nps_g[i]+0.05, val_g[i],
                 f'{genders_l[i]}\n(N={cnt_g[i]})',
                 va='center', fontsize=9, color=GRAY)
    ax2.axhline(0, color=GRAY_LL, lw=0.8)
    ax2.axvline(8.0, color=GRAY_LL, lw=0.8, ls='--')
    ax2.set_xlabel('NPS', fontsize=9, color=GRAY)
    ax2.set_ylabel('Valence', fontsize=9, color=GRAY)
    ax2.set_xlim(7.5, 8.5); ax2.set_ylim(-0.10, 0.01)
    ax2.set_title('(b) 성별 NPS × Valence × 건수',
                  fontsize=9, fontweight='bold', color=GRAY)
    save('fig6.png')

    # ═══ Fig 7: 상담사 응대 패턴 ════════════════════════════════════════════
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    pats   = ['공감 멘트 (28%)', '선제적 안내 (20%)', '구체적 솔루션 (12%)']
    vals_p = [28, 20, 12]
    ax1.barh(pats[::-1], vals_p[::-1],
             color=[NAVY, NAVY_L, NAVY_LL], edgecolor='white', height=0.5)
    for i, v in enumerate(vals_p[::-1]):
        ax1.text(v+0.5, i, f'{v}%', va='center', fontsize=10,
                 fontweight='bold', color=GRAY)
    ax1.set_xlabel('%', fontsize=9, color=GRAY)
    ax1.set_title('(a) 상담사 응대 유형', fontsize=9, fontweight='bold', color=GRAY)

    nps_lb = ['비추천', '중립', '추천']
    emp = [26, 28, 28]; pro = [20, 19, 20]
    x3 = np.arange(3); w = 0.3
    ax2.bar(x3-w/2, emp, w, label='공감 멘트', color=NAVY_L, edgecolor='white')
    ax2.bar(x3+w/2, pro, w, label='선제적 안내', color=ORANGE, edgecolor='white')
    ax2.set_xticks(x3); ax2.set_xticklabels(nps_lb, fontsize=9)
    ax2.set_ylabel('%', fontsize=9, color=GRAY)
    ax2.legend(fontsize=8)
    ax2.set_title('(b) NPS별 응대 비율 (차이 없음)',
                  fontsize=9, fontweight='bold', color=GRAY)
    save('fig7.png')

    # ═══ Fig 8: Baseline 보정 (덤벨 차트) ══════════════════════════════════
    fig, ax = plt.subplots(figsize=(8, 3.8))
    labels = ['감사 표현', '인사 응대', '짧은 응답 (대기 수락)', '마무리 인사']
    before = [0.076, -0.134, 0.180, -0.241]
    after  = [0.298, -0.037, -0.001, -0.425]
    y      = np.arange(4)

    for i in range(4):
        ax.plot([before[i], after[i]], [y[i], y[i]],
                '-', color=GRAY_LL, lw=2, zorder=1)
    ax.scatter(before, y, c=GRAY_L, s=100, zorder=5,
               edgecolors='white', lw=2, label='보정 전 (절대값)')
    ax.scatter(after, y, c=NAVY, s=100, zorder=5,
               edgecolors='white', lw=2, label='보정 후 (Baseline)')
    for i in range(4):
        ax.text(before[i], y[i]+0.2, f'{before[i]:+.3f}',
                ha='center', fontsize=8, color=GRAY_L)
        ax.text(after[i], y[i]-0.25, f'{after[i]:+.3f}',
                ha='center', fontsize=8, fontweight='bold', color=NAVY)
    ax.axvline(0, color=GRAY_LL, lw=0.8)
    ax.set_yticks(y); ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlabel('Audio Valence', fontsize=9, color=GRAY)
    ax.set_xlim(-0.55, 0.4)
    ax.legend(fontsize=8, loc='lower left')
    save('fig8.png')

    # ═══ Fig 10: 제품군별 NPS ════════════════════════════════════════════════
    fig, ax = plt.subplots(figsize=(7, 3.5))
    prods = ['에어컨\n(N=76)', '주방가전\n(N=331)', '생활가전\n(N=163)',
             'PC/모니터\n(N=17)', 'TV/AV\n(N=87)']
    nps_p = [7.59, 7.81, 7.96, 8.35, 8.36]
    c_p   = [ORANGE, ORANGE_L, GRAY_L, NAVY_LL, NAVY_L]
    bars  = ax.bar(prods, nps_p, color=c_p, edgecolor='white', width=0.55)
    for b, v in zip(bars, nps_p):
        ax.text(b.get_x()+b.get_width()/2, v+0.04,
                f'{v:.2f}', ha='center', fontsize=9, fontweight='bold', color=GRAY)
    ax.set_ylabel('평균 NPS', fontsize=9, color=GRAY)
    ax.set_ylim(7, 9)
    ax.axhline(8.0, color=GRAY_LL, lw=0.8, ls='--')
    ax.tick_params(axis='x', labelsize=8)
    save('fig10.png')

    print(f'  피겨 생성 완료 → {fig_dir}')


# ═════════════════════════════════════════════════════════════════════════════
# PART 2: Word 보고서 생성
# ═════════════════════════════════════════════════════════════════════════════

def generate_report(fig_dir, output_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    NAVY_RGB = RGBColor(33,  33,  33)
    GRAY_RGB = RGBColor(100, 100, 100)

    # ── 헬퍼: 셀 음영 ────────────────────────────────────────────────────────
    def shade(cell, h):
        tc = cell._element.get_or_add_tcPr()
        tc.append(tc.makeelement(qn('w:shd'),
                                 {qn('w:fill'): h, qn('w:val'): 'clear'}))

    # ── 헬퍼: 표 ─────────────────────────────────────────────────────────────
    def tbl(doc, headers, rows):
        n_cols = len(headers)
        t = doc.add_table(rows=1 + len(rows), cols=n_cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = 'Table Grid'

        pw = 15.5  # A4 사용 가능 너비 cm
        col_max = []
        for ci in range(n_cols):
            lens = [len(str(headers[ci]))]
            for rd in rows:
                if ci < len(rd):
                    lens.append(max(len(line) for line in str(rd[ci]).split('\n')))
            col_max.append(max(lens) ** 0.55)

        total  = sum(col_max) or 1
        widths = [max(col_max[ci] / total * pw, 2.0) for ci in range(n_cols)]
        w_sum  = sum(widths)
        widths = [w / w_sum * pw for w in widths]

        for ci in range(n_cols):
            for row in t.rows:
                row.cells[ci].width = Cm(widths[ci])

        center = set(range(n_cols - 1)) if n_cols > 1 else set()

        def _cell(cell, text, header=False, ci=0):
            for p in cell.paragraphs:
                p.clear()
            p   = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(51, 51, 51)
            if header:
                run.bold = True
            pf = p.paragraph_format
            pf.space_before  = Pt(1)
            pf.space_after   = Pt(1)
            pf.line_spacing  = 1.0
            if header or ci in center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc_pr = cell._element.get_or_add_tcPr()
            mar   = tc_pr.makeelement(qn('w:tcMar'), {})
            for side in ['top', 'bottom']:
                mar.append(mar.makeelement(
                    qn(f'w:{side}'), {qn('w:w'): '55', qn('w:type'): 'dxa'}))
            for side in ['left', 'right']:
                mar.append(mar.makeelement(
                    qn(f'w:{side}'), {qn('w:w'): '120', qn('w:type'): 'dxa'}))
            tc_pr.append(mar)

        for i, h in enumerate(headers):
            _cell(t.rows[0].cells[i], h, header=True, ci=i)
            shade(t.rows[0].cells[i], 'F2F2F2')

        for ri, rd in enumerate(rows):
            for ci, v in enumerate(rd):
                _cell(t.rows[ri+1].cells[ci], v, ci=ci)

        spacer = doc.add_paragraph('')
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after  = Pt(2)
        for r in spacer.runs:
            r.font.size = Pt(2)
        return t

    # ── 헬퍼: 그림 삽입 ──────────────────────────────────────────────────────
    def fig(doc, fname, caption, w=5.8):
        path = os.path.join(fig_dir, fname)
        if not os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'[그림 파일 없음: {fname}]')
            r.font.color.rgb = RGBColor(200, 0, 0)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(w))
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY_RGB
        r.italic = True
        c.paragraph_format.space_after = Pt(14)

    # ── 헬퍼: 굵은 제목 + 본문 단락 ─────────────────────────────────────────
    def bp(doc, title, body):
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        p.add_run(body)

    def section_heading(doc, text, level=1):
        doc.add_heading(text, level=level)

    # ── 편의 변수 ─────────────────────────────────────────────────────────────
    total_cust  = STATS['n_customer']
    n_calls     = STATS['n_calls']
    short       = STATS['short']
    short_pct   = STATS['short_pct']
    changed     = STATS['changed']
    meaningful  = total_cust - short
    changed_pct = changed / meaningful
    corr        = STATS['nps_corr']

    # ── 문서 생성 ─────────────────────────────────────────────────────────────
    doc = Document()

    s = doc.styles['Normal']
    s.font.name = '맑은 고딕'
    s.font.size = Pt(10.5)
    s.paragraph_format.line_spacing = 1.5
    s._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']
        hs.font.name = '맑은 고딕'
        hs.font.color.rgb = NAVY_RGB
        hs._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ══ 표지 ════════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('고객 상담 음성 감성 분석 보고서')
    r.font.size = Pt(28); r.bold = True; r.font.color.rgb = NAVY_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('STT 텍스트 분석 vs STT + 음성 프로소디 융합 분석\n인사이트 리포트')
    r.font.size = Pt(14); r.font.color.rgb = GRAY_RGB

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('2026.04')
    r.font.size = Pt(12); r.font.color.rgb = GRAY_RGB
    doc.add_page_break()

    # ══ Executive Summary ════════════════════════════════════════════════════
    section_heading(doc, 'Executive Summary', level=1)

    doc.add_paragraph(
        f'{n_calls}건의 실제 상담 음성 전체를 대상으로 STT 텍스트 분석과 음성 프로소디 분석의 '
        f'결과 차이를 비교하였다. 전체 고객 발화 {total_cust:,}건 중 {short:,}건'
        f'({short_pct:.0%})은 짧은 응답으로 텍스트 분석만으로는 감정 파악이 불가능하였으며, '
        f'의미 있는 문장 {meaningful:,}건 중 {changed:,}건({changed_pct:.0%})에서 음성 추가 시 '
        f'감정 판정이 변경되었다. 텍스트와 음성을 종합하여 산출한 감성 점수(상세 정의는 2.2절 참조)는 '
        f'실제 NPS 점수와도 유의미한 상관(r={corr:.3f})을 보여, '
        f'음성 분석이 기존 텍스트 분석을 보완하는 독립적 가치가 있음을 확인하였다.'
    )
    doc.add_paragraph()

    tbl(doc,
        ['핵심 지표', '수치', '의미'],
        [
            ['분석 대상',      f'{n_calls}건 (WAV 매칭)',             '전체 698건 중 98.3%'],
            ['고객 발화',      f'{total_cust:,}건',                   '상담사 발화 제외'],
            ['짧은 응답 비율', f'{short:,}건 ({short_pct:.0%})',       '텍스트 분석 불가 → 음성으로 커버'],
            ['판정 변경률',    f'{changed:,}/{meaningful:,}건 ({changed_pct:.0%})',
             '음성 추가로 감정 그룹 변경'],
            ['NPS-감성 상관',  f'r = {corr:.3f}',
             '유의미한 양의 상관 — 분석 타당성 검증 (N=686)'],
            ['최다 감정 그룹', '안정/중립 34.8%',                    '5그룹 중 가장 큰 비중'],
            ['종료 단계 감성', '긍정(+) 방향으로 전환',              '유일한 긍정 구간 (탐색 단계 최저 부정)'],
        ]
    )
    doc.add_page_break()

    # ══ 1. 분석 배경 및 목적 ════════════════════════════════════════════════
    section_heading(doc, '1. 분석 배경 및 목적', level=1)

    doc.add_paragraph(
        '콜센터 상담 품질 측정의 전통적 방법은 NPS(순추천지수) 설문과 모니터링 요원의 '
        '샘플 청취에 의존한다. 그러나 NPS 응답률은 통상 30~40% 수준에 불과하고, '
        '샘플 모니터링은 전체 통화의 극히 일부만을 검토한다. 이에 따라 대다수 통화의 '
        '품질 데이터는 누락된 채로 남아, 상담사 코칭과 서비스 개선의 근거가 부족한 상황이다.'
    )
    doc.add_paragraph(
        '본 분석은 두 가지 기술적 접근을 결합하여 이 공백을 메운다. 첫째, STT(Speech-to-Text) '
        '기반 텍스트 감성 분석으로 발화 내용의 의미를 파악한다. 둘째, 음성 프로소디 분석으로 '
        '텍스트로는 전달되지 않는 어조·강도·억양의 변화를 포착한다. 두 결과가 충돌할 경우 '
        'LLM이 맥락을 고려한 최종 판정을 내림으로써, 단일 모달리티 대비 '
        '더 정확하고 설명 가능한 감성 데이터를 생산한다.'
    )
    doc.add_paragraph(
        '분석의 핵심 목적은 세 가지이다: '
        '(1) STT 텍스트 분석과 음성 프로소디 융합 분석의 판정 차이 정량화, '
        '(2) NPS와 음성 감성의 상관관계 및 보완적 가치 검증, '
        '(3) 상담사 코칭·모니터링·실시간 대시보드 구축을 위한 실무 인사이트 도출.'
    )
    doc.add_page_break()

    # ══ 2. 분석 방법 ════════════════════════════════════════════════════════
    section_heading(doc, '2. 분석 방법', level=1)

    doc.add_paragraph(
        '분석 파이프라인은 4단계로 구성된다. 텍스트와 음성을 각각 독립적으로 분석한 후, '
        'LLM이 두 결과를 교차검증하여 최종 판정을 내리는 구조이다. '
        'LLM은 STT 원문을 직접 읽지 않으며, BERT의 감정 확률과 음성 감성 수치만을 보고 판단하여 '
        '판정 편향을 최소화한다.'
    )

    tbl(doc,
        ['단계', '방법', '핵심 기능'],
        [
            ['Stage 1\nForced Alignment',
             'stable-ts\n(Whisper base)',
             'STT 텍스트를 오디오에 단어 단위 정렬\n→ 발화 시작/종료 시각 확보 (±0.1초)'],
            ['Stage 2\n텍스트 감정',
             'BERT 60-class\n→ 5그룹 재매핑',
             '고객 발화의 텍스트 감정 분류\n→ 라벨 수 편향 보정 (48/60 부정 → 정규화)'],
            ['Stage 3\n음성 특징',
             'librosa\n10개 지표',
             'F0, Energy, ZCR, Jitter, Shimmer, HNR 추출\n→ 고객 통화 내 평균 대비 상대적 변화 (Baseline 보정)'],
            ['Stage 4\nLLM 교차검증',
             'LLM',
             'BERT vs 음성 충돌 시 어느 쪽을 신뢰할지 판정\n→ 최종 감정 그룹 + 감성 점수 + 자연어 근거'],
        ]
    )

    doc.add_paragraph()
    section_heading(doc, '2.1 음성 감성 지수의 정의', level=2)
    doc.add_paragraph(
        '본 분석에서 산출하는 "음성 감성 지수"는 고객 발화의 감정 방향(긍정/부정)을 '
        '수치화한 자체 지표이다. 산출 과정은 다음과 같다.'
    )
    tbl(doc,
        ['구분', '내용'],
        [
            ['산출 입력',
             '① BERT 텍스트 감정 분류 결과 (5개 그룹 확률)\n'
             '② 음성 프로소디 특징 10개 (F0, Energy, ZCR, Jitter, Shimmer, HNR 등)\n'
             '③ LLM 교차검증 (①과 ②가 충돌할 때 어느 쪽을 신뢰할지 판정)'],
            ['스케일',
             '-1(매우 부정) ~ 0(중립) ~ +1(매우 긍정)\n'
             '양수(+)는 긍정 방향, 음수(-)는 부정 방향을 의미'],
            ['기준점',
             '해당 고객의 통화 내 평균을 0으로 설정 (Baseline 보정)\n'
             '→ "다른 사람과 비교"가 아니라 "본인의 평소 톤 대비 변화"를 측정'],
            ['해석 방법',
             '절대 점수보다 "패턴"과 "상대 비교"에 의미가 있음\n'
             '예: 탐색 단계에서 부정(-) → 종료 단계에서 긍정(+)으로 회복'],
        ]
    )

    doc.add_paragraph()
    section_heading(doc, '2.3 감성 지수의 신뢰성 검증', level=2)
    doc.add_paragraph(
        '감성 지수가 실제 고객 만족도(NPS)와 일치하는지 확인하여 분석의 타당성을 검증하였다.'
    )
    tbl(doc,
        ['검증 항목', '결과', '해석'],
        [
            ['NPS와의 상관관계',
             f'Pearson r = {STATS["nps_corr"]:.3f} (N=686)',
             '통계적으로 유의미한 양의 상관.\n감성 지수가 높을수록 NPS도 높은 경향.'],
            ['NPS 그룹별 감성 지수 차이',
             '비추천(1~5점): 부정(-) 방향\n추천(9~10점): 상대적 긍정',
             '비추천 고객은 추천 고객 대비\n감성 점수가 약 2배 낮게 나타남.'],
            ['단계별 패턴 일치',
             '탐색 단계 최저 → 종료 단계 최고',
             '문제 호소(탐색)에서 부정,\n해결 후(종료) 긍정 — 상식적 흐름과 일치.'],
        ]
    )
    doc.add_paragraph(
        '이상의 검증 결과를 바탕으로, 본 보고서에서 제시하는 감성 지수의 절대 수치보다는 '
        '그룹 간 상대 비교, 단계별 추이 패턴, 감정 전환 지점에서의 방향 변화에 '
        '초점을 맞추어 해석한다.'
    )

    doc.add_paragraph()
    section_heading(doc, '2.3 Baseline 보정', level=2)
    doc.add_paragraph(
        '음성 감성 지수 산출의 핵심은 Baseline 보정이다. 모든 고객에 동일한 절대 기준을 적용하면 '
        '화자의 고유한 발화 습관이 감정 상태와 무관하게 부정으로 판정되는 오류가 발생한다. '
        'Baseline 보정은 해당 고객의 통화 내 전체 평균을 기준점으로 설정하고, '
        '각 발화가 평소 대비 얼마나 달라졌는지만 측정한다.'
    )

    fig(doc, 'fig8.png',
        'Fig. 1  Baseline 보정 전후 비교 — 보정 전에는 이름 소개 발화도 부정(-)으로 오판되지만, '
        '보정 후에는 "본인의 평소 톤 대비 변화 없음"으로 교정되어 중립(0) 수준으로 안정화됨.',
        5.5)

    doc.add_page_break()

    # ══ 3. 핵심 인사이트 ════════════════════════════════════════════════════
    section_heading(doc, '3. 핵심 인사이트', level=1)

    # ── 3.1 STT Only vs STT+Audio 비교 ────────────────────────────────────
    section_heading(doc, '3.1 STT Only vs STT+Audio 비교', level=2)

    fig(doc, 'fig4.png',
        'Fig. 2  STT Only vs STT+Audio 판정 변경 패턴 수평 바 차트 — '
        f'의미 있는 발화 {meaningful:,}건 중 {changed:,}건({changed_pct:.0%})에서 판정 변경 발생.',
        6.0)

    doc.add_paragraph(
        f'44%의 고객 발화는 텍스트만으로 감정 파악이 불가능하다. '
        f'"네", "예", "알겠습니다" 등 짧은 응답 {short:,}건({short_pct:.0%})이 이에 해당하며, '
        f'음성 분석을 추가해야 비로소 이 영역의 감정이 포착된다. '
        f'나아가 의미 있는 문장 {meaningful:,}건 중에서도 {changed:,}건({changed_pct:.0%})은 '
        f'음성 추가 시 판정이 바뀐다. 변경 패턴은 크게 3가지로 분류된다.'
    )

    tbl(doc,
        ['변경 유형', '추정 빈도', '대표 사례', '변경 전 → 후', '음성 근거'],
        [
            ['BERT 오판 보정',  '~40%',
             '이름 소개, 단순 확인',
             '감사/만족 → 안정/중립',
             '음성 baseline 대비 변화 없음\n→ 감정 없는 발화로 보정'],
            ['숨은 감정 감지',  '~35%',
             '형식적 인사, 마무리',
             '(분석불가) → 불안/걱정',
             '에너지 급감(부정(-) 방향)\nF0 불안정 → 위축 감지'],
            ['감정 강도 조정',  '~25%',
             '우려 표현, 요청',
             '불안/걱정(약) → 불안/걱정(강)',
             '텍스트·음성 모두 부정(-) 방향\n→ 융합 시 부정 강도 심화'],
        ]
    )

    doc.add_paragraph()

    fig(doc, 'fig1.png',
        'Fig. 3  대표 사례 — 고객 감정 변화 궤적 (배송 문의, NPS 10점). '
        '점 색상은 감정 그룹, 배경 음영은 상담 단계. '
        'NPS 최고점 상담에서도 통화 중 부정(-) 방향 감성 구간이 반복 출현.',
        6.0)

    doc.add_paragraph(
        '위 차트는 단일 통화에서 고객의 감정이 발화 단위로 어떻게 변화하는지를 시각화한 것이다. '
        'NPS 10점(최고)인 상담임에도 불구하고 통화 중간에 불안/걱정 구간이 반복적으로 나타난다. '
        'NPS는 통화 종료 후 "전체적인 인상"을 회고적으로 평가하므로, 통화 과정의 감정 변동을 '
        '반영하지 못한다. 음성 감성 분석은 이러한 "순간순간의 감정"을 포착하여 NPS를 보완한다.'
    )

    tbl(doc,
        ['시간 구간', '발화 요약', '감정 그룹', '감성 방향', '해석'],
        [
            ['0~17초 (초기)',    '인사, 제품 문의',             '감사/만족 → 안정/중립', '긍정(+) → 중립(0)', '긍정적 시작, 일반 인사 교환'],
            ['16~45초 (탐색)',  '기사 연락 불가 호소 → 대기', '불안/걱정 → 안정/중립', '부정(-) → 중립(0)', '문제 호소 시 부정 하락. 대기 수락 후 안정화'],
            ['86~104초 (해결)', '배송 일정 확인 → 감사',       '혼란/당황 → 감사/만족', '부정(-) → 긍정(+)', '구체적 일정 안내 후 긍정 전환'],
            ['116~136초 (종료)', '지연 우려 → 수용',           '불안/걱정 → 불안/걱정', '부정(-) 지속',      '지연 가능성 언급 후 재하락. 형식적 마무리'],
        ]
    )

    doc.add_paragraph()

    fig(doc, 'fig2.png',
        'Fig. 4  (a) 분석 방법 분포 — 짧은 응답 44.1%는 음성만으로 판정. '
        f'(b) 감정 그룹별 분포 (N={total_cust:,}) — 안정/중립 34.8%, 불안/걱정 20.1% 순.',
        5.8)

    doc.add_paragraph(
        f'전체 고객 발화의 {short_pct:.0%}가 짧은 응답("네", "예", "알겠습니다" 등)으로, '
        f'텍스트 분석만으로는 감정 판별이 원천적으로 불가능하다. '
        f'음성 분석을 추가하면 이 영역이 커버되어, 분석 가능 발화가 '
        f'{1-short_pct:.0%} → 100%로 확대된다. '
        f'감정 그룹 분포는 안정/중립(34.8%), 불안/걱정(20.1%), 감사/만족(18.6%), '
        f'불만/짜증(13.8%), 혼란/당황(12.6%)의 순으로 나타났다.'
    )

    doc.add_page_break()

    # ── 3.2 단계별 감성 추이 ─────────────────────────────────────────────
    section_heading(doc, '3.2 단계별 감성 추이', level=2)

    fig(doc, 'fig3.png',
        'Fig. 5  상담 단계별 평균 감성 점수 바 차트 — 탐색 단계에서 고객 감성이 가장 부정적이며, '
        '종료 단계에서 긍정(+) 방향으로 회복되는 패턴이 관찰됨.',
        5.5)

    doc.add_paragraph(
        '상담을 5단계(초기 → 탐색 → 해결시도 → 결과제시 → 종료)로 구분하여 '
        '평균 감성 점수 변화를 분석하면, 탐색 단계에서 고객 감성이 가장 부정적이며 '
        '종료 단계에서만 긍정(+) 방향으로 회복되는 패턴이 관찰된다. '
        '탐색→해결시도 단계에서 감정 전환이 가장 빈번하게 발생한다. '
        '이 패턴은 실제 NPS 점수와도 유의미한 상관(r=0.262)을 보여, 분석의 타당성을 뒷받침한다.'
    )

    tbl(doc,
        ['상담 단계', '감성 방향', '상대적 수준', '해석'],
        [
            ['초기',     '부정(-) 방향', '소폭 부정',   '인사·확인 단계. 감정적으로 중립에 가까우나 소폭 부정'],
            ['탐색',     '부정(-) 방향', '최저점',      '고객이 문제 상황을 서술하는 단계. 전 구간 중 감성 최저'],
            ['해결시도', '부정(-) 방향', '완화 시작',   '상담사 대응 중. 탐색 단계 대비 부정이 완화되기 시작'],
            ['결과제시', '부정(-) 방향', '추가 완화',   '해결책 제시 단계. 추가 완화, 그러나 여전히 부정(−)'],
            ['종료',     '긍정(+) 방향', '유일한 긍정', '마무리 인사. 전 구간 중 유일한 긍정(+) 구간'],
        ]
    )

    bp(doc, '핵심 인사이트: ',
       '탐색 단계가 감성 최저점이다. 고객이 문제를 서술하는 시점에 상담사가 즉각 '
       '공감을 표현하느냐 여부가 이후 감정 회복 속도에 결정적 영향을 미친다. '
       '결과제시 단계까지도 감성이 부정(-) 방향인 것은, 해결책이 제시되더라도 고객이 '
       '"결과를 받아들이기까지" 심리적 처리 시간이 필요함을 시사한다. '
       '이 패턴은 추천 고객(NPS 9~10)과 비추천 고객(NPS 1~5) 모두에서 공통적으로 나타나며, '
       '차이는 패턴의 회복 속도와 최저점의 깊이에서 발생한다.')

    doc.add_paragraph()

    # ── 3.3 NPS-감성 상관관계 ────────────────────────────────────────────
    section_heading(doc, '3.3 NPS-감성 상관관계', level=2)

    fig(doc, 'fig5.png',
        'Fig. 6  NPS 그룹별(비추천 1~5 / 중립 6~8 / 추천 9~10) 평균 감성 점수 바 차트 — '
        '비추천 그룹이 추천 그룹 대비 감성 점수가 약 2배 낮게 나타나며, 계단형 차이 관찰.',
        5.5)

    doc.add_paragraph(
        f'NPS 점수와 통화 중 고객 평균 감성 지수 사이의 Pearson 상관계수는 '
        f'r={corr:.3f}로, 유의미한 양의 상관관계가 확인되었다 (N={n_calls}). '
        f'비추천 고객은 추천 고객 대비 감성 점수가 약 2배 낮게 나타났으며, '
        f'이는 실제 NPS와도 일치하는 경향으로 분석의 타당성을 뒷받침한다.'
    )

    tbl(doc,
        ['NPS 구간', '건수', '감성 방향', '추천 그룹 대비', '해석'],
        [
            ['1~5점 (비추천)', '104건', '부정(-) 방향 — 가장 낮음',  '약 2배 낮음', '전반적 부정 감성. 문제 미해결 경향'],
            ['6~8점 (중립)',   '325건', '부정(-) 방향 — 중간',       '중간 수준',   '감정 변동 있으나 최종 수용'],
            ['9~10점 (추천)', '257건', '부정(-) 방향 — 상대적 높음', '기준',        '상대적 긍정. 그러나 과정 중 부정 구간 존재'],
        ]
    )

    doc.add_paragraph()
    bp(doc, '인사이트: ',
       'NPS 9~10점 최고 만족 고객에서도 통화 중 감성이 부정(-) 방향이다. '
       'NPS는 "결과적 만족"을 반영하지만, 통화 과정의 불편함은 별도로 존재한다. '
       '또한 r=0.262는 강한 상관이 아님에 주목해야 한다. 이는 NPS와 음성 감성이 '
       '"동일한 것을 다른 방법으로 측정하는 것"이 아니라, 서로 다른 측면을 측정하는 '
       '독립적인 지표임을 의미한다. NPS가 높아도 통화 중 감성이 낮을 수 있고, '
       'NPS가 낮아도 특정 구간에서 감성이 개선될 수 있다. '
       '이 독립성이 음성 감성 분석의 핵심 가치이다.')

    doc.add_page_break()

    # ── 3.4 감정 전환 패턴 ───────────────────────────────────────────────
    section_heading(doc, '3.4 감정 전환 패턴', level=2)

    doc.add_paragraph(
        f'전체 {n_calls}건의 통화에서 감정 전환이 총 7,877회 발생하였으며, '
        f'통화당 평균 11.5회의 감정 전환이 나타났다. '
        f'이 중 부정→긍정(N→P) 전환은 1,591회(20%), 긍정→부정(P→N) 전환은 1,462회(19%)로, '
        f'상담 중 감정이 빈번하게 양방향으로 변동함을 보여준다. '
        f'특히 탐색→해결시도 단계 구간에서 감정 전환이 가장 빈번하게 발생한다.'
    )

    tbl(doc,
        ['전환 유형', '빈도', '비율', '의미'],
        [
            ['부정→긍정 (N→P)', '1,591회', '20%', '상담사 개입 또는 문제 해결 후 개선'],
            ['긍정→부정 (P→N)', '1,462회', '19%', '새로운 문제 발생 또는 기대 불충족'],
            ['기타 전환',       '4,824회', '61%', '유사 감정 그룹 간 이동'],
            ['통화당 평균',     '11.5회',  '—',   '높은 감정 변동성 — 정적 측정의 한계'],
        ]
    )

    bp(doc, '인사이트: ',
       '통화당 평균 11.5회의 감정 전환은 고객 감정이 매우 역동적임을 보여준다. '
       '"통화 평균 감성"만으로는 이 변동을 포착할 수 없다.')

    doc.add_paragraph()
    section_heading(doc, '3.5 감정 전환 트리거 분석', level=2)
    doc.add_paragraph(
        '686건 전체에서 부정→긍정 전환 1,607건의 직전 상담사 발화를 분석한 결과, '
        '전환을 유발한 상담사 응대 유형은 다음과 같다.'
    )
    tbl(doc,
        ['상담사 응대 유형', '건수', '비율', '대표 발화 예시'],
        [
            ['확인/조회',    '264건', '16%', '"확인해보겠습니다", "조회해볼게요"'],
            ['정보 안내',    '168건', '10%', '"냉장고 뒤쪽 기계실을 열어서..."'],
            ['구체적 솔루션','129건', '8%',  '"기사 방문", "배송", "교환 처리"'],
            ['공감/사과',    '125건', '8%',  '"죄송합니다", "불편하셨겠습니다"'],
            ['감사/인사',    '125건', '8%',  '"감사합니다", "기다려주셔서"'],
        ]
    )
    doc.add_paragraph(
        '전환의 34%가 해결시도 단계에서 발생하며, 이는 상담사가 구체적 대응을 시작하는 시점과 일치한다. '
        '주목할 점은 공감/사과(8%)보다 확인/조회(16%)가 2배 더 효과적이라는 것이다. '
        '"감정적 공감"보다 "문제를 실제로 확인하고 있다는 행동적 신호"가 '
        '고객 감정 회복에 더 크게 기여하는 것으로 해석된다.'
    )

    doc.add_paragraph()
    doc.add_paragraph(
        '반면, 긍정→부정 전환 1,473건에서는 다음 유형의 상담사 발화가 트리거로 작용하였다.'
    )
    tbl(doc,
        ['트리거 유형', '건수', '비율', '해석'],
        [
            ['비용 안내',       '131건', '9%',  '유상 수리, 추가 비용 등 예상과 다른 금액 정보'],
            ['지연/대기 안내',  '123건', '8%',  '배송 지연 가능성, 처리 소요 시간 안내'],
            ['불가/제한 안내',  '59건',  '4%',  '"확인이 안 됩니다", "해당 서비스 불가"'],
            ['확인 요청',       '162건', '11%', '추가 정보 요청 (번호, 주소 등) → 절차 번거로움'],
        ]
    )
    doc.add_paragraph(
        '고객 기대와 다른 정보(비용, 지연, 불가)가 제시되는 시점에서 감정 악화가 집중 발생한다. '
        '특히 해결시도 단계(34%)와 탐색 단계(29%)에서 긍정→부정 전환이 가장 빈번하다. '
        '이는 상담사가 문제를 파악하고 솔루션을 제시하는 과정에서 '
        '고객 기대와의 간극이 드러나기 때문으로 보인다.'
    )

    doc.add_page_break()

    # ══ 4. 고객-상담사 상호작용 패턴 ════════════════════════════════════════
    section_heading(doc, '4. 고객-상담사 상호작용 패턴', level=1)

    fig(doc, 'fig7.png',
        'Fig. 7  (a) 상담사 응대 패턴 분포 — 공감 28%, 선제 20%, 솔루션 12%. '
        '(b) NPS 그룹별 응대율 비교 — 비추천/중립/추천 간 공감률 차이 없음(26~28%).',
        6.0)

    doc.add_paragraph(
        '고객의 평균 발화 길이는 3.1초이며 상담사는 5.9초로, '
        '상담사가 약 1.9배 길게 발화한다. 이는 고객이 짧은 응답과 감탄사 위주로 반응하고, '
        '상담사가 안내·설명 중심으로 발화하는 콜센터 상담의 구조적 특성을 반영한다.'
    )

    tbl(doc,
        ['발화 주체', '평균 발화 길이', '주요 패턴', '특이사항'],
        [
            ['고객',   '3.1초', '"네", "예", "알겠습니다" 등 짧은 호응 44.1%',
             '텍스트 감성 분석 불가 구간이 전체의 절반 수준'],
            ['상담사', '5.9초', '안내·설명 중심 발화',
             '공감 28%, 선제 안내 20%, 솔루션 12%'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        '주목할 점은 NPS 상·하위 그룹 간 상담사의 공감 멘트 비율에 유의미한 차이가 없다는 것이다. '
        '비추천 그룹과 추천 그룹 모두 공감 응대율이 26~28% 수준으로 동일하다.'
    )

    bp(doc, '인사이트: ',
       '상담사 공감 멘트 빈도는 NPS와 무관하다 — 타이밍이 중요하다. '
       '공감 표현을 얼마나 자주 하는지보다, 고객이 실제로 불안을 호소하는 바로 그 시점에 '
       '공감 멘트를 제공했는지가 중요하다. 이를 검증하려면 "부정→긍정 전환 직전 상담사 발화"의 '
       '패턴을 분석해야 하며, 이는 음성 감성 분석의 발화 단위 시간 매핑(Forced Alignment)이 '
       '있어야 가능한 분석이다.')

    doc.add_paragraph()
    bp(doc, '발화 구조적 함의: ',
       '고객의 실제 감정은 긴 서사적 발화보다 짧은 호응(~44%)에서 더 많이 표출된다. '
       '"네 네 네"의 톤이 밝으면 수용, 어두우면 불안을 나타내며, 이는 텍스트만으로는 구분할 수 없다. '
       '음성 분석은 이 "말 사이의 감정"을 포착하는 데 핵심적인 역할을 한다.')

    doc.add_page_break()

    # ══ 5. 고객 세그먼트별 분석 ══════════════════════════════════════════════
    section_heading(doc, '5. 고객 세그먼트별 분석', level=1)

    fig(doc, 'fig6.png',
        'Fig. 8  (a) 연령대별 평균 NPS 및 감성 점수 — 40대 NPS 최저(7.76), 75세 이상 최고(8.31). '
        '(b) 성별 NPS 및 감성 점수 — 남성 NPS 8.19 / 여성 NPS 7.87.',
        6.0)

    section_heading(doc, '5.1 연령대별·성별 분석', level=2)

    doc.add_paragraph(
        '40대 고객이 NPS 가장 낮고(7.76), 75세 이상이 가장 높다(8.31). '
        '한편 음성 감성 점수에서는 65~74세 그룹이 가장 낮은 수준을 기록하였으며, '
        '이는 NPS가 상대적으로 높음에도 불구하고 통화 중 실제 감정은 더 부정적임을 보여주는 '
        '흥미로운 역전 패턴이다. 이 패턴은 실제 NPS 점수와도 일치하는 경향이 있어 '
        '분석의 타당성을 뒷받침한다.'
    )

    tbl(doc,
        ['연령대', '평균 NPS', '감성 방향', '특이사항'],
        [
            ['20~39세',   '8.21', '부정(-) — 상대적 낮음',   'NPS·감성 모두 상위. 빠른 문제 해결 선호'],
            ['40~49세',   '7.76', '부정(-) — 중간',          'NPS 최저. 높은 기대치와 구체적 요구'],
            ['50~64세',   '7.95', '부정(-) — 중간',          '중간 수준. 설명·안내에 민감'],
            ['65~74세',   '8.11', '부정(-) — 가장 낮음',     'NPS 높으나 감성 최저 — 친절함에 만족하나 내면 불안'],
            ['75세 이상', '8.31', '부정(-) — 중간 수준',     'NPS 최고. 상담사 응대 자체에 높은 만족'],
        ]
    )

    doc.add_paragraph()
    bp(doc, '인사이트: ',
       '65~74세 그룹은 NPS와 음성 감성 점수 사이에 가장 큰 괴리가 있다. '
       '이는 고령 고객이 NPS 설문에서는 "상담사가 친절했다"는 인상으로 높은 점수를 주지만, '
       '통화 중 실제 감정에서는 상대적으로 더 많은 불안과 혼란을 경험함을 의미한다. '
       '음성 감성 분석 없이는 이 괴리를 발견하기 어렵다.')

    doc.add_paragraph()
    tbl(doc,
        ['성별', '평균 NPS', '감성 방향', '건수', '특이사항'],
        [
            ['남성', '8.19', '부정(-) — 상대적 낮음', '263건', 'NPS 소폭 높으나 음성 감성은 더 부정적'],
            ['여성', '7.87', '부정(-) — 상대적 높음', '407건', 'NPS 낮으나 음성 감성은 상대적 긍정'],
        ]
    )

    bp(doc, '인사이트: ',
       '남성이 NPS에서 높은 점수를 주지만 통화 중 음성 감성은 더 낮다. '
       '이는 남성이 NPS 설문에서 "결과 중심"으로 평가하는 반면, '
       '여성은 "과정 중심"으로 평가하는 성향 차이를 반영할 수 있다. '
       '또한 통화 건수에서 여성(407건)이 남성(263건)보다 55% 많아, '
       '고객 구성의 성별 불균형도 고려해야 한다.')

    doc.add_page_break()

    section_heading(doc, '5.2 제품군별 분석', level=2)

    fig(doc, 'fig10.png',
        'Fig. 9  제품군별 평균 NPS 바 차트 — TV/AV(8.36), PC(8.35) 높음. '
        '에어컨(7.59) 최저, 주방가전(7.81) 두 번째로 낮음.',
        5.5)

    doc.add_paragraph(
        '에어컨과 주방가전이 만족도 낮다. 제품군별 NPS 격차가 0.77점(에어컨 7.59 vs TV/AV 8.36)으로 '
        '상당하며, 이 격차는 제품 특성보다 상담 해결 구조의 차이에서 기인한다.'
    )

    tbl(doc,
        ['제품군', '평균 NPS', '상담 특성', '낮은/높은 이유'],
        [
            ['에어컨',   '7.59', '계절성 폭주, 설치·A/S 연계',
             '외부 기사 일정 등 상담사 통제 불가 요인 多'],
            ['주방가전', '7.81', '배송·설치·부품 복합',
             '물리적 서비스 연계로 즉시 해결 어려움'],
            ['생활가전', '7.96', '일반 가전 A/S',
             '중간 수준. 다양한 제품 포함'],
            ['PC',       '8.35', '원격 설정·소프트웨어',
             '즉시 해결 가능 비율 높아 만족도 높음'],
            ['TV/AV',    '8.36', '화질·설정·연결 문제',
             '원격 안내로 빠른 해결. 대기 없음'],
        ]
    )

    bp(doc, '인사이트: ',
       '에어컨·주방가전 상담의 낮은 NPS는 상담사 역량이 아닌 구조적 한계에서 비롯된다. '
       '음성 감성 분석에서 이 제품군 상담의 "해결시도" 단계 감성 점수 추이를 별도 분석하면, '
       '고객이 해결 불가 상황을 통보받는 정확한 시점을 특정하여 '
       '"해결 불가 안내 스크립트 개선"에 활용할 수 있다.')

    doc.add_page_break()

    # ══ 6. 상담 품질 지표 종합 ═══════════════════════════════════════════════
    section_heading(doc, '6. 상담 품질 지표 종합', level=1)

    doc.add_paragraph(
        '13개 상담 품질 지표를 LLM 교차검증으로 정량화하였다. '
        '지표는 고객 측 요소, 상담사 측 요소, 상호작용 요소로 구분된다.'
    )

    tbl(doc,
        ['지표명', '평균 점수', '구분', '수준'],
        [
            ['고객_협조도',    '80.3', '고객',    '양호 (60+)'],
            ['설명명확성',     '77.0', '상담사',  '양호 (60+)'],
            ['문제객관성',     '70.3', '고객',    '양호 (60+)'],
            ['문제구체성',     '61.2', '고객',    '양호 (60+)'],
            ['공감표현',       '60.5', '상담사',  '양호 (60+)'],
            ['솔루션구체성',   '59.0', '상담사',  '보통 (40~60)'],
            ['해결진척도',     '52.3', '상호작용', '보통 (40~60)'],
            ['해결의지',       '50.6', '상담사',  '보통 (40~60)'],
            ['감정회복력',     '50.2', '상호작용', '보통 (40~60)'],
            ['다음단계명확성', '42.5', '상담사',  '개선 필요 (40 미만 근접)'],
            ['주도성',         '35.4', '상담사',  '개선 필요 (40 미만)'],
            ['마찰도',         '33.9', '상호작용', '개선 필요 (40 미만)'],
            ['감정강도',       '12.8', '고객',    '주의 (낮을수록 좋음 — 해석 주의)'],
        ]
    )

    doc.add_paragraph()
    bp(doc, '핵심 개선 포인트: ',
       '상담사의 "주도성"(35.4)과 "다음단계 명확성"(42.5)이 낮다. '
       '이는 상담사가 고객의 요청에 반응적으로 대응하는 경향이 있으며, '
       '"다음에 어떤 절차가 진행되는지"를 명확히 안내하는 부분에서 개선 여지가 있다. '
       '음성 감성 분석에서 상담사의 선제적 안내 직후 긍정 전환이 발생하는 패턴이 '
       '확인되었으므로, "다음 단계 선제 안내" 스크립트 강화가 효과적일 것으로 판단된다.')

    bp(doc, '감정 강도 12.8점의 해석: ',
       '고객의 감정 강도가 낮다는 것은 대부분의 상담이 감정적으로 격앙되지 않은 상태에서 '
       '진행된다는 의미이다. 역설적으로 이는 "미세한 감정 변화"를 포착하는 것이 더 중요해짐을 '
       '의미한다. 격앙된 고객의 불만은 텍스트만으로도 감지되지만, '
       '감정 강도가 낮은 고객의 은밀한 불안은 음성 프로소디 분석으로만 포착 가능하다.')

    doc.add_page_break()

    # ══ 7. 대표 케이스 분석 ══════════════════════════════════════════════════
    section_heading(doc, '7. 대표 케이스 분석', level=1)

    doc.add_paragraph(
        '아래는 음성 분석이 가장 큰 차이를 만드는 3가지 유형의 대표 사례이다. '
        '개인정보 보호를 위해 통화 식별 정보는 일체 포함하지 않는다.'
    )

    section_heading(doc, '7.1 동일한 짧은 응답, 다른 감정', level=2)
    doc.add_paragraph(
        '텍스트로는 모두 "네 네"이지만, 음성 톤에 따라 전혀 다른 감정으로 판정된다. '
        '이는 텍스트 분석의 근본적 한계이며, 음성 분석만이 제공할 수 있는 고유한 정보이다.'
    )
    tbl(doc,
        ['시점', '발화', '음성 특징 (baseline 대비)', '판정', '근거'],
        [
            ['45.1초',  '"네 네"',    'F0 안정, 에너지 평소 수준',     '안정/중립 — 긍정(+) 방향', '차분한 수락'],
            ['31.3초',  '"네 네 네"', 'F0 불안정, 유성구간 44%↓',      '불안/걱정 — 부정(-) 방향', '말이 끊기는 패턴'],
            ['129.5초', '"네 네"',    '에너지 급감(부정(-)), ZCR 증가', '불안/걱정 — 부정(-) 방향', 'baseline 대비 위축'],
        ]
    )

    section_heading(doc, '7.2 형식적 표현 뒤의 숨은 감정', level=2)
    doc.add_paragraph(
        '"수고하세요"를 텍스트로만 보면 중립이지만, '
        '음성에서 baseline 대비 에너지가 크게 떨어지면 내재된 불편함을 시사한다.'
    )
    tbl(doc,
        ['시점', '발화', 'STT Only', 'STT+Audio', '음성 근거'],
        [
            ['135.6초', '마무리 인사', '(분석불가)', '불안/걱정 — 부정(-) 방향',
             '에너지 급감, F0 불안정\n→ 형식적이나 내재된 불편함'],
            ['33.7초',  '이름 소개',  '감사/만족 — 긍정(+) 방향', '안정/중립 — 중립(0)',
             'baseline 대비 변화 없음\n→ BERT 오판을 음성이 보정'],
        ]
    )

    section_heading(doc, '7.3 텍스트-음성 충돌 시 LLM 교차검증', level=2)
    doc.add_paragraph(
        'BERT와 음성이 다른 방향을 가리킬 때, LLM 교차검증이 맥락을 고려하여 최종 판정을 내린다.'
    )
    tbl(doc,
        ['발화', 'BERT 판정', '음성 감성 방향', 'LLM 최종', '판정 근거'],
        [
            ['배송 일정 확인\n"그분이 전화 오신다고요"',
             '혼란/당황', '긍정(+) 방향', '혼란/당황 — 부정(-) 방향',
             '반복적 "네" + 질문 패턴\n→ BERT의 당황 판정 우선'],
            ['시간 희망 표현\n"오전에 올 수 있으면..."',
             '안정/중립', '중립(0) 수준', '안정/중립 — 긍정(+) 방향',
             '양쪽 일치\n→ 조심스러운 요청으로 판정'],
        ]
    )

    doc.add_page_break()

    # ══ 8. 결론 및 시사점 ════════════════════════════════════════════════════
    section_heading(doc, '8. 결론 및 시사점', level=1)

    section_heading(doc, '8.1 핵심 발견 요약', level=2)

    doc.add_paragraph(
        f'686건 상담 음성의 전수 분석을 통해 STT 텍스트 분석만으로는 고객 감정의 '
        f'상당 부분을 포착할 수 없음이 정량적으로 확인되었다. '
        f'전체 고객 발화의 {short_pct:.0%}({short:,}건)가 텍스트 분석 불가 구간이며, '
        f'의미 있는 문장에서도 {changed_pct:.0%}({changed:,}건)의 판정이 '
        f'음성 추가 시 변경된다. 이 두 수치를 합산하면, STT 텍스트 분석만으로는 '
        f'전체 고객 발화의 약 60% 이상에서 감정 정보를 놓치거나 오판하는 셈이다.'
    )

    doc.add_paragraph(
        '발견된 핵심 패턴을 정리하면 다음과 같다. 첫째, 탐색 단계에서 고객 감성이 '
        '가장 부정적이며, 종료 단계에서 회복되는 패턴이 관찰된다. '
        '고객이 문제를 서술하는 탐색 단계가 상담사 개입의 골든 타임임을 확인하였다. '
        '둘째, NPS와 통화 중 감성의 상관(r=0.262)이 중간 수준으로, '
        '두 지표는 서로 다른 측면을 측정하는 독립적 정보임이 검증되었다. '
        '셋째, NPS 9~10점 최고 만족 고객에서도 통화 중 감성이 부정(-) 방향이며, '
        '특히 NPS 고점 상담의 개별 궤적을 보면 부정 구간이 반복적으로 출현한다. '
        '이는 "결과 만족"과 "과정 감정"이 분리됨을 뜻한다.'
    )

    doc.add_paragraph(
        '넷째, 연령대별로는 40대 고객의 NPS가 7.76으로 가장 낮고, '
        '제품군별로는 에어컨(7.59)과 주방가전(7.81)이 상대적으로 낮은 만족도를 보인다. '
        '이는 해당 세그먼트에서 상담 과정이나 서비스 구조에 개선 여지가 있음을 시사한다.'
    )
    doc.add_paragraph(
        '다섯째, 상담사의 공감 응대율은 NPS 그룹 간 차이가 거의 없다 '
        '(비추천 26%, 중립 28%, 추천 28%). '
        '공감 멘트를 "얼마나 자주" 하느냐보다, 고객이 불안을 호소하는 바로 그 시점에 '
        '적절히 공감하느냐가 만족도에 더 큰 영향을 미친다는 해석이 가능하다.'
    )
    doc.add_paragraph(
        '여섯째, 통화당 평균 11.5회의 감정 전환이 발생하며, 이는 상담 중 고객 감정이 '
        '매우 역동적으로 변화함을 보여준다. 단순히 통화 전체의 평균 감성만으로는 '
        '이러한 변동을 포착할 수 없으므로, 발화 단위의 실시간 감성 추적이 필요하다.'
    )

    doc.add_paragraph()

    section_heading(doc, '8.2 다차원 인사이트', level=2)

    bp(doc, '인사이트 1 — 측정의 사각지대 해소: ',
       f'콜센터 품질 관리에서 가장 큰 사각지대는 NPS 미응답 통화(전체의 60~70%)이다. '
       f'음성 감성 분석은 모든 통화에 대해 자동화된 감성 점수를 산출하므로, '
       f'NPS 설문 응답 여부와 무관하게 전체 통화의 품질을 추적할 수 있다. '
       f'또한 r=0.262의 상관관계를 활용하면 음성 감성으로부터 NPS를 역산 추정하는 '
       f'모델 개발이 가능하며, 이는 추후 회귀 분석을 통해 구체화될 수 있다.')

    doc.add_paragraph()
    bp(doc, '인사이트 2 — 실시간 개입 트리거: ',
       '통화 중 부정 감성이 특정 임계값 이하로 지속될 때(예: 감성 점수가 부정(-) 방향 임계 이하이고 '
       '30초 이상 지속) 관리자에게 알림을 보내는 실시간 개입 시스템을 구축할 수 있다. '
       '탐색 단계에서의 감성 급락이 이후 NPS 저하와 연관되는지를 분석하여 '
       '고위험 통화를 조기에 식별하는 것이 목표이다.')

    doc.add_paragraph()
    bp(doc, '인사이트 3 — 감정 전환 매핑 기반 코칭: ',
       '1,591건의 부정→긍정 전환 직전 상담사 발화를 분류하면, '
       '"어떤 유형의 발화가 감정 개선을 유도하는가"에 대한 데이터 기반 답변이 가능하다. '
       '현재 분석에서 선제적 안내(20%)가 공감(28%)보다 적게 사용되지만, '
       '단계별 감성 개선 효과는 선제 안내 이후 구간에서 더 뚜렷하게 나타나는 경향이 있다. '
       '이를 검증하여 "선제 안내 타이밍 최적화" 코칭 프로그램을 개발할 수 있다.')

    doc.add_paragraph()
    bp(doc, '인사이트 4 — 제품군별 맞춤 스크립트: ',
       '에어컨(NPS 7.59)과 TV/AV(NPS 8.36)의 0.77점 격차는 제품의 품질 차이가 아닌 '
       '상담 구조의 차이에서 비롯된다. 에어컨 상담에서 해결 불가 상황을 '
       '"어떻게 안내하느냐"가 NPS에 직접 영향을 미친다. '
       '음성 감성 분석으로 "해결 불가 통보 시점의 감성 하락 폭"을 제품군별로 비교하면, '
       '어느 제품군의 안내 스크립트 개선이 가장 시급한지를 우선순위화할 수 있다.')

    doc.add_paragraph()
    bp(doc, '인사이트 5 — 고령 고객 UX 특수성: ',
       '65~74세 그룹의 NPS-감성 괴리(NPS 높음 vs 감성 점수 낮음)는 이 연령대가 '
       '"상담사에 대한 예의"로 높은 NPS를 주면서도 내면적으로 더 많은 불안과 혼란을 '
       '경험함을 시사한다. 이 그룹의 불안/걱정 발화 빈도와 혼란/당황 발화 빈도를 분리 분석하여, '
       '고령 고객 특화 상담 프로토콜(더 느린 설명 속도, 단계별 확인 질문 추가 등)의 '
       '효과를 측정할 수 있다.')

    doc.add_paragraph()

    section_heading(doc, '8.3 비즈니스 적용 방향', level=2)

    doc.add_paragraph(
        '분석 결과를 실무에 적용하는 우선순위를 단기·중기·장기로 구분한다.'
    )

    tbl(doc,
        ['시기', '적용 방향', '기대 효과', '필요 조건'],
        [
            ['단기\n(0~3개월)',
             '전체 통화 자동 감성 스코어링\n+ 품질 대시보드 구축',
             'NPS 미응답 70% 통화 품질 가시화\n일별·상담사별 감성 트렌드 모니터링',
             '현재 파이프라인 운영화\n처리 속도 최적화 (현재 15~20초/콜)'],
            ['단기\n(0~3개월)',
             '탐색 단계 집중 공감 교육\n("골든 타임" 개념 도입)',
             '탐색 단계 감성 점수 개선\n부정→긍정 전환 속도 향상',
             '탐색 단계 자동 감지 기능\n상담사 교육 자료 업데이트'],
            ['중기\n(3~6개월)',
             '부정→긍정 전환 직전 발화 분석\n→ 코칭 포인트 데이터화',
             '"효과적 발화 패턴" 교과서 개발\n상담사 개인별 맞춤 피드백',
             '1,591건 전환 사례 수동 레이블링\n상담사 식별 데이터 연계'],
            ['중기\n(3~6개월)',
             '제품군별 해결 불가 안내 스크립트\n음성 감성 A/B 테스트',
             '에어컨·주방가전 NPS 0.3~0.5점 개선\n고객 불안 표현 감소',
             '스크립트 버전 관리 시스템\n그룹별 통화 무작위 배정'],
            ['장기\n(6개월+)',
             'NPS 예측 모델 개발\n(감성 점수 + 단계 패턴 → NPS)',
             'NPS 미응답 통화 품질 추정\n고위험 통화 사전 예측',
             '충분한 학습 데이터 (2,000건+)\nMLOps 인프라 구축'],
            ['장기\n(6개월+)',
             '실시간 개입 알림 시스템\n(감성 임계값 기반)',
             '고위험 통화 에스컬레이션 자동화\nVOC 사전 차단',
             'Stereo 녹음 또는\nSpeaker Diarization 도입'],
        ]
    )

    doc.add_paragraph()

    section_heading(doc, '8.4 한계 및 개선 방향', level=2)

    doc.add_paragraph(
        '본 분석은 여러 가지 기술적·방법론적 한계를 내포한다. '
        '이를 명확히 인식하고 향후 개선 방향을 제시한다.'
    )

    tbl(doc,
        ['영역', '현재 상태', '한계', '개선 방향'],
        [
            ['음성 감정 모델',
             '룰 기반 휴리스틱\n(F0+Energy+ZCR 등)',
             '분노 vs 기쁨 구분 어려움\n(유사 음향 패턴)',
             'Pretrained SER 모델\n(wav2vec2 / HuBERT)'],
            ['BERT 도메인',
             'GoEmotions\n(Reddit 댓글)',
             '"비통한" 등 상담 무관\n라벨 혼입',
             '상담 데이터 Fine-tuning\n(레이블 500건+)'],
            ['Mono 오디오',
             '단일 채널 녹음',
             '상담사/고객 음성 혼합\n발화 경계 ±1초 오차',
             'Stereo 녹음 또는\nSpeaker Diarization'],
            ['LLM 비용·일관성',
             'LLM\n콜당 ~$0.02',
             '대규모 처리 시 비용\n판정 일관성 변동',
             'temperature=0 고정\n판정 캐싱 전략'],
            ['통계적 유의성',
             'N=686\n단일 기간 데이터',
             '시즌·캠페인 효과 미분리\n제품군별 소표본',
             '분기별 반복 측정\n층화 샘플링 설계'],
            ['인과관계',
             '상관관계 분석\n(r=0.262)',
             '인과 방향 미확인\n혼재 변수 존재',
             '준실험 설계\n상담사별 A/B 테스트'],
        ]
    )

    doc.add_paragraph()

    section_heading(doc, '8.5 방법론적 고려사항', level=2)

    doc.add_paragraph(
        'Baseline 보정은 분석 신뢰도를 높이는 핵심 장치이지만, '
        '통화 길이가 짧거나(60초 미만) 발화 수가 적은(5건 미만) 경우 '
        'Baseline 자체가 불안정해져 보정 효과가 제한된다. '
        f'본 분석에서 {short_pct:.0%}의 짧은 응답이 포함된 통화에서 이 문제가 발생할 수 있으며, '
        '이를 해결하기 위해 최소 발화 수 미만 통화에 대해서는 '
        '별도의 가중치 보정 또는 제외 기준을 적용할 필요가 있다.'
    )

    doc.add_paragraph(
        'LLM 교차검증에서 LLM의 판정은 temperature=0으로 설정하더라도 '
        '완전한 결정론적 출력을 보장하지 않는다. 동일한 입력에 대해 판정이 달라질 경우를 '
        '대비하여, 중요 통화(NPS 비추천 그룹 등)에 대해서는 판정 캐싱과 감사 로그 유지가 '
        '권장된다. 또한 LLM의 "자연어 근거"는 정성적 참고 자료로 활용되어야 하며, '
        '단독 의사결정 근거로 사용해서는 안 된다.'
    )

    doc.add_paragraph(
        f'NPS와 음성 감성 지수 간 상관(r={corr:.3f})의 해석에 주의가 필요하다. '
        '설명분산(r²)은 약 6.9%로, NPS 변동의 93.1%는 음성 감성 이외의 요인에 의해 결정된다. '
        '따라서 음성 감성 단독으로 NPS를 대체하거나 예측하려는 시도는 과도한 확대 해석이다. '
        '음성 감성 분석의 가치는 NPS 대체가 아닌, NPS가 측정하지 못하는 "과정의 감정"을 '
        '독립적으로 보완하는 데 있다.'
    )

    doc.add_paragraph()

    section_heading(doc, '8.6 종합 결론', level=2)

    doc.add_paragraph(
        '본 분석은 콜센터 상담 품질 측정에서 음성 프로소디 분석이 텍스트 분석의 '
        '구조적 한계를 보완하는 독립적 가치를 지님을 686건의 실데이터로 검증하였다. '
        '핵심 메시지는 세 가지이다.'
    )

    doc.add_paragraph(
        '첫째, 텍스트 분석만으로는 전체 고객 감정의 60% 이상을 포착하지 못한다. '
        f'{short_pct:.0%}의 짧은 응답 사각지대와 {changed_pct:.0%}의 오판·누락을 합산하면, '
        '텍스트 단독 분석은 감정 정보의 절반 이상을 놓친다. '
        '이는 기존 텍스트 기반 품질 관리 시스템의 측정 타당성 문제를 제기한다.'
    )

    doc.add_paragraph(
        '둘째, 음성 감성 지수는 NPS를 대체하는 지표가 아니라, NPS가 포착하지 못하는 '
        '영역을 보완하는 지표이다. 두 지표 간 상관(r=0.262)이 높지 않은 것은, '
        '측정하는 대상 자체가 다르기 때문이다. NPS는 통화 종료 후 "결과에 대한 회고적 평가"이고, '
        '음성 감성은 "통화 과정에서의 실시간 감정 변화"를 측정한다. '
        '두 지표를 함께 활용하면, '
        '"NPS는 높지만 통화 과정에서 반복적으로 부정적 구간이 있었던 통화"와 '
        '"NPS는 낮지만 상담사 응대로 감정이 점차 회복된 통화"를 '
        '구분하여 각각 다른 후속 조치를 취할 수 있다.'
    )

    doc.add_paragraph(
        '셋째, 감성 분석의 진짜 가치는 "평균 점수"가 아닌 "변화 패턴"에 있다. '
        '통화당 평균 11.5회의 감정 전환은 상담 감정이 매우 역동적임을 보여준다. '
        '탐색 단계의 감성 최저점, 부정→긍정 전환 타이밍, '
        '상담 단계별 감성 궤적 등의 패턴 정보가 상담사 코칭과 서비스 개선에 실질적 근거를 제공한다.'
    )

    doc.add_paragraph(
        '분석 방법론의 완성도 측면에서 Baseline 보정, Forced Alignment 기반 '
        '발화 단위 시간 매핑, LLM 교차검증의 결합은 현재 시점에서 실용적으로 구현 가능한 '
        '최선의 아키텍처로 판단된다. 향후 Stereo 녹음 도입과 SER(Speech Emotion Recognition) '
        '전용 딥러닝 모델 적용으로 정확도를 추가 향상시킬 수 있다.'
    )

    # ── 저장 ─────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f'보고서 저장: {output_path}')
    return output_path


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print('=== 음성 감성 분석 보고서 생성 시작 ===')

    # 1. 임시 디렉토리에 피겨 생성
    fig_dir = tempfile.mkdtemp(prefix='voice_report_figs_')
    print(f'\n[1/2] 피겨 생성 중 (임시 디렉토리: {fig_dir})')
    generate_figures(fig_dir)

    # 2. Word 보고서 생성
    print(f'\n[2/2] Word 보고서 생성 중...')
    generate_report(fig_dir, OUTPUT_DOCX)

    print(f'\n=== 완료 ===')
    print(f'출력 파일: {OUTPUT_DOCX}')
